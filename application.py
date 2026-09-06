# ============================================================
# VISION EDITION — streamlit_app.py
# Version 2.0 — Multi-dossiers / Accès dirigeant / Rôles
# © 2025 Nicolas CUISSET — Mémoire d'expertise comptable
# ============================================================

import streamlit as st
import pandas as pd
import numpy as np
import json
import hashlib
import datetime as _dt
from io import BytesIO
import plotly.express as px
import plotly.graph_objects as go
import qrcode
from PIL import Image
import anthropic
import smtplib
import openpyxl
from openpyxl.styles import PatternFill as _OPFill, Font as _OPFont, Alignment as _OPAlign, Border as _OPBorder, Side as _OPSide
from openpyxl.worksheet.datavalidation import DataValidation as _OPDV
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

APP_URL = "https://outilaccompagnementmaisonsedition-lyvgltfbwtqo4m9tdmzofu.streamlit.app/"

# ============================================================
# HELPERS FORMAT FR
# ============================================================
def fmt_fr(x, decimals=0):
    try:
        s = f"{float(x):,.{decimals}f}"
    except (ValueError, TypeError):
        return str(x)
    return s.replace(",", "§").replace(".", ",").replace("§", " ")

# ============================================================
# GESTION DES UTILISATEURS ET DOSSIERS (JSON session)
# ============================================================
# Structure en session_state["cabinet"] :
# {
#   "ec_users": { "login": {"pw_hash": ..., "name": ..., "role": "ec"} },
#   "dossiers": {
#     "id_dossier": {
#       "nom": "Editions de l'Argonaute",
#       "dirigeant": "Thomas Bernard",
#       "exercice": "2025/2026",
#       "param_comptes": {...},
#       "mapping": {...},
#       "labels_indirect": {...},
#       "dirigeant_users": { "login": {"pw_hash": ..., "name": ...} }
#     }
#   }
# }

def init_cabinet():
    """Initialise la structure cabinet si absente."""
    if "cabinet" not in st.session_state:
        st.session_state["cabinet"] = {
            "ec_users": {
                "aurore": {
                    "pw_hash": _hash("12345"),
                    "name": "Aurore Demoulin",
                    "role": "ec"
                }
            },
            "dossiers": {}
        }

def _hash(pw):
    return hashlib.sha256(pw.encode()).hexdigest()

def verifier_login(username, password, role_attendu=None):
    """Vérifie un login EC ou dirigeant selon le dossier actif."""
    init_cabinet()
    cab = st.session_state["cabinet"]
    # Vérif EC
    ec = cab["ec_users"].get(username)
    if ec and ec["pw_hash"] == _hash(password):
        if role_attendu in (None, "ec"):
            return {"ok": True, "role": "ec", "name": ec["name"], "dossier_id": None}
    # Vérif dirigeant dans tous les dossiers
    for did, dos in cab["dossiers"].items():
        du = dos.get("dirigeant_users", {}).get(username)
        if du and du["pw_hash"] == _hash(password):
            return {"ok": True, "role": "dirigeant", "name": du["name"],
                    "dossier_id": did, "dossier_nom": dos["nom"]}
    return {"ok": False}


# ============================================================
# SYSTÈME DE NOTIFICATIONS
# ============================================================

def envoyer_email(destinataire, sujet, corps, nom_cabinet=NOM_CABINET):
    """Envoie un email via Gmail SMTP. Requiert GMAIL_ADDRESS et
    GMAIL_APP_PASSWORD dans les secrets Streamlit."""
    try:
        gmail_addr = st.secrets.get("GMAIL_ADDRESS", "")
        gmail_pw   = st.secrets.get("GMAIL_APP_PASSWORD", "")
        if not gmail_addr or not gmail_pw:
            return False, "Identifiants Gmail non configurés dans les secrets Streamlit."

        msg = MIMEMultipart("alternative")
        msg["Subject"] = f"[VISION EDITION] {sujet}"
        msg["From"]    = f"{nom_cabinet} <{gmail_addr}>"
        msg["To"]      = destinataire

        corps_html = f"""
        <html><body style="font-family:Arial,sans-serif;max-width:600px;margin:auto;padding:20px">
        <div style="background:#1F4E79;padding:20px;border-radius:8px 8px 0 0">
            <h2 style="color:white;margin:0">📚 VISION EDITION</h2>
            <p style="color:#D6E4F0;margin:5px 0 0 0;font-size:13px">{nom_cabinet}</p>
        </div>
        <div style="background:#f8f9fa;padding:24px;border:1px solid #e0e0e0;border-radius:0 0 8px 8px">
            {corps.replace(chr(10), '<br>')}
            <hr style="margin:20px 0;border:none;border-top:1px solid #e0e0e0">
            <p style="color:#888;font-size:12px">
                Ce message a été envoyé automatiquement depuis VISION EDITION.<br>
                Pour consulter votre tableau de bord, connectez-vous à l'application.
            </p>
        </div>
        </body></html>
        """
        msg.attach(MIMEText(corps, "plain", "utf-8"))
        msg.attach(MIMEText(corps_html, "html", "utf-8"))

        with smtplib.SMTP_SSL("smtp.gmail.com", 465) as smtp:
            smtp.login(gmail_addr, gmail_pw)
            smtp.sendmail(gmail_addr, destinataire, msg.as_string())
        return True, "Email envoyé avec succès."
    except smtplib.SMTPAuthenticationError:
        return False, "Erreur d'authentification Gmail. Vérifiez GMAIL_APP_PASSWORD dans les secrets."
    except smtplib.SMTPException as e:
        return False, f"Erreur SMTP : {e}"
    except Exception as e:
        return False, f"Erreur : {e}"

def ajouter_notification(dossier_id, message, type_notif="info"):
    """Ajoute une notification in-app dans le dossier."""
    dos = get_dossier(dossier_id)
    if "notifications" not in dos:
        dos["notifications"] = []
    dos["notifications"].append({
        "message": message,
        "type": type_notif,
        "date": _dt.datetime.now().strftime("%d/%m/%Y à %Hh%M"),
        "lu": False
    })
    save_dossier(dossier_id, dos)

def marquer_notifications_lues(dossier_id):
    """Marque toutes les notifications comme lues."""
    dos = get_dossier(dossier_id)
    for n in dos.get("notifications", []):
        n["lu"] = True
    save_dossier(dossier_id, dos)

def nb_notifs_non_lues(dossier_id):
    dos = get_dossier(dossier_id)
    return sum(1 for n in dos.get("notifications", []) if not n["lu"])

def get_dossier(dossier_id):
    init_cabinet()
    return st.session_state["cabinet"]["dossiers"].get(dossier_id, {})

def save_dossier(dossier_id, data):
    init_cabinet()
    st.session_state["cabinet"]["dossiers"][dossier_id] = data

# ============================================================
# CONFIG PAGE
# ============================================================
# ============================================================
# IDENTITÉ DU CABINET — modifier ici pour personnaliser
# ============================================================
NOM_CABINET = "CAB ÉDITION"
SLOGAN_CABINET = "Expert-comptable · Maisons d'édition indépendantes"

st.set_page_config(
    page_title=f"VISION EDITION — {NOM_CABINET}",
    page_icon="📚",
    layout="wide"
)

# ============================================================
# INITIALISATION SESSION
# ============================================================
init_cabinet()
for k, v in [
    ("login", False), ("role", None), ("username", ""), ("name", ""),
    ("dossier_id", None), ("messages_agent", []), ("mode_anonyme", False)
]:
    if k not in st.session_state:
        st.session_state[k] = v

# ============================================================
# AUTHENTIFICATION
# ============================================================
if not st.session_state["login"]:
    col1, col2, col3 = st.columns([1, 1.2, 1])
    with col2:
        st.markdown(f"## 📚 VISION EDITION")
        st.markdown(f"**{NOM_CABINET}** — *{SLOGAN_CABINET}*")
        st.divider()
        username_input = st.text_input("Identifiant", placeholder="Votre identifiant")
        password_input = st.text_input("Mot de passe", type="password")
        if st.button("🔑 Connexion", use_container_width=True, type="primary"):
            res = verifier_login(username_input, password_input)
            if res["ok"]:
                st.session_state["login"] = True
                st.session_state["role"] = res["role"]
                st.session_state["username"] = username_input
                st.session_state["name"] = res["name"]
                st.session_state["dossier_id"] = res.get("dossier_id")
                st.rerun()
            else:
                st.error("❌ Identifiants incorrects")
        st.divider()
        st.markdown("**Accès mobile**")
        qr = qrcode.QRCode(version=1, box_size=4, border=2)
        qr.add_data(APP_URL)
        qr.make(fit=True)
        img_qr = qr.make_image(fill_color="black", back_color="white")
        buf = BytesIO(); img_qr.save(buf, format="PNG")
        st.image(buf.getvalue(), width=120)
    st.stop()

# ============================================================
# SIDEBAR SELON RÔLE
# ============================================================
role = st.session_state["role"]

with st.sidebar:
    st.markdown(f"**{NOM_CABINET}**")
    st.caption(SLOGAN_CABINET)
    st.divider()
    st.markdown(f"👤 **{st.session_state['name']}**")
    if role == "ec":
        st.caption("Expert-comptable")
    elif role == "dirigeant":
        dossier_actif = get_dossier(st.session_state["dossier_id"])
        st.caption(f"Dirigeant — {dossier_actif.get('nom', '')}")
        nb_nl = nb_notifs_non_lues(st.session_state["dossier_id"])
        if nb_nl > 0:
            st.warning(f"🔔 {nb_nl} message(s) non lu(s) de votre cabinet")
    st.divider()

    if role == "ec":
        # Sélection dossier
        cab = st.session_state["cabinet"]
        dossiers = cab["dossiers"]
        if dossiers:
            options_dos = {did: d["nom"] for did, d in dossiers.items()}
            did_sel = st.selectbox(
                "📁 Dossier client",
                options=list(options_dos.keys()),
                format_func=lambda x: options_dos[x],
                key="sidebar_dossier_sel"
            )
            if st.session_state["dossier_id"] != did_sel:
                st.session_state["dossier_id"] = did_sel
                # Réinitialiser les données du dossier précédent
                for k in ["df_pivot", "df_pivot_brut", "df_comptables", "param_comptes",
                           "df_source_mappe", "repartition_active"]:
                    st.session_state.pop(k, None)
                st.rerun()
        else:
            st.info("Aucun dossier. Créez-en un.")
            st.session_state["dossier_id"] = None

        pages_ec = [
            "🏠 Accueil",
            "👥 Gestion des dossiers",
            "📂 Import des données",
            "⚙️ Paramétrage analytique",
            "📈 Tableau de bord éditorial",
            "📖 Analyse par titre",
            "🎯 Simulateur de rentabilité",
            "💰 Trésorerie prévisionnelle",
            "✍️ Droits d'auteurs",
            "📦 Retours & Remises",
            "📊 Synthèse financière",
            "🤖 Assistant IA",
            "📄 Rapport de pilotage",
            "📋 Plan d'action"
        ]
        page = st.selectbox("Navigation", pages_ec)
        st.divider()
        st.session_state["mode_anonyme"] = st.checkbox(
            "🕶️ Mode démonstration", value=st.session_state.get("mode_anonyme", False)
        )

    elif role == "dirigeant":
        pages_dir = [
            "🏠 Mon tableau de bord",
            "📖 Mes titres",
            "💰 Ma trésorerie",
            "📦 Retours & Remises",
            "✍️ Droits d'auteurs",
            "🤖 Mon assistant IA"
        ]
        page = st.selectbox("Navigation", pages_dir)

    st.divider()
    if st.button("🚪 Déconnexion", use_container_width=True):
        for k in list(st.session_state.keys()):
            del st.session_state[k]
        st.rerun()

# ============================================================
# HELPERS ANALYTIQUE (identiques à v1 — extraits clés)
# ============================================================
COMPTE_CHARGES_INDIRECTES_REPARTIES = "CHARGES INDIRECTES REPARTIES"
COMPTE_PRODUITS_INDIRECTS_REPARTIS  = "PRODUITS INDIRECTS REPARTIS"

def filtrer_isbn_reels(df):
    if "Code_Analytique" not in df.columns:
        return df.iloc[0:0]
    label_ci = st.session_state.get("labels_indirect", {}).get("charges", "CHARGES INDIRECTES")
    label_pi = st.session_state.get("labels_indirect", {}).get("produits", "PRODUITS INDIRECTS")
    labels_exclus = {label_ci.upper(), label_pi.upper(), ""}
    code = df["Code_Analytique"].astype(str)
    mask = (~code.str.upper().isin(labels_exclus)) & (code.str.strip() != "")
    if "Famille_Analytique" in df.columns and df["Famille_Analytique"].astype(str).str.upper().eq("EDITION").any():
        mask = mask & (df["Famille_Analytique"].astype(str).str.upper() == "EDITION")
    return df[mask]

def mask_retours(df_scope, params):
    prefixes_retours = tuple(params.get("retours") or [])
    if not prefixes_retours:
        return pd.Series(False, index=df_scope.index)
    mask = df_scope["Compte"].astype(str).str.startswith(prefixes_retours)
    prefixes_remises = tuple(params.get("remises") or [])
    if prefixes_remises:
        mask = mask & (~df_scope["Compte"].astype(str).str.startswith(prefixes_remises))
    return mask

def mask_remises(df_scope, params):
    prefixes_remises = tuple(params.get("remises") or [])
    if not prefixes_remises:
        return pd.Series(False, index=df_scope.index)
    return df_scope["Compte"].astype(str).str.startswith(prefixes_remises)

def mask_provisions_reprises(df_scope, params):
    prefixes = tuple(params.get("provisions_reprises") or [])
    if not prefixes:
        return pd.Series(False, index=df_scope.index)
    return df_scope["Compte"].astype(str).str.startswith(prefixes)

def mask_ventes(df_scope, params):
    prefixes_ventes = tuple(params.get("ventes") or [])
    mask_conf = (df_scope["Compte"].astype(str).str.startswith(prefixes_ventes)
                 if prefixes_ventes else pd.Series(False, index=df_scope.index))
    mask_rep  = df_scope["Compte"].astype(str) == COMPTE_PRODUITS_INDIRECTS_REPARTIS
    is_7 = df_scope["Compte"].astype(str).str.startswith("7")
    mask_autres = (is_7 & (~mask_retours(df_scope, params))
                   & (~mask_remises(df_scope, params))
                   & (~mask_provisions_reprises(df_scope, params)))
    return mask_conf | mask_rep | mask_autres

def mask_charges(df_scope, params):
    prefixes = tuple(params.get("charges") or [])
    mask = (df_scope["Compte"].astype(str).str.startswith(prefixes)
            if prefixes else pd.Series(False, index=df_scope.index))
    return mask | (df_scope["Compte"].astype(str) == COMPTE_CHARGES_INDIRECTES_REPARTIES)

def normaliser_codes_isbn(df, col="Code_Analytique"):
    if col not in df.columns:
        return df
    df = df.copy()
    codes = df[col].astype(str)
    label_ci = st.session_state.get("labels_indirect", {}).get("charges", "CHARGES INDIRECTES")
    label_pi = st.session_state.get("labels_indirect", {}).get("produits", "PRODUITS INDIRECTS")
    labels_reserves = {label_ci.upper(), label_pi.upper(), "", "NAN"}
    mask_ean = (~codes.str.strip().str.upper().isin(labels_reserves)) & codes.str.contains(" - ", regex=False)
    if not mask_ean.any():
        return df
    ean_num = codes[mask_ean].str.split(" - ").str[0].str.strip()
    canonique = codes[mask_ean].groupby(ean_num).agg(lambda s: max(s.unique(), key=len))
    df.loc[mask_ean, col] = ean_num.map(canonique)
    return df

def obtenir_mapping_anonymisation(df):
    titres = sorted(filtrer_isbn_reels(df)["Code_Analytique"].astype(str).unique().tolist())
    cache = st.session_state.get("anonymisation_cache")
    if cache and cache.get("titres_source") == titres:
        return cache["mapping"]
    mapping = {t: f"T{i+1}" for i, t in enumerate(titres)}
    st.session_state["anonymisation_cache"] = {"titres_source": titres, "mapping": mapping}
    return mapping

def label_affiche(code, df_pour_mapping=None):
    if not st.session_state.get("mode_anonyme"):
        return code
    mapping = st.session_state.get("anonymisation_cache", {}).get("mapping")
    if mapping is None and df_pour_mapping is not None:
        mapping = obtenir_mapping_anonymisation(df_pour_mapping)
    return (mapping or {}).get(code, code)

def calculer_indicateurs_titres(df, params, titres):
    df_i = df[df["Code_Analytique"].isin(titres)]
    def par_compte(prefix_list, col, exclude=None):
        if not prefix_list:
            return pd.Series(0.0, index=titres)
        mask = df_i["Compte"].astype(str).str.startswith(tuple(prefix_list))
        if exclude:
            mask = mask & (~df_i["Compte"].astype(str).str.startswith(tuple(exclude)))
        return df_i[mask].groupby("Code_Analytique")[col].sum().reindex(titres, fill_value=0.0)

    ventes  = par_compte(params["ventes"], "Crédit")
    ventes_d = par_compte(params.get("ventes_distributeur") or params["ventes"], "Crédit")
    retours = (par_compte(params["retours"], "Débit", exclude=params.get("remises"))
               - par_compte(params["retours"], "Crédit", exclude=params.get("remises")))
    remises = par_compte(params["remises"], "Débit") - par_compte(params["remises"], "Crédit")
    pstock  = params.get("stock") or ["603"]
    charges = (par_compte(params["charges"], "Débit", exclude=pstock)
               - par_compte(params["charges"], "Crédit", exclude=pstock))
    vstock  = par_compte(pstock, "Débit") - par_compte(pstock, "Crédit")
    prov_rep = params.get("provisions_reprises") or []
    net_prov = (par_compte(prov_rep, "Crédit") - par_compte(prov_rep, "Débit"))
    charges  = charges - net_prov
    mask_cfi = df_i["Compte"].astype(str) == COMPTE_CHARGES_INDIRECTES_REPARTIES
    cf = df_i[mask_cfi].groupby("Code_Analytique")["Débit"].sum().reindex(titres, fill_value=0.0)

    res = pd.DataFrame({"Code_Analytique": titres})
    res["Ventes HT"]  = ventes.values
    res["Retours"]    = retours.values
    res["Remises"]    = remises.values
    res["CA net"]     = res["Ventes HT"] - res["Retours"] - res["Remises"]
    res["Charges variables"] = charges.values
    res["Marge brute"] = res["CA net"] - res["Charges variables"]
    res["Variation de stock"] = vstock.values
    res["Charges fixes imputées"] = cf.values
    res["Résultat net"] = res["Marge brute"] - res["Variation de stock"] - res["Charges fixes imputées"]
    res["Taux retour (%)"] = np.where(ventes_d.values != 0, res["Retours"] / ventes_d.values * 100, 0)
    res["Taux remise (%)"] = np.where(ventes_d.values != 0, res["Remises"] / ventes_d.values * 100, 0)
    res["CA distributeur"] = ventes_d.values

    def _signal(row):
        if row["Résultat net"] > 0 and row["Taux retour (%)"] < 20: return "🟢"
        if row["Résultat net"] > 0 and row["Taux retour (%)"] < 35: return "🟡"
        return "🔴"
    res["Signal"] = res.apply(_signal, axis=1)
    return res

def get_client_ai():
    try:
        return anthropic.Anthropic(api_key=st.secrets["ANTHROPIC_API_KEY"])
    except Exception:
        return None

SYSTEM_PROMPT_EC = """Tu es un expert-comptable spécialisé dans l'accompagnement des maisons d'édition indépendantes françaises.
Tu connais parfaitement les normes comptables françaises (PCG) appliquées à l'édition, les droits d'auteurs, retours éditeurs, provisions, distributeurs.
Tu assistes l'utilisateur dans l'analyse de ses données comptables analytiques via VISION EDITION.
Réponds toujours en français, de façon concise et orientée action."""

SYSTEM_PROMPT_DIRIGEANT = f"""Tu es un assistant de pilotage mis à disposition par le cabinet {NOM_CABINET}.
Tu aides le dirigeant à comprendre ses données de gestion : chiffre d'affaires, taux de retour, rentabilité par ISBN, trésorerie.
Réponds de façon simple, claire et sans jargon comptable excessif.
Tu n'as accès qu'aux données de la maison d'édition de l'utilisateur.

Si la question dépasse le périmètre des données disponibles (conseil stratégique, question juridique, fiscalité, comptabilité générale), réponds exactement :
"Pour cette question, je vous invite à contacter directement votre cabinet {NOM_CABINET}."

Réponds toujours en français."""


# ============================================================
# PAGE : ACCUEIL (EC)
# ============================================================
# ============================================================
def get_df_dirigeant():
    """Récupère le pivot du dossier du dirigeant connecté."""
    did = st.session_state.get("dossier_id")
    if not did:
        st.warning("⚠️ Aucun dossier associé à votre compte.")
        st.stop()
    dos = get_dossier(did)
    if "df_pivot" not in st.session_state:
        st.warning("⚠️ Votre expert-comptable n'a pas encore généré les données de pilotage. "
                   "Contactez votre cabinet.")
        st.stop()
    return st.session_state["df_pivot"].copy(), st.session_state.get("param_comptes",{}), dos




# ============================================================
# FICHE TITRE — dialogue modal
# ============================================================
def _dialog_deco(title, width="large"):
    if hasattr(st, "dialog"):
        return st.dialog(title, width=width)
    def _wrap(func):
        def _inner(*args, **kwargs):
            with st.expander(f"🪟 {title}", expanded=True):
                return func(*args, **kwargs)
        return _inner
    return _wrap

@_dialog_deco("📖 Fiche titre", width="large")
def afficher_fiche_titre(isbn_sel, df, params):
    df_t    = df[df["Code_Analytique"] == isbn_sel]
    df_v_   = df_t[df_t["Compte"].astype(str).str.startswith(tuple(params["ventes"]))]
    df_vd_  = df_t[df_t["Compte"].astype(str).str.startswith(tuple(params.get("ventes_distributeur") or params["ventes"]))]
    df_r_   = df_t[mask_retours(df_t, params)]
    df_rem_ = df_t[mask_remises(df_t, params)]
    pstock  = tuple(params.get("stock") or ["603"])
    df_c_   = df_t[df_t["Compte"].astype(str).str.startswith(tuple(params["charges"]))
                   & (~df_t["Compte"].astype(str).str.startswith(pstock))]
    df_stk_ = df_t[df_t["Compte"].astype(str).str.startswith(pstock)]
    df_cfi_ = df_t[df_t["Compte"].astype(str) == COMPTE_CHARGES_INDIRECTES_REPARTIES]
    pref_p  = tuple(params.get("provisions_reprises") or [])
    df_pv_  = df_t[df_t["Compte"].astype(str).str.startswith(pref_p)] if pref_p else df_t.iloc[0:0]
    net_pv  = df_pv_["Credit"].sum() - df_pv_["Debit"].sum() if False else df_pv_["Crédit"].sum() - df_pv_["Débit"].sum()

    ventes_ht = df_v_["Crédit"].sum()
    ventes_d  = df_vd_["Crédit"].sum()
    retours_m = df_r_["Débit"].sum() - df_r_["Crédit"].sum()
    remises_m = df_rem_["Débit"].sum() - df_rem_["Crédit"].sum()
    ca_net_t  = ventes_ht - retours_m - remises_m
    charges_v = (df_c_["Débit"].sum() - df_c_["Crédit"].sum()) - net_pv
    var_stk   = df_stk_["Débit"].sum() - df_stk_["Crédit"].sum()
    marge_b   = ca_net_t - charges_v
    cf        = df_cfi_["Débit"].sum()
    res_net   = marge_b - var_stk - cf
    taux_ret  = (retours_m / ventes_d * 100) if ventes_d else 0
    taux_rem  = (remises_m / ventes_d * 100) if ventes_d else 0

    if res_net > 0 and taux_ret < 20:
        signal, bg_s, fg_s = "🟢 Titre rentable", "#d1fae5", "#065f46"
    elif res_net > 0 and taux_ret < 35:
        signal, bg_s, fg_s = "🟡 Rentabilité à surveiller", "#fef3c7", "#92400e"
    else:
        signal, bg_s, fg_s = "🔴 Titre en difficulté", "#fee2e2", "#991b1b"

    st.markdown(f"""<div style='padding:14px 18px;border-radius:12px;background:{bg_s};
        color:{fg_s};font-weight:600;font-size:16px;margin-bottom:14px'>
        {label_affiche(isbn_sel,df)} — {signal}</div>""", unsafe_allow_html=True)

    k1,k2,k3,k4 = st.columns(4)
    k1.metric("Taux de retour", f"{taux_ret:.1f} %")
    k2.metric("Taux de remise", f"{taux_rem:.1f} %")
    k3.metric("Marge brute", f"{fmt_fr(marge_b)} €")
    k4.metric("Résultat net", f"{fmt_fr(res_net)} €")

    if cf == 0 and not st.session_state.get("repartition_active"):
        st.caption("ℹ️ Aucune charge fixe imputée : répartition des charges indirectes non activée.")
    if abs(var_stk) > 0.5:
        st.caption(f"ℹ️ Variation de stock : {fmt_fr(var_stk)} € — isolée de la marge brute, incluse dans le résultat net.")

    st.markdown("#### Mini SIG — Soldes intermédiaires de gestion")
    detail_charges = params.get("detail_charges")
    mixte  = params.get("contenu_fabrication_mixte") or {}
    mx_cpt = tuple(mixte.get("comptes") or [])
    mx_kw  = [m.strip().upper() for m in (mixte.get("mots_cles_fabrication") or []) if m.strip()]

    charge_rows = []
    if detail_charges:
        total_det = 0.0; cpts_det = []
        for nom_nat, prefixes in detail_charges.items():
            if prefixes and any(str(p).startswith(str(sp)) or str(sp).startswith(str(p))
                                for p in prefixes for sp in list(pstock)):
                continue
            pref_eff = (list(prefixes) + list(params.get("provisions_reprises") or [])
                        if nom_nat == "Provision pour retour" else list(prefixes))
            df_nat_p = df_t[df_t["Compte"].astype(str).str.startswith(tuple(pref_eff))] if pref_eff else df_t.iloc[0:0]
            df_nat_m = df_t.iloc[0:0]
            if nom_nat in ("Contenu","Fabrication") and mx_cpt:
                df_mx = df_t[df_t["Compte"].astype(str).str.startswith(mx_cpt)]
                if not df_mx.empty and mx_kw:
                    mask_fab = df_mx["Libellé"].astype(str).str.upper().str.contains("|".join(mx_kw),regex=True)
                else:
                    mask_fab = pd.Series(False,index=df_mx.index)
                df_nat_m = df_mx[mask_fab] if nom_nat=="Fabrication" else df_mx[~mask_fab]
                cpts_det.extend(list(mx_cpt))
            df_nat = pd.concat([df_nat_p,df_nat_m]) if not df_nat_m.empty else df_nat_p
            val    = df_nat["Débit"].sum() - df_nat["Crédit"].sum()
            cpts_det.extend(pref_eff)
            charge_rows.append((f"− {nom_nat}", -val, "deduction"))
            total_det += val
        reste = charges_v - total_det
        if abs(reste) > 0.5:
            charge_rows.append(("− Autres charges directes", -reste, "deduction"))
    else:
        charge_rows = [("− Charges variables", -charges_v, "deduction")]

    rows_sig = (
        [("Ventes HT", ventes_ht, "base"),
         ("− Retours", -retours_m, "deduction"),
         ("− Remises", -remises_m, "deduction"),
         ("= CA net", ca_net_t, "subtotal")]
        + charge_rows
        + [("= Marge brute", marge_b, "subtotal")]
        + ([("− Variation de stock", -var_stk, "deduction")] if abs(var_stk) > 0.5 else [])
        + [("− Charges fixes imputées", -cf, "deduction"),
           ("= Résultat net du titre", res_net, "total")]
    )

    html_rows = ""
    for lib_s, mont_s, style_s in rows_sig:
        if style_s == "subtotal":
            rs = "background:#eef2ff;font-weight:600;"
        elif style_s == "total":
            c_ = "#065f46" if mont_s>=0 else "#991b1b"
            f_ = "#d1fae5" if mont_s>=0 else "#fee2e2"
            rs = f"background:{f_};font-weight:700;color:{c_};"
        elif style_s == "deduction":
            rs = "color:#b91c1c;"
        else:
            rs = "font-weight:500;"
        html_rows += (f"<tr style='{rs}'>"
                      f"<td style='padding:6px 12px;border-bottom:1px solid #eee'>{lib_s}</td>"
                      f"<td style='padding:6px 12px;border-bottom:1px solid #eee;text-align:right'>"
                      f"{fmt_fr(mont_s)} €</td></tr>")
    st.markdown(f"<table style='width:100%;border-collapse:collapse;font-size:14px'>"
                f"{html_rows}</table>", unsafe_allow_html=True)

    _sm = {"base":"absolute","deduction":"relative","subtotal":"total","total":"total"}
    fig_f = go.Figure(go.Waterfall(
        orientation="v", measure=[_sm[r[2]] for r in rows_sig],
        x=[r[0] for r in rows_sig], y=[r[1] for r in rows_sig],
        text=[f"{fmt_fr(r[1])} €" for r in rows_sig], textposition="outside",
        connector={"line":{"color":"gray","width":0.5}},
        decreasing={"marker":{"color":"#EF4444"}},
        increasing={"marker":{"color":"#10B981"}},
        totals={"marker":{"color":"#3B82F6"}}
    ))
    fig_f.update_layout(height=300+len(rows_sig)*10, margin=dict(t=10),
                        xaxis_tickangle=-30 if len(rows_sig)>6 else 0)
    st.plotly_chart(fig_f, use_container_width=True)

    df_ev = df_t.copy()
    df_ev["Date"] = pd.to_datetime(df_ev["Date"], errors="coerce")
    df_ev["Mois"] = df_ev["Date"].dt.to_period("M").astype(str)
    ev_v = df_ev[df_ev["Compte"].astype(str).str.startswith(tuple(params["ventes"]))].groupby("Mois")["Crédit"].sum().reset_index()
    ev_r = df_ev[mask_retours(df_ev,params)].groupby("Mois")["Débit"].sum().reset_index()
    ev   = ev_v.merge(ev_r,on="Mois",how="left").fillna(0)
    ev.columns = ["Mois","Ventes","Retours"]
    if not ev.empty:
        st.markdown("#### Évolution mensuelle")
        fig_ev = px.line(ev,x="Mois",y=["Ventes","Retours"],markers=True,height=240)
        fig_ev.update_layout(margin=dict(t=10),legend_title="")
        st.plotly_chart(fig_ev,use_container_width=True)

    buf_ft = BytesIO()
    df_cr_ft = pd.DataFrame({"Poste":[r[0] for r in rows_sig],"Montant (EUR)":[r[1] for r in rows_sig]})
    with pd.ExcelWriter(buf_ft,engine="openpyxl") as writer:
        df_cr_ft.to_excel(writer,index=False,sheet_name="Compte_Resultat")
        if not ev.empty: ev.to_excel(writer,index=False,sheet_name="Evolution_mensuelle")
    buf_ft.seek(0)
    st.download_button("📥 Exporter la fiche titre (Excel)", buf_ft,
        file_name=f"Fiche_{isbn_sel.replace('/','-')}.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        key=f"exp_fiche_{isbn_sel}")


# HELPER check_pivot (défini avant le bloc if/elif)
# ============================================================
def check_pivot():
    if not st.session_state.get("dossier_id"):
        st.warning("⚠️ Sélectionnez un dossier.")
        st.stop()
    if "df_pivot" not in st.session_state:
        st.warning("⚠️ Générez d'abord le socle analytique.")
        st.stop()
    dos = get_dossier(st.session_state["dossier_id"])
    st.info(f"📁 **{dos['nom']}** — {dos.get('exercice','')}")
    return st.session_state["df_pivot"].copy(), st.session_state["param_comptes"]

if role == "ec" and page == "🏠 Accueil":
    st.title("📚 VISION EDITION")
    st.markdown(f"*Bienvenue, **{st.session_state['name']}***")
    st.divider()

    col_main, col_qr = st.columns([3, 1])
    with col_main:
        st.markdown("### Comment démarrer ?")
        c1, c2, c3, c4 = st.columns(4)
        for col, num, icon, titre, desc in [
            (c1,"1","👥","Créer un dossier","Ajoutez votre client maison d'édition"),
            (c2,"2","📂","Importer","Chargez l'export comptable analytique"),
            (c3,"3","⚙️","Paramétrer","Mappez les colonnes et comptes"),
            (c4,"4","📈","Analyser","Explorez les tableaux de bord"),
        ]:
            with col:
                st.markdown(f"""<div style='text-align:center;padding:16px;background:#f8f9fa;
                border-radius:12px;border:1px solid #e0e0e0'>
                <div style='font-size:28px'>{icon}</div>
                <div style='font-size:11px;color:#888;margin:4px 0'>Étape {num}</div>
                <div style='font-weight:600;font-size:14px'>{titre}</div>
                <div style='font-size:12px;color:#666;margin-top:4px'>{desc}</div>
                </div>""", unsafe_allow_html=True)

        st.divider()
        st.markdown("### Outil Analytique Diffuseur")
        st.info("L'outil **Analytique Diffuseur** est une application distincte déployée séparément. "
                "Elle traite le relevé mensuel de votre diffuseur et génère les écritures comptables "
                "prêtes à importer dans votre logiciel avant utilisation de VISION EDITION.")

    with col_qr:
        st.markdown("### Accès mobile")
        qr2 = qrcode.QRCode(version=1, box_size=4, border=2)
        qr2.add_data(APP_URL); qr2.make(fit=True)
        img2 = qr2.make_image(fill_color="black", back_color="white")
        b2 = BytesIO(); img2.save(b2, format="PNG")
        st.image(b2.getvalue(), width=160)
        st.caption("Scanner pour accès mobile")

# ============================================================
# PAGE : GESTION DES DOSSIERS
# ============================================================
elif role == "ec" and page == "👥 Gestion des dossiers":
    st.header("👥 Gestion des dossiers clients")
    cab = st.session_state["cabinet"]

    tab_dossiers, tab_notifs, tab_ec_users, tab_import_export = st.tabs([
        "📁 Dossiers", "📨 Notifications", "👤 Utilisateurs EC", "💾 Import / Export config"
    ])

    # ── Onglet Dossiers ──
    with tab_dossiers:
        st.subheader("Créer un nouveau dossier")
        with st.form("form_nouveau_dossier"):
            c1, c2 = st.columns(2)
            nom_dos    = c1.text_input("Nom de la maison d'édition *", placeholder="Editions de l'Argonaute")
            dirigeant  = c1.text_input("Nom du dirigeant *", placeholder="Thomas Bernard")
            exercice   = c2.text_input("Exercice", value="2025/2026")
            periode_d  = c2.text_input("Début de période", value="01/04/2025")
            periode_f  = c2.text_input("Fin de période", value="31/03/2026")
            email_dir_form = c1.text_input("Email du dirigeant (pour notifications)",
                placeholder="thomas.bernard@editions-argonaute.fr")
            login_dir  = c1.text_input("Login dirigeant (pour accès client)", placeholder="argonaute")
            pw_dir     = c2.text_input("Mot de passe dirigeant", type="password")
            submitted = st.form_submit_button("➕ Créer ce dossier", type="primary")

            if submitted:
                if not nom_dos or not dirigeant:
                    st.error("Le nom de la maison d'édition et le dirigeant sont obligatoires.")
                else:
                    import uuid
                    did = str(uuid.uuid4())[:8]
                    dossier = {
                        "nom": nom_dos,
                        "dirigeant": dirigeant,
                        "exercice": exercice,
                        "periode_debut": periode_d,
                        "periode_fin": periode_f,
                        "email_dirigeant": email_dir_form,
                        "param_comptes": {},
                        "mapping": {},
                        "labels_indirect": {"charges": "CHARGES INDIRECTES", "produits": "PRODUITS INDIRECTS"},
                        "dirigeant_users": {},
                        "notifications": []
                    }
                    if login_dir and pw_dir:
                        dossier["dirigeant_users"][login_dir] = {
                            "pw_hash": _hash(pw_dir),
                            "name": dirigeant
                        }
                    save_dossier(did, dossier)
                    st.success(f"✅ Dossier **{nom_dos}** créé (ID : {did})")
                    if login_dir:
                        st.info(f"Accès dirigeant créé : login = **{login_dir}**")
                    st.rerun()

        st.divider()
        st.subheader("Dossiers existants")
        dossiers = cab["dossiers"]
        if not dossiers:
            st.info("Aucun dossier créé pour l'instant.")
        else:
            for did, dos in dossiers.items():
                with st.expander(f"📁 {dos['nom']} — {dos.get('exercice','')}"):
                    c1, c2, c3 = st.columns(3)
                    c1.write(f"**Dirigeant :** {dos.get('dirigeant','')}")
                    c2.write(f"**Période :** {dos.get('periode_debut','')} au {dos.get('periode_fin','')}")
                    c3.write(f"**ID :** {did}")

                    # Accès dirigeants
                    du = dos.get("dirigeant_users", {})
                    if du:
                        st.write("**Accès dirigeants :**", ", ".join(du.keys()))
                    else:
                        st.caption("Aucun accès dirigeant configuré.")

                    # Ajout accès dirigeant
                    with st.form(f"form_dirigeant_{did}"):
                        dc1, dc2 = st.columns(2)
                        new_login = dc1.text_input("Nouveau login dirigeant", key=f"nl_{did}")
                        new_pw    = dc2.text_input("Mot de passe", type="password", key=f"np_{did}")
                        new_name  = dc1.text_input("Nom affiché", value=dos.get("dirigeant",""), key=f"nn_{did}")
                        if st.form_submit_button("Ajouter accès dirigeant"):
                            if new_login and new_pw:
                                dos["dirigeant_users"][new_login] = {
                                    "pw_hash": _hash(new_pw),
                                    "name": new_name or dos.get("dirigeant","")
                                }
                                save_dossier(did, dos)
                                st.success(f"Accès {new_login} ajouté.")
                                st.rerun()

                    if st.button(f"🗑️ Supprimer ce dossier", key=f"del_{did}"):
                        del st.session_state["cabinet"]["dossiers"][did]
                        if st.session_state["dossier_id"] == did:
                            st.session_state["dossier_id"] = None
                        st.rerun()

    # ── Onglet Utilisateurs EC ──
    with tab_ec_users:
        st.subheader("Collaborateurs du cabinet")
        with st.form("form_ec_user"):
            c1, c2, c3 = st.columns(3)
            ec_login = c1.text_input("Login")
            ec_pw    = c2.text_input("Mot de passe", type="password")
            ec_name  = c3.text_input("Nom affiché")
            if st.form_submit_button("➕ Ajouter collaborateur"):
                if ec_login and ec_pw:
                    cab["ec_users"][ec_login] = {
                        "pw_hash": _hash(ec_pw),
                        "name": ec_name or ec_login,
                        "role": "ec"
                    }
                    st.success(f"Collaborateur {ec_login} ajouté.")
                    st.rerun()

        st.divider()
        for login_u, info in cab["ec_users"].items():
            col1, col2 = st.columns([4,1])
            col1.write(f"👤 **{info['name']}** (login : `{login_u}`)")
            if login_u != st.session_state["username"]:
                if col2.button("Supprimer", key=f"del_ec_{login_u}"):
                    del cab["ec_users"][login_u]
                    st.rerun()

    # ── Onglet Notifications ──
    with tab_notifs:
        st.subheader("Envoyer une notification au dirigeant")

        dossiers_notif = st.session_state["cabinet"]["dossiers"]
        if not dossiers_notif:
            st.info("Créez d'abord un dossier client.")
        else:
            did_notif = st.selectbox(
                "Dossier client",
                options=list(dossiers_notif.keys()),
                format_func=lambda x: dossiers_notif[x]["nom"],
                key="notif_dossier_sel"
            )
            dos_notif = dossiers_notif[did_notif]

            st.markdown("**Message in-app** *(visible à la connexion du dirigeant)*")
            with st.form("form_notif"):
                type_notif = st.selectbox("Type de message", [
                    "✅ info — Mise à jour disponible",
                    "⚠️ warning — Point d'attention",
                    "🔴 error — Action requise"
                ], key="notif_type")
                msg_notif = st.text_area("Message",
                    placeholder="Ex : Votre tableau de bord de juillet est disponible. "
                                "Votre taux de retour est de 18% — dans la norme sectorielle.",
                    height=100, key="notif_msg")

                st.divider()
                st.markdown("**Notification par email** *(optionnel)*")
                email_dir = st.text_input("Email du dirigeant",
                    value=dos_notif.get("email_dirigeant",""),
                    placeholder="thomas.bernard@editions-argonaute.fr",
                    key="notif_email")
                envoyer_mail = st.checkbox("Envoyer aussi par email", value=True, key="notif_send_mail")

                submitted_notif = st.form_submit_button("📨 Envoyer la notification", type="primary")

                if submitted_notif:
                    if not msg_notif.strip():
                        st.error("Le message ne peut pas être vide.")
                    else:
                        # Sauvegarder email si renseigné
                        if email_dir:
                            dos_notif["email_dirigeant"] = email_dir
                            save_dossier(did_notif, dos_notif)

                        # Notification in-app
                        type_key = "info"
                        if "warning" in type_notif: type_key = "warning"
                        elif "error" in type_notif: type_key = "error"
                        ajouter_notification(did_notif, msg_notif.strip(), type_key)
                        st.success("✅ Notification in-app enregistrée.")

                        # Email
                        if envoyer_mail and email_dir:
                            nom_dos_notif = dos_notif.get("nom","votre maison d'édition")
                            corps_mail = f"""Bonjour,

Votre expert-comptable vous adresse un message concernant {nom_dos_notif} :

{msg_notif.strip()}

Pour consulter votre tableau de bord, connectez-vous à VISION EDITION :
{APP_URL}

Cordialement,
CAB EDITION"""
                            ok, msg_ret = envoyer_email(
                                email_dir,
                                f"Mise à jour — {nom_dos_notif}",
                                corps_mail
                            )
                            if ok:
                                st.success(f"✅ Email envoyé à {email_dir}")
                            else:
                                st.warning(f"⚠️ Notification in-app OK mais email non envoyé : {msg_ret}")
                        elif envoyer_mail and not email_dir:
                            st.warning("⚠️ Aucun email renseigné — notification in-app uniquement.")

            # Historique notifications
            st.divider()
            st.subheader("Historique des notifications")
            dos_hist = get_dossier(did_notif)
            notifs_hist = dos_hist.get("notifications", [])
            if not notifs_hist:
                st.info("Aucune notification envoyée pour ce dossier.")
            else:
                for i, n in enumerate(reversed(notifs_hist)):
                    statut = "🔵 Non lu" if not n["lu"] else "✅ Lu"
                    with st.expander(f"{statut} — {n['date']} — {n['message'][:50]}..."):
                        st.write(f"**Message :** {n['message']}")
                        st.write(f"**Type :** {n['type']} | **Statut :** {statut}")
                if st.button("🗑️ Effacer l'historique", key="clear_notifs"):
                    dos_hist["notifications"] = []
                    save_dossier(did_notif, dos_hist)
                    st.rerun()

    # ── Onglet Import/Export config ──
    with tab_import_export:
        st.subheader("Sauvegarder la configuration cabinet")
        st.caption("Exportez la configuration complète (dossiers + paramètres) en JSON. "
                   "Réimportez-la à la prochaine session pour retrouver tous vos dossiers.")

        config_export = {
            "ec_users": {
                k: {"name": v["name"], "role": v["role"], "pw_hash": v["pw_hash"]}
                for k, v in cab["ec_users"].items()
            },
            "dossiers": cab["dossiers"]
        }

        st.download_button(
            "💾 Exporter la configuration cabinet (JSON)",
            data=json.dumps(config_export, ensure_ascii=False, indent=2),
            file_name=f"config_cabinet_{_dt.date.today().isoformat()}.json",
            mime="application/json",
            use_container_width=True
        )

        st.divider()
        st.subheader("Recharger une configuration")
        fichier_config = st.file_uploader("Fichier JSON de configuration", type=["json"])
        if fichier_config:
            try:
                cfg = json.load(fichier_config)
                if st.button("📤 Appliquer cette configuration", type="primary"):
                    if "ec_users" in cfg:
                        cab["ec_users"].update(cfg["ec_users"])
                    if "dossiers" in cfg:
                        cab["dossiers"].update(cfg["dossiers"])
                    st.success("✅ Configuration chargée avec succès.")
                    st.rerun()
            except Exception as e:
                st.error(f"Fichier invalide : {e}")


# ============================================================
# PAGE : IMPORT DES DONNÉES (EC)
# ============================================================
elif role == "ec" and page == "📂 Import des données":
    st.header("📂 Import des données analytiques")

    if not st.session_state.get("dossier_id"):
        st.warning("⚠️ Sélectionnez ou créez un dossier client dans **👥 Gestion des dossiers**.")
        st.stop()

    dos = get_dossier(st.session_state["dossier_id"])
    st.info(f"📁 Dossier actif : **{dos['nom']}** — {dos.get('exercice','')}")

    tab1, tab2 = st.tabs(["📁 Importer mon fichier", "🎭 Données de démonstration"])

    with tab1:
        st.info("Importez votre export comptable analytique au format Excel (.xlsx). "
                "Ce fichier est le grand livre analytique exporté depuis votre logiciel comptable "
                "après import des écritures générées par **Analytique Diffuseur**.")
        fichier = st.file_uploader("Sélectionnez votre fichier Excel", type=["xlsx"])
        if fichier:
            try:
                df = pd.read_excel(fichier, header=0)
                df.columns = df.columns.str.strip()
                st.session_state["df_comptables"] = df
                st.success(f"✅ {df.shape[0]} lignes chargées — {df.shape[1]} colonnes")
                st.write("**Colonnes :**", list(df.columns))
                st.dataframe(df.head(10))
                st.info("➡️ Passez dans **⚙️ Paramétrage analytique** pour configurer les colonnes.")
            except Exception as e:
                st.error(f"❌ Erreur : {e}")

    with tab2:
        st.info("Données fictives pour explorer tous les modules.")
        if st.button("🎭 Charger les données de démonstration", type="primary"):
            np.random.seed(42)
            isbns = [f"978-2-{i:04d}-{j:04d}-{k}" for i,j,k in [
                (1234,1001,1),(1234,1002,8),(1234,1003,5),(1234,1004,2),
                (1234,1005,9),(1234,1006,6),(1234,1007,3),(1234,1008,0)]]
            titres_d = ["Le Dernier Manuscrit","Mémoires du Vent","Sous les Toits de Paris",
                        "L'Héritage Silencieux","Chroniques du Nord","La Lumière d'Août",
                        "Terres Inconnues","Les Mots du Soir"]
            rows = []
            dates = pd.date_range("2024-01-01","2024-12-31",freq="MS")
            for date in dates:
                for isbn,titre in zip(isbns,titres_d):
                    v = np.random.randint(500,8000)
                    rows.append({"Compte":"701100","Débit":0,"Crédit":round(v,2),
                                 "Code_Analytique":isbn,"Famille_Analytique":"EDITION",
                                 "Libellé":f"Ventes {titre}","Date":date})
                    if np.random.random()<0.6:
                        rows.append({"Compte":"709100","Débit":round(v*np.random.uniform(0.05,0.30),2),"Crédit":0,
                                     "Code_Analytique":isbn,"Famille_Analytique":"EDITION",
                                     "Libellé":f"Retours {titre}","Date":date})
                    rows.append({"Compte":"607100","Débit":round(v*np.random.uniform(0.15,0.40),2),"Crédit":0,
                                 "Code_Analytique":isbn,"Famille_Analytique":"EDITION",
                                 "Libellé":f"Charges {titre}","Date":date})
                for compte,libelle,montant in [("615000","Loyer",1200),("641000","Salaires",4500)]:
                    rows.append({"Compte":compte,"Débit":montant,"Crédit":0,
                                 "Code_Analytique":"CHARGES INDIRECTES","Famille_Analytique":"EDITION",
                                 "Libellé":libelle,"Date":date})
            df_demo = pd.DataFrame(rows)
            df_demo["Date"] = pd.to_datetime(df_demo["Date"])
            st.session_state["df_comptables"] = df_demo
            st.session_state["df_source_mappe"] = df_demo
            pivot = df_demo.groupby(["Compte","Famille_Analytique","Code_Analytique","Date","Libellé"],
                                    as_index=False).agg({"Débit":"sum","Crédit":"sum"})
            st.session_state["df_pivot"] = pivot
            st.session_state["df_pivot_brut"] = pivot.copy()
            st.session_state["param_comptes"] = {
                "ventes":["701"],"ventes_distributeur":["701"],"retours":["709"],
                "remises":["7091"],"charges":["6"],"provisions_reprises":["781"],
                "charges_imputees":"Oui"
            }
            st.session_state["labels_indirect"] = {"charges":"CHARGES INDIRECTES","produits":"PRODUITS INDIRECTS"}
            st.session_state["familles_cols"] = ["Famille_Analytique"]
            st.session_state["codes_cols"] = ["Code_Analytique"]
            st.session_state["noms_familles_actives"] = ["EDITION"]
            st.session_state["repartition_active"] = False
            st.success("✅ Données de démonstration chargées !")
            st.info("➡️ Accédez à **📈 Tableau de bord éditorial**.")

# ============================================================
# PAGE : PARAMÉTRAGE ANALYTIQUE (EC) — simplifié
# ============================================================
elif role == "ec" and page == "⚙️ Paramétrage analytique":
    st.header("⚙️ Paramétrage analytique")

    if not st.session_state.get("dossier_id"):
        st.warning("⚠️ Sélectionnez un dossier.")
        st.stop()

    dos = get_dossier(st.session_state["dossier_id"])
    st.info(f"📁 **{dos['nom']}** — {dos.get('exercice','')}")

    if "df_comptables" not in st.session_state:
        st.warning("⚠️ Importez d'abord vos données.")
        st.stop()

    df = st.session_state["df_comptables"].copy()
    columns = list(df.columns)

    # Recharger config dossier si disponible
    saved_params = dos.get("param_comptes", {})
    saved_mapping = dos.get("mapping", {})

    with st.expander("📥 Recharger une configuration JSON sauvegardée", expanded=False):
        fconfig = st.file_uploader("Fichier JSON de configuration", type=["json"], key="up_config_param")
        if fconfig:
            try:
                cfg = json.load(fconfig)
                if st.button("Appliquer"):
                    cfg_inv = {v:k for k,v in (cfg.get("mapping",{})).items()}
                    for k_ss, col_cfg in [
                        ("map_compte_col","Compte"),("map_debit_col","Débit"),
                        ("map_credit_col","Crédit"),("map_date_col","Date")]:
                        if cfg_inv.get(col_cfg) in columns:
                            st.session_state[k_ss] = cfg_inv[col_cfg]
                    p = cfg.get("param_comptes",{})
                    st.session_state["map_ventes_comptes"] = ",".join(p.get("ventes",["701"]))
                    st.session_state["map_retours_comptes"] = ",".join(p.get("retours",["709"]))
                    st.session_state["map_remises_comptes"] = ",".join(p.get("remises",["7091"]))
                    st.session_state["map_charges_comptes"] = ",".join(p.get("charges",["6"]))
                    st.session_state["map_provisions_reprises_comptes"] = ",".join(p.get("provisions_reprises",["781"]))
                    st.session_state["map_ventes_distributeur_comptes"] = ",".join(p.get("ventes_distributeur",["7011"]))
                    st.success("✅ Configuration rechargée.")
                    st.rerun()
            except Exception as e:
                st.error(f"Fichier invalide : {e}")

    col1, col2 = st.columns(2)
    with col1:
        st.subheader("Mapping des colonnes de base")
        compte_col  = st.selectbox("Colonne Compte", columns, key="map_compte_col")
        debit_col   = st.selectbox("Colonne Débit", columns, key="map_debit_col")
        credit_col  = st.selectbox("Colonne Crédit", columns, key="map_credit_col")
        date_col    = st.selectbox("Colonne Date", columns, key="map_date_col")
        libelle_col = st.selectbox("Libellé (optionnel)", [""]+columns, key="map_libelle_col")
        journal_col = st.selectbox("Code journal (optionnel)", [""]+columns, key="map_journal_col",
            help="Recommandé pour le module Trésorerie : permet d'exclure les écritures AN.")
    with col2:
        st.subheader("Comptes comptables")
        ventes_comptes  = st.text_input("Comptes ventes (CA large)", value=",".join(saved_params.get("ventes",["701"])), key="map_ventes_comptes")
        ventes_dist     = st.text_input("Comptes ventes distributeur (base taux retour/remise)", value=",".join(saved_params.get("ventes_distributeur",["7011"])), key="map_ventes_distributeur_comptes",
            help="Sous-compte strict du diffuseur (ex: 7011 pour BLDD). Base exclusive du taux de retour/remise.")
        retours_comptes = st.text_input("Comptes retours", value=",".join(saved_params.get("retours",["709"])), key="map_retours_comptes")
        remises_comptes = st.text_input("Comptes remises", value=",".join(saved_params.get("remises",["7091"])), key="map_remises_comptes")
        charges_comptes = st.text_input("Comptes charges", value=",".join(saved_params.get("charges",["6"])), key="map_charges_comptes")
        prov_comptes    = st.text_input("Comptes reprises sur provisions (ex. retour)", value=",".join(saved_params.get("provisions_reprises",["781"])), key="map_provisions_reprises_comptes",
            help="Reprises sur provisions (ex. 781) — nettées contre les charges, pas comptées en CA.")
        charges_imputees = st.radio("Charges déjà imputées par section ?", ["Oui","Non"], key="map_charges_imputees")

    # ── Détail des charges par nature (mini SIG) ──
    with st.expander("📐 Détail des charges par nature (optionnel — mini SIG détaillé par titre)"):
        st.caption(
            "Renseignez ici les comptes correspondant à chaque nature de charge directe. "
            "Si cette section reste vide, la fiche titre affiche une seule ligne agrégée "
            "« Charges variables ». Dès qu'au moins une nature est renseignée, la fiche titre "
            "affiche un compte de résultat en cascade détaillé par nature de charge."
        )
        col_d1, col_d2 = st.columns(2)
        with col_d1:
            cpt_stock   = st.text_input("Variation de stock", value=saved_params.get("detail_charges",{}).get("Variation de stock",[""])[0] if saved_params.get("detail_charges") else "", key="cpt_variation_stock", help="Ex. 603")
            cpt_droits  = st.text_input("Droits d'auteur", value="", key="cpt_droits_auteur_detail", help="Ex. 604300000")
            cpt_comm    = st.text_input("Commercialisation", value="", key="cpt_commercialisation", help="Ex. 645106,6228")
            cpt_struct  = st.text_input("Structure / gérant", value="", key="cpt_structure", help="Ex. 6411,6451")
        with col_d2:
            cpt_prov_d  = st.text_input("Provision pour retour (dotation)", value="", key="cpt_dotations", help="Ex. 6810")
            cpt_contenu = st.text_input("Contenu (préparation éditoriale, prépresse)", value="", key="cpt_contenu")
            cpt_fab     = st.text_input("Fabrication (impression, façonnage)", value="", key="cpt_fabrication", help="Ex. 605")
            cpt_mixte   = st.text_input("Compte mixte contenu/fabrication (scindé par mot-clé)", value="", key="cpt_mixte_contenu_fab", help="Ex. 604 — scindé automatiquement par mot-clé dans le libellé")
            cpt_mots_cles = st.text_input("Mots-clés fabrication", value="REPROGRAPHIE,IMPRIM,PRINT,FACONNAGE", key="cpt_mots_cles_fab")

    st.subheader("Familles analytiques")
    nb_familles = st.number_input("Nombre de familles", min_value=1, max_value=4, value=1, step=1, key="map_nb_familles")
    familles_mapping = []
    noms_suggestion = ["EDITION","COMMUNICATION","Types de dépenses / revenus","AUTEUR"]
    for i in range(int(nb_familles)):
        fc1,fc2,fc3 = st.columns(3)
        nom_f = fc1.text_input(f"Famille {i+1}", value=noms_suggestion[i] if i<4 else "", key=f"nom_famille_{i}")
        fam_col = fc2.selectbox(f"Colonne famille {i+1}", [""]+columns, key=f"famille_col_{i}")
        cod_col = fc3.selectbox(f"Colonne catégorie {i+1}", [""]+columns, key=f"code_col_{i}")
        familles_mapping.append({"nom":nom_f,"famille_col":fam_col,"code_col":cod_col})

    col_li1,col_li2 = st.columns(2)
    label_ci = col_li1.text_input("Libellé charges indirectes", value="CHARGES INDIRECTES", key="map_label_ci")
    label_pi = col_li2.text_input("Libellé produits indirects", value="PRODUITS INDIRECTS", key="map_label_pi")

    if st.button("⚙️ Générer le socle analytique", type="primary"):
        mapping = {compte_col:"Compte",debit_col:"Débit",credit_col:"Crédit",date_col:"Date"}
        if libelle_col: mapping[libelle_col] = "Libellé"
        if journal_col: mapping[journal_col] = "Journal"
        familles_cols, codes_cols, noms_familles_actives = [], [], []
        for i,fam in enumerate(familles_mapping):
            suffix = "" if i==0 else f"_{i+1}"
            col_f_out = f"Famille_Analytique{suffix}"
            col_c_out = f"Code_Analytique{suffix}"
            if fam["famille_col"]: mapping[fam["famille_col"]] = col_f_out
            if fam["code_col"]:    mapping[fam["code_col"]]    = col_c_out
            familles_cols.append(col_f_out); codes_cols.append(col_c_out)
            noms_familles_actives.append(fam["nom"] or f"Famille {i+1}")

        df.rename(columns=mapping, inplace=True)
        for col in familles_cols+codes_cols+["Libellé","Journal"]:
            if col not in df.columns: df[col] = ""
            else: df[col] = df[col].fillna("")
        df["Date"]   = pd.to_datetime(df["Date"], errors="coerce")
        df["Débit"]  = pd.to_numeric(df["Débit"], errors="coerce").fillna(0)
        df["Crédit"] = pd.to_numeric(df["Crédit"], errors="coerce").fillna(0)
        df["Compte"] = df["Compte"].astype(str).str.strip()

        def sc(s): return [c.strip() for c in s.split(",") if c.strip()]

        detail_charges = {
            "Variation de stock":    sc(st.session_state.get("cpt_variation_stock","")),
            "Droits d'auteur":       sc(st.session_state.get("cpt_droits_auteur_detail","")),
            "Commercialisation":     sc(st.session_state.get("cpt_commercialisation","")),
            "Structure/gérant":      sc(st.session_state.get("cpt_structure","")),
            "Provision pour retour": sc(st.session_state.get("cpt_dotations","")),
            "Contenu":               sc(st.session_state.get("cpt_contenu","")),
            "Fabrication":           sc(st.session_state.get("cpt_fabrication","")),
        }
        mixte_comptes  = sc(st.session_state.get("cpt_mixte_contenu_fab",""))
        mots_cles_fab_ = sc(st.session_state.get("cpt_mots_cles_fab","REPROGRAPHIE,IMPRIM,PRINT,FACONNAGE"))
        detail_actif   = any(v for v in detail_charges.values()) or bool(mixte_comptes)

        params = {
            "ventes":  sc(ventes_comptes),
            "ventes_distributeur": sc(ventes_dist) or sc(ventes_comptes),
            "retours": sc(retours_comptes),
            "remises": sc(remises_comptes),
            "charges": sc(charges_comptes),
            "provisions_reprises": sc(prov_comptes),
            "charges_imputees": st.session_state.get("map_charges_imputees","Oui"),
            "detail_charges": detail_charges if detail_actif else None,
            "contenu_fabrication_mixte": {"comptes": mixte_comptes, "mots_cles_fabrication": mots_cles_fab_} if mixte_comptes else None,
            "stock": sc(st.session_state.get("cpt_variation_stock","")) or ["603"],
        }

        group_cols = ["Compte"]+familles_cols+codes_cols+["Date"]
        if "Libellé" in df.columns: group_cols.append("Libellé")
        if "Journal" in df.columns: group_cols.append("Journal")
        pivot = df.groupby(group_cols, as_index=False).agg({"Débit":"sum","Crédit":"sum"})
        pivot = normaliser_codes_isbn(pivot)

        st.session_state["df_source_mappe"] = df
        st.session_state["df_pivot"] = pivot
        st.session_state["df_pivot_brut"] = pivot.copy()
        st.session_state["param_comptes"] = params
        st.session_state["labels_indirect"] = {"charges":label_ci.strip(),"produits":label_pi.strip()}
        st.session_state["familles_cols"] = familles_cols
        st.session_state["codes_cols"] = codes_cols
        st.session_state["noms_familles_actives"] = noms_familles_actives
        st.session_state["repartition_active"] = False

        # Sauvegarder dans le dossier
        dos["param_comptes"] = params
        dos["mapping"] = mapping
        dos["labels_indirect"] = st.session_state["labels_indirect"]
        save_dossier(st.session_state["dossier_id"], dos)

        st.success("✅ Socle analytique généré et paramètres sauvegardés dans le dossier.")

        config_out = {"mapping":mapping,"param_comptes":params,
                      "familles_cols":familles_cols,"codes_cols":codes_cols,
                      "noms_familles_actives":noms_familles_actives,
                      "labels_indirect":st.session_state["labels_indirect"]}
        st.download_button("💾 Sauvegarder la configuration (JSON)",
            data=json.dumps(config_out,ensure_ascii=False,indent=2),
            file_name=f"config_{dos['nom'][:20].replace(' ','_')}.json",
            mime="application/json")
        st.dataframe(pivot.head(15))

    # ── Répartition des charges indirectes (hors bouton, toujours visible) ──
    if "df_pivot_brut" in st.session_state:
        st.divider()
        st.subheader("📐 Répartition des charges et produits indirects")
        st.markdown(
            "Les charges de structure non imputées à un titre ont été codées comme **charges indirectes**. "
            "Vous pouvez les répartir au **nombre de titres actifs** (plutôt qu'au CA) pour ne pas masquer "
            "les titres non rentables derrière les titres porteurs."
        )
        pivot_brut = st.session_state["df_pivot_brut"]
        label_ci   = st.session_state.get("labels_indirect",{}).get("charges","CHARGES INDIRECTES")
        label_pi   = st.session_state.get("labels_indirect",{}).get("produits","PRODUITS INDIRECTS")
        masque_ci  = pivot_brut["Code_Analytique"].astype(str).str.strip() == label_ci
        masque_pi  = pivot_brut["Code_Analytique"].astype(str).str.strip() == label_pi
        total_ci   = (pivot_brut[masque_ci]["Débit"] - pivot_brut[masque_ci]["Crédit"]).sum()
        total_pi   = (pivot_brut[masque_pi]["Crédit"] - pivot_brut[masque_pi]["Débit"]).sum()
        titres_act = sorted(filtrer_isbn_reels(pivot_brut)["Code_Analytique"].astype(str).unique().tolist())
        nb_titres  = len(titres_act)

        col_r1,col_r2,col_r3 = st.columns(3)
        col_r1.metric("Charges indirectes détectées", f"{fmt_fr(total_ci,2)} €")
        col_r2.metric("Produits indirects détectés",  f"{fmt_fr(total_pi,2)} €")
        col_r3.metric("Nombre de titres actifs",       nb_titres)

        repartir = st.radio(
            "Souhaitez-vous répartir les charges et produits indirects sur les titres actifs ?",
            ["Non, je garde une ligne 'indirecte' globale","Oui, répartir sur les titres actifs"],
            index=0, key="repartir_radio"
        )

        if repartir.startswith("Oui"):
            if nb_titres == 0:
                st.error("❌ Aucun titre actif détecté — répartition impossible.")
            else:
                pivot_rep = pivot_brut[~masque_ci & ~masque_pi].copy()
                nouvelles = []

                total_ci_d = pivot_brut[masque_ci]["Débit"].sum()
                total_ci_c = pivot_brut[masque_ci]["Crédit"].sum()
                if total_ci_d != 0 or total_ci_c != 0:
                    pcd = round(total_ci_d / nb_titres, 2)
                    pcc = round(total_ci_c / nb_titres, 2)
                    for isbn in titres_act:
                        nouvelles.append({
                            "Compte": COMPTE_CHARGES_INDIRECTES_REPARTIES,
                            "Famille_Analytique": "EDITION",
                            "Code_Analytique": isbn,
                            "Date": pd.NaT,
                            "Libellé": f"Quote-part charges indirectes ({nb_titres} titres actifs)",
                            "Débit": pcd, "Crédit": pcc,
                        })

                total_pi_d = pivot_brut[masque_pi]["Débit"].sum()
                total_pi_c = pivot_brut[masque_pi]["Crédit"].sum()
                if total_pi_d != 0 or total_pi_c != 0:
                    ppd = round(total_pi_d / nb_titres, 2)
                    ppc = round(total_pi_c / nb_titres, 2)
                    for isbn in titres_act:
                        nouvelles.append({
                            "Compte": COMPTE_PRODUITS_INDIRECTS_REPARTIS,
                            "Famille_Analytique": "EDITION",
                            "Code_Analytique": isbn,
                            "Date": pd.NaT,
                            "Libellé": f"Quote-part produits indirects ({nb_titres} titres actifs)",
                            "Débit": ppd, "Crédit": ppc,
                        })

                if nouvelles:
                    df_nouv = pd.DataFrame(nouvelles)
                    for c in pivot_rep.columns:
                        if c not in df_nouv.columns: df_nouv[c] = None
                    pivot_rep = pd.concat([pivot_rep, df_nouv[pivot_rep.columns]], ignore_index=True)

                st.session_state["df_pivot"] = pivot_rep
                st.session_state["repartition_active"] = True
                st.session_state["repartition_detail"] = {
                    "nb_titres_actifs": nb_titres,
                    "part_charge":  round(total_ci / nb_titres, 2) if nb_titres else 0,
                    "part_produit": round(total_pi / nb_titres, 2) if nb_titres else 0,
                }
                st.success(
                    f"✅ Répartition effectuée sur {nb_titres} titres actifs : "
                    f"{fmt_fr(round(total_ci/nb_titres,2),2)} € de charges et "
                    f"{fmt_fr(round(total_pi/nb_titres,2),2)} € de produits par titre."
                )
        else:
            st.session_state["df_pivot"] = pivot_brut
            st.session_state["repartition_active"] = False
            st.info("Les charges et produits indirects restent sur une ligne globale. "
                    "Vous pouvez modifier ce choix à tout moment sans régénérer le socle.")


# ── TABLEAU DE BORD ──
elif role == "ec" and page == "📈 Tableau de bord éditorial":
    df, params = check_pivot()
    dos = get_dossier(st.session_state["dossier_id"])
    st.header(f"📈 Tableau de bord — {dos['nom']}")
    if st.session_state.get("repartition_active"):
        st.caption("ℹ️ Charges indirectes réparties sur les titres actifs.")

    df["Date"] = pd.to_datetime(df["Date"], errors="coerce")
    years = sorted(df["Date"].dt.year.dropna().unique().tolist())
    annee = st.selectbox("Année", ["Toutes"]+[str(y) for y in years])
    if annee != "Toutes":
        df = df[df["Date"].dt.year == int(annee)]

    df_date_ok = df.dropna(subset=["Date"]).copy()
    df_date_ok["Mois"] = df_date_ok["Date"].dt.to_period("M").astype(str)

    df_v   = df[mask_ventes(df, params)]
    df_r   = df[mask_retours(df, params)]
    df_rem = df[mask_remises(df, params)]
    df_c   = df[mask_charges(df, params)]
    pref_dist = tuple(params.get("ventes_distributeur") or params["ventes"])
    df_dist = df[df["Compte"].astype(str).str.startswith(pref_dist)]
    ca_dist = df_dist["Crédit"].sum()

    df_prov = df[mask_provisions_reprises(df, params)]
    net_prov = df_prov["Crédit"].sum() - df_prov["Débit"].sum()

    ca_brut       = df_v["Crédit"].sum()
    corrections   = df_v["Débit"].sum()
    total_retours = df_r["Débit"].sum() - df_r["Crédit"].sum()
    total_remises = df_rem["Débit"].sum() - df_rem["Crédit"].sum()
    ca_net        = ca_brut - total_retours - total_remises - corrections
    charges_tot   = (df_c["Débit"].sum() - df_c["Crédit"].sum()) - net_prov
    resultat      = ca_net - charges_tot
    taux_ret      = (total_retours / ca_dist * 100) if ca_dist else 0
    taux_rem      = (total_remises / ca_dist * 100) if ca_dist else 0

    k1,k2,k3,k4,k5 = st.columns(5)
    k1.metric("CA brut", f"{fmt_fr(ca_brut)} €")
    k2.metric("CA net", f"{fmt_fr(ca_net)} €", delta=f"-{fmt_fr(total_retours+total_remises+corrections)} €")
    k3.metric("Taux de retour", f"{taux_ret:.1f} %",
              delta_color="inverse", delta="⚠️ Élevé" if taux_ret > 25 else "✅ Normal")
    k4.metric("Charges totales", f"{fmt_fr(charges_tot)} €")
    k5.metric("Résultat net", f"{fmt_fr(resultat)} €")

    col_g1, col_g2 = st.columns(2)
    with col_g1:
        st.subheader("Évolution mensuelle")
        df_v_d = df_date_ok[mask_ventes(df_date_ok, params)]
        df_r_d = df_date_ok[mask_retours(df_date_ok, params)]
        tv = df_v_d.groupby("Mois")["Crédit"].sum().reset_index().rename(columns={"Crédit":"CA brut"})
        tr = df_r_d.groupby("Mois")["Débit"].sum().reset_index().rename(columns={"Débit":"Retours"})
        trend = tv.merge(tr, on="Mois", how="left").fillna(0)
        trend["CA net"] = trend["CA brut"] - trend["Retours"]
        fig = go.Figure()
        fig.add_trace(go.Bar(x=trend["Mois"], y=trend["CA brut"], name="CA brut", marker_color="#3B82F6"))
        fig.add_trace(go.Bar(x=trend["Mois"], y=trend["Retours"], name="Retours", marker_color="#EF4444"))
        fig.add_trace(go.Scatter(x=trend["Mois"], y=trend["CA net"], name="CA net",
                                 mode="lines+markers", line=dict(color="#10B981",width=2)))
        fig.update_layout(barmode="overlay", height=320, margin=dict(t=20), legend=dict(orientation="h"))
        st.plotly_chart(fig, use_container_width=True)

    with col_g2:
        st.subheader("Top 10 titres")
        df_isbn = filtrer_isbn_reels(df)
        top = df_isbn.groupby("Code_Analytique", as_index=False).agg({"Crédit":"sum","Débit":"sum"})
        top["Résultat"] = top["Crédit"] - top["Débit"]
        top10 = top.nlargest(10,"Résultat").copy()
        top10["Titre"] = top10["Code_Analytique"].apply(lambda c: label_affiche(c, df))
        fig2 = px.bar(top10, x="Titre", y="Résultat", color="Résultat",
                      color_continuous_scale=["#EF4444","#F59E0B","#10B981"], height=320)
        fig2.update_layout(showlegend=False, margin=dict(t=20))
        st.plotly_chart(fig2, use_container_width=True)

# ── ANALYSE PAR TITRE ──
elif role == "ec" and page == "📖 Analyse par titre":
    df, params = check_pivot()
    dos = get_dossier(st.session_state["dossier_id"])
    st.header(f"📖 Analyse par titre — {dos['nom']}")
    if st.session_state.get("repartition_active"):
        st.caption("ℹ️ Les charges/produits indirects ont été répartis sur les titres actifs.")

    titres = sorted(filtrer_isbn_reels(df)["Code_Analytique"].astype(str).unique().tolist())
    if not titres:
        st.warning("Aucun ISBN/code analytique détecté.")
        st.stop()

    indic = calculer_indicateurs_titres(df, params, titres)

    # ── Repères rapides ──
    st.subheader("🧭 Repères rapides")
    col_s, col_t, col_f = st.columns(3)

    def _liste_titres(container, sous_titre, df_tri, col_montant, prefix_key):
        container.markdown(f"**{sous_titre}**")
        for _, row in df_tri.iterrows():
            c1, c2 = container.columns([4, 1])
            c1.markdown(f"{row['Signal']} **{label_affiche(row['Code_Analytique'],df)}** — {fmt_fr(row[col_montant])} €")
            if c2.button("Voir", key=f"{prefix_key}_{row['Code_Analytique']}", use_container_width=True):
                afficher_fiche_titre(row["Code_Analytique"], df, params)

    _liste_titres(col_s, "📊 Plus significatifs (CA brut)", indic.sort_values("Ventes HT",ascending=False).head(5), "Ventes HT", "signif")
    _liste_titres(col_t, "🏆 Plus rentables (résultat net)", indic.sort_values("Résultat net",ascending=False).head(5), "Résultat net", "rent")
    _liste_titres(col_f, "⚠️ Plus difficiles (résultat net)", indic.sort_values("Résultat net",ascending=True).head(5), "Résultat net", "diff")

    st.divider()
    st.divider()
    st.markdown("**Ouvrir la fiche détaillée d'un titre**")
    col_sel_pa, col_btn_pa = st.columns([4,1])
    with col_sel_pa:
        isbn_sel = st.selectbox("Titre (ISBN)", titres, label_visibility="collapsed",
                                format_func=lambda c: label_affiche(c,df))
    with col_btn_pa:
        if st.button("📖 Ouvrir la fiche", type="primary", use_container_width=True,
                     key="btn_ouvrir_fiche"):
            afficher_fiche_titre(isbn_sel, df, params)

    # ── Calculs fiche titre ──
    df_t    = df[df["Code_Analytique"] == isbn_sel]
    df_v_   = df_t[df_t["Compte"].astype(str).str.startswith(tuple(params["ventes"]))]
    df_vd_  = df_t[df_t["Compte"].astype(str).str.startswith(tuple(params.get("ventes_distributeur") or params["ventes"]))]
    df_r_   = df_t[mask_retours(df_t, params)]
    df_rem_ = df_t[mask_remises(df_t, params)]
    pstock  = tuple(params.get("stock") or ["603"])
    df_c_   = df_t[df_t["Compte"].astype(str).str.startswith(tuple(params["charges"]))
                   & (~df_t["Compte"].astype(str).str.startswith(pstock))]
    df_stock_ = df_t[df_t["Compte"].astype(str).str.startswith(pstock)]
    df_cfi    = df_t[df_t["Compte"].astype(str) == COMPTE_CHARGES_INDIRECTES_REPARTIES]

    pref_prov = tuple(params.get("provisions_reprises") or [])
    df_prov_t = df_t[df_t["Compte"].astype(str).str.startswith(pref_prov)] if pref_prov else df_t.iloc[0:0]
    net_prov_t = df_prov_t["Crédit"].sum() - df_prov_t["Débit"].sum()

    ventes_ht   = df_v_["Crédit"].sum()
    ventes_dist = df_vd_["Crédit"].sum()
    retours_m   = df_r_["Débit"].sum() - df_r_["Crédit"].sum()
    remises_m   = df_rem_["Débit"].sum() - df_rem_["Crédit"].sum()
    ca_net_t    = ventes_ht - retours_m - remises_m
    charges_v   = (df_c_["Débit"].sum() - df_c_["Crédit"].sum()) - net_prov_t
    var_stock   = df_stock_["Débit"].sum() - df_stock_["Crédit"].sum()
    marge_b     = ca_net_t - charges_v
    cf          = df_cfi["Débit"].sum()
    res_net     = marge_b - var_stock - cf
    taux_ret_t  = (retours_m / ventes_dist * 100) if ventes_dist else 0
    taux_rem_t  = (remises_m / ventes_dist * 100) if ventes_dist else 0

    # Signal
    if res_net > 0 and taux_ret_t < 20:
        signal_t, bg_t, fg_t = "🟢 Titre rentable", "#d1fae5", "#065f46"
    elif res_net > 0 and taux_ret_t < 35:
        signal_t, bg_t, fg_t = "🟡 Rentabilité à surveiller", "#fef3c7", "#92400e"
    else:
        signal_t, bg_t, fg_t = "🔴 Titre en difficulté", "#fee2e2", "#991b1b"

    st.markdown(f"""<div style='padding:14px 18px;border-radius:12px;background:{bg_t};
        color:{fg_t};font-weight:600;font-size:16px;margin-bottom:14px'>
        {label_affiche(isbn_sel,df)} — {signal_t}</div>""", unsafe_allow_html=True)

    k1,k2,k3,k4 = st.columns(4)
    k1.metric("Taux de retour", f"{taux_ret_t:.1f} %")
    k2.metric("Taux de remise", f"{taux_rem_t:.1f} %")
    k3.metric("Marge brute", f"{fmt_fr(marge_b)} €")
    k4.metric("Résultat net", f"{fmt_fr(res_net)} €")

    if cf == 0 and not st.session_state.get("repartition_active"):
        st.caption("ℹ️ Aucune charge fixe imputée : la répartition des charges indirectes n'a pas été activée.")
    if abs(var_stock) > 0.5:
        st.caption(f"ℹ️ Variation de stock : {fmt_fr(var_stock)} € — isolée de la marge brute, incluse dans le résultat net.")

    # ── Mini SIG ──
    st.markdown("#### Mini SIG — Soldes intermédiaires de gestion")

    detail_charges = params.get("detail_charges")
    mixte          = params.get("contenu_fabrication_mixte") or {}
    mixte_comptes  = tuple(mixte.get("comptes") or [])
    mots_cles_fab  = [m.strip().upper() for m in (mixte.get("mots_cles_fabrication") or []) if m.strip()]

    charge_rows = []
    if detail_charges:
        total_detail = 0.0
        comptes_couverts = []
        for nom_nat, prefixes in detail_charges.items():
            if prefixes and any(str(p)==str(sp) or str(p).startswith(str(sp)) or str(sp).startswith(str(p))
                                for p in prefixes for sp in list(pstock)):
                continue
            pref_eff = (list(prefixes) + list(params.get("provisions_reprises") or [])
                        if nom_nat == "Provision pour retour" else list(prefixes))
            df_nat_pref = df_t[df_t["Compte"].astype(str).str.startswith(tuple(pref_eff))] if pref_eff else df_t.iloc[0:0]
            df_nat_mixte = df_t.iloc[0:0]
            if nom_nat in ("Contenu","Fabrication") and mixte_comptes:
                df_mx = df_t[df_t["Compte"].astype(str).str.startswith(mixte_comptes)]
                if not df_mx.empty and mots_cles_fab:
                    lib_up = df_mx["Libellé"].astype(str).str.upper()
                    mask_fab = lib_up.str.contains("|".join(mots_cles_fab), regex=True)
                else:
                    mask_fab = pd.Series(False, index=df_mx.index)
                df_nat_mixte = df_mx[mask_fab] if nom_nat=="Fabrication" else df_mx[~mask_fab]
                comptes_couverts.extend(list(mixte_comptes))
            if pref_eff or (nom_nat in ("Contenu","Fabrication") and mixte_comptes):
                df_nat = pd.concat([df_nat_pref, df_nat_mixte]) if not df_nat_mixte.empty else df_nat_pref
                val    = df_nat["Débit"].sum() - df_nat["Crédit"].sum()
                comptes_couverts.extend(pref_eff)
            else:
                val = 0.0
            charge_rows.append((f"− {nom_nat}", -val, "deduction"))
            total_detail += val
        reste = charges_v - total_detail
        if abs(reste) > 0.5:
            charge_rows.append(("− Autres charges directes", -reste, "deduction"))
    else:
        charge_rows = [("− Charges variables", -charges_v, "deduction")]

    rows_sig = (
        [("Ventes HT", ventes_ht, "base"),
         ("− Retours", -retours_m, "deduction"),
         ("− Remises", -remises_m, "deduction"),
         ("= CA net",  ca_net_t,  "subtotal")]
        + charge_rows
        + [("= Marge brute", marge_b, "subtotal")]
        + ([("− Variation de stock", -var_stock, "deduction")] if abs(var_stock) > 0.5 else [])
        + [("− Charges fixes imputées", -cf, "deduction"),
           ("= Résultat net du titre", res_net, "total")]
    )

    # Tableau HTML
    html_rows = ""
    for libelle_sig, montant_sig, style_sig in rows_sig:
        if style_sig == "subtotal":
            row_style = "background:#eef2ff;font-weight:600;"
        elif style_sig == "total":
            col_ = "#065f46" if montant_sig >= 0 else "#991b1b"
            fill_ = "#d1fae5" if montant_sig >= 0 else "#fee2e2"
            row_style = f"background:{fill_};font-weight:700;color:{col_};"
        elif style_sig == "deduction":
            row_style = "color:#b91c1c;"
        else:
            row_style = "font-weight:500;"
        html_rows += (f"<tr style='{row_style}'>"
                      f"<td style='padding:7px 12px;border-bottom:1px solid #eee'>{libelle_sig}</td>"
                      f"<td style='padding:7px 12px;border-bottom:1px solid #eee;text-align:right'>"
                      f"{fmt_fr(montant_sig)} €</td></tr>")
    st.markdown(f"""<table style='width:100%;border-collapse:collapse;font-size:14px;
        border-radius:8px;overflow:hidden'>{html_rows}</table>""", unsafe_allow_html=True)

    # Waterfall
    _style_mesure = {"base":"absolute","deduction":"relative","subtotal":"total","total":"total"}
    measures_sig  = [_style_mesure[r[2]] for r in rows_sig]
    fig_sig = go.Figure(go.Waterfall(
        orientation="v", measure=measures_sig,
        x=[r[0] for r in rows_sig], y=[r[1] for r in rows_sig],
        text=[f"{fmt_fr(r[1])} €" for r in rows_sig],
        textposition="outside",
        connector={"line":{"color":"gray","width":0.5}},
        decreasing={"marker":{"color":"#EF4444"}},
        increasing={"marker":{"color":"#10B981"}},
        totals={"marker":{"color":"#3B82F6"}}
    ))
    nb_natures = len(rows_sig)
    fig_sig.update_layout(
        height=300 + nb_natures * 15,
        margin=dict(t=10), yaxis_title="€",
        xaxis_tickangle=-30 if nb_natures > 6 else 0
    )
    st.plotly_chart(fig_sig, use_container_width=True)

    # Évolution mensuelle
    df_t_evol = df_t.copy()
    df_t_evol["Date"] = pd.to_datetime(df_t_evol["Date"], errors="coerce")
    df_t_evol["Mois"] = df_t_evol["Date"].dt.to_period("M").astype(str)
    evol_v = df_t_evol[df_t_evol["Compte"].astype(str).str.startswith(tuple(params["ventes"]))].groupby("Mois")["Crédit"].sum().reset_index()
    evol_r = df_t_evol[mask_retours(df_t_evol,params)].groupby("Mois")["Débit"].sum().reset_index()
    evol = evol_v.merge(evol_r, on="Mois", how="left").fillna(0)
    evol.columns = ["Mois","Ventes","Retours"]
    if not evol.empty:
        st.markdown("#### Évolution mensuelle")
        fig_evol = px.line(evol, x="Mois", y=["Ventes","Retours"], markers=True, height=260)
        fig_evol.update_layout(margin=dict(t=10), legend_title="")
        st.plotly_chart(fig_evol, use_container_width=True)

    # Export Excel fiche titre
    buf_ft = BytesIO()
    df_cr_ft = pd.DataFrame({"Poste":[r[0] for r in rows_sig],"Montant (€)":[r[1] for r in rows_sig]})
    with pd.ExcelWriter(buf_ft, engine="openpyxl") as writer:
        df_cr_ft.to_excel(writer, index=False, sheet_name="Compte_Resultat")
        if not evol.empty:
            evol.to_excel(writer, index=False, sheet_name="Evolution_mensuelle")
    buf_ft.seek(0)
    st.download_button("📥 Exporter la fiche titre (Excel)", buf_ft,
        file_name=f"Fiche_titre_{isbn_sel.replace('/','-')}.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        key=f"export_fiche_{isbn_sel}")

# ── SIMULATEUR ──
elif role == "ec" and page == "🎯 Simulateur de rentabilité":
    df, params = check_pivot()
    dos = get_dossier(st.session_state["dossier_id"])
    st.header(f"🎯 Simulateur de rentabilité — {dos['nom']}")
    titres = sorted(filtrer_isbn_reels(df)["Code_Analytique"].astype(str).unique().tolist())
    if not titres:
        st.warning("Aucun ISBN détecté.")
        st.stop()
    isbn_sim = st.selectbox("Titre à simuler", titres, format_func=lambda c: label_affiche(c,df))
    df_t = df[df["Code_Analytique"]==isbn_sim]
    df_v_ = df_t[df_t["Compte"].astype(str).str.startswith(tuple(params["ventes"]))]
    df_vd = df_t[df_t["Compte"].astype(str).str.startswith(tuple(params.get("ventes_distributeur") or params["ventes"]))]
    df_r_ = df_t[mask_retours(df_t,params)]
    df_rem_ = df_t[mask_remises(df_t,params)]
    pstock = tuple(params.get("stock") or ["603"])
    df_c_ = df_t[df_t["Compte"].astype(str).str.startswith(tuple(params["charges"]))
                  & (~df_t["Compte"].astype(str).str.startswith(pstock))]
    df_cfi = df_t[df_t["Compte"].astype(str)==COMPTE_CHARGES_INDIRECTES_REPARTIES]

    vb  = df_v_["Crédit"].sum(); vd = df_vd["Crédit"].sum()
    rb  = df_r_["Débit"].sum()-df_r_["Crédit"].sum()
    rem = df_rem_["Débit"].sum()-df_rem_["Crédit"].sum()
    cv  = df_c_["Débit"].sum()-df_c_["Crédit"].sum()
    cf  = df_cfi["Débit"].sum()
    tr  = (rb/vd*100) if vd else 0

    col1,col2 = st.columns(2)
    with col1:
        tr_hyp = st.slider("Taux de retour hypothétique (%)", 0.0, 100.0, float(round(tr,1)), 0.5)
        var_ch = st.slider("Variation charges variables (%)", -50, 50, 0, 1)
    with col2:
        nb_t = (st.session_state.get("repartition_detail") or {}).get("nb_titres_actifs", len(titres))
        nb_t_h = st.slider("Nombre titres actifs (clé répartition)", 1, max(nb_t*2,nb_t+5), int(nb_t))
        total_ci = (st.session_state.get("repartition_detail") or {}).get("part_charge",0)*nb_t

    ret_h = vd*tr_hyp/100; cv_h = cv*(1+var_ch/100)
    ca_h  = vb-ret_h-rem; mb_h = ca_h-cv_h
    cf_h  = (total_ci/nb_t_h) if (nb_t_h and st.session_state.get("repartition_active")) else cf
    res_h = mb_h-cf_h

    st.divider()
    c1,c2,c3,c4 = st.columns(4)
    c1.metric("CA net (simulé)",f"{fmt_fr(ca_h)} €",delta=f"{fmt_fr(ca_h-(vb-rb-rem))} €")
    c2.metric("Marge brute (simulée)",f"{fmt_fr(mb_h)} €",delta=f"{fmt_fr(mb_h-(vb-rb-rem-cv))} €")
    c3.metric("Résultat net (simulé)",f"{fmt_fr(res_h)} €",delta=f"{fmt_fr(res_h-(vb-rb-rem-cv-cf))} €")
    c4.metric("Charges fixes (simulées)",f"{fmt_fr(cf_h)} €")

    if vd:
        seuil = (vb-rem-cv_h-cf_h)/vd*100
        seuil = max(0,seuil)
        marge_s = seuil-tr_hyp
        if marge_s>=0:
            st.success(f"✅ Seuil de rentabilité : taux de retour max **{seuil:.1f}%** (marge de {marge_s:.1f} pts)")
        else:
            st.error(f"❌ Déficitaire — seuil d'équilibre à **{seuil:.1f}%** (vs {tr_hyp:.1f}% simulé)")


# ── TRÉSORERIE (EC) ──
elif role == "ec" and page == "💰 Trésorerie prévisionnelle":
    df, params = check_pivot()
    dos = get_dossier(st.session_state["dossier_id"])
    st.header(f"💰 Trésorerie prévisionnelle — {dos['nom']}")
    df["Date"] = pd.to_datetime(df["Date"], errors="coerce")
    df["Débit"] = pd.to_numeric(df["Débit"], errors="coerce").fillna(0)
    df["Crédit"] = pd.to_numeric(df["Crédit"], errors="coerce").fillna(0)
    if "Journal" not in df.columns: df["Journal"] = ""
    df = df.dropna(subset=["Date"])
    if df.empty:
        st.warning("Aucune écriture datée.")
        st.stop()

    col1, col2 = st.columns(2)
    with col1:
        date_debut = st.date_input("Date de départ", df["Date"].min())
        tresorerie_ouv = st.number_input("Trésorerie à l'ouverture (€)", value=0.0, step=100.0)
    with col2:
        horizon = st.slider("Horizon projection (mois)", 0, 24, 6)
        comptes_banque = tuple(x.strip() for x in st.text_input("Comptes trésorerie","512,530,580").split(",") if x.strip())
        journaux_exclus = tuple(x.strip() for x in st.text_input("Journaux exclus (AN)","AN").split(",") if x.strip())

    mask_an = df["Journal"].isin(journaux_exclus) if journaux_exclus else pd.Series(False, index=df.index)
    df_flux = df[~mask_an & (df["Date"]>=pd.to_datetime(date_debut))].copy()
    df_flux["Mois"] = df_flux["Date"].dt.to_period("M")

    if df_flux.empty:
        st.warning("Aucune écriture après la date de départ.")
        st.stop()

    mois = sorted(df_flux["Mois"].unique())
    encaissements = df_flux[df_flux["Compte"].astype(str).str.startswith("41")].groupby("Mois")["Crédit"].sum()
    decaissements = df_flux[df_flux["Compte"].astype(str).str.startswith(("40","42","43","44"))].groupby("Mois")["Débit"].sum()
    flux_net = encaissements.sub(decaissements, fill_value=0).reindex(mois, fill_value=0)
    treso = tresorerie_ouv + flux_net.cumsum()

    st.session_state["treso_real"] = treso
    st.session_state["treso_ouverture"] = tresorerie_ouv

    m1,m2,m3 = st.columns(3)
    m1.metric("Trésorerie ouverture", f"{fmt_fr(tresorerie_ouv)} €")
    m2.metric("Trésorerie clôture (réalisé)", f"{fmt_fr(treso.iloc[-1])} €")
    m3.metric("Flux net généré", f"{fmt_fr(flux_net.sum())} €")

    _MOIS_FR = {1:"Jan",2:"Fév",3:"Mar",4:"Avr",5:"Mai",6:"Jui",7:"Jul",8:"Aoû",9:"Sep",10:"Oct",11:"Nov",12:"Déc"}
    def ml(p): return f"{_MOIS_FR.get(p.month,str(p.month))} {p.year}"

    fig_t = go.Figure()
    fig_t.add_trace(go.Scatter(x=[ml(m) for m in treso.index], y=treso.values,
                               name="Réalisé", line=dict(color="#111827",width=2.5)))

    if horizon > 0:
        base_enc = encaissements.iloc[-3:].mean() if len(encaissements)>=3 else (encaissements.mean() if len(encaissements) else 0)
        base_dec = decaissements.iloc[-3:].mean() if len(decaissements)>=3 else (decaissements.mean() if len(decaissements) else 0)
        col_sc1,col_sc2,col_sc3,col_sc4 = st.columns(4)
        t_opt  = col_sc1.number_input("Croissance encaissements optimiste (%/mois)", value=4.0, step=0.5)/100
        t_cent = col_sc2.number_input("Croissance encaissements central (%/mois)", value=2.0, step=0.5)/100
        t_pess = col_sc3.number_input("Croissance encaissements pessimiste (%/mois)", value=0.0, step=0.5)/100
        t_ch   = col_sc4.number_input("Évolution charges (%/mois)", value=1.0, step=0.5)/100
        for nom, tc, col_hex in [("Optimiste",t_opt,"#10B981"),("Central",t_cent,"#3B82F6"),("Pessimiste",t_pess,"#EF4444")]:
            futurs = [mois[-1]+i for i in range(1,horizon+1)]
            enc_f = base_enc; dec_f = base_dec; proj = []
            for m_f in futurs:
                enc_f *= (1+tc); dec_f *= (1+t_ch); proj.append(enc_f-dec_f)
            treso_proj = treso.iloc[-1] + pd.Series(proj, index=futurs).cumsum()
            treso_full = pd.concat([treso.iloc[[-1]], treso_proj])
            fig_t.add_trace(go.Scatter(x=[ml(m) for m in treso_full.index], y=treso_full.values,
                                       name=nom, line=dict(color=col_hex,width=2,dash="dot")))

    fig_t.add_hline(y=0, line_dash="dash", line_color="gray")
    fig_t.update_layout(height=380, margin=dict(t=20), xaxis_title="", yaxis_title="€",
                        legend=dict(orientation="h"))
    st.plotly_chart(fig_t, use_container_width=True)

# ── DROITS D'AUTEURS (EC) ──
elif role == "ec" and page == "✍️ Droits d'auteurs":
    df, params = check_pivot()
    dos = get_dossier(st.session_state["dossier_id"])
    st.header(f"✍️ Droits d'auteurs — {dos['nom']}")
    st.info("Module complet disponible — configurez les comptes ci-dessous.")
    st.info(
        "**Comment fonctionne ce module ?** "
        "Les droits d'auteurs déjà comptabilisés sont lus directement dans votre grand livre. "
        "Un simulateur est disponible pour estimer avant provision."
    )

    ong1, ong2, ong3, ong4 = st.tabs([
        "📋 Référentiel contrats",
        "🧮 Simulateur",
        "📒 Réel (comptabilisé)",
        "📄 Relevés par auteur"
    ])

    # ── Initialisation référentiel ──
    if "royalties_referentiel" not in st.session_state:
        st.session_state["royalties_referentiel"] = []

    with ong1:
        st.subheader("Saisie des contrats auteurs")
        col_a, col_b = st.columns(2)
        with col_a:
            inp_auteur = st.text_input("Nom de l'auteur", key="da_auteur")
            inp_isbn   = st.text_input("ISBN (code analytique exact du socle)", key="da_isbn")
            inp_titre  = st.text_input("Titre du livre", key="da_titre")
        with col_b:
            inp_part   = st.number_input("Part de cet auteur (%)", min_value=0.0, max_value=100.0, value=100.0, step=1.0, key="da_part")
            inp_statut = st.selectbox("Statut fiscal", ["Traitements et salaires (option assimilé)","BNC (droits d'auteur)"], key="da_statut")
            inp_taux   = st.number_input("Taux forfaitaire (%)", value=10.0, step=0.5, key="da_taux")
        if st.button("➕ Ajouter ce contrat", key="da_add"):
            if inp_auteur and inp_isbn:
                st.session_state["royalties_referentiel"].append({
                    "auteur": inp_auteur, "isbn": inp_isbn.strip(),
                    "titre": inp_titre or inp_isbn, "part": inp_part,
                    "statut_fiscal": inp_statut,
                    "paliers": [{"seuil": 0, "taux": inp_taux}]
                })
                st.success(f"✅ Contrat ajouté : {inp_auteur} / {inp_titre or inp_isbn}")
        st.divider()
        ref = st.session_state["royalties_referentiel"]
        if ref:
            rows_ref = [{"Auteur":c["auteur"],"ISBN":c["isbn"],"Titre":c["titre"],
                         "Part (%)":c["part"],"Taux (%)":c["paliers"][0]["taux"],
                         "Statut":c.get("statut_fiscal","")} for c in ref]
            st.dataframe(pd.DataFrame(rows_ref), use_container_width=True, hide_index=True)
            idx_s = st.number_input("Supprimer le contrat n°", min_value=0, max_value=max(0,len(ref)-1), step=1, key="da_del_idx")
            if st.button("🗑️ Supprimer", key="da_del"):
                st.session_state["royalties_referentiel"].pop(int(idx_s)); st.rerun()
        else:
            st.info("Aucun contrat enregistré.")

    with ong2:
        st.subheader("Simulateur — droits d'auteur & précompte URSSAF")
        if "df_pivot" not in st.session_state or not st.session_state["royalties_referentiel"]:
            st.warning("⚠️ Générez d'abord le socle et saisissez au moins un contrat dans l'onglet Référentiel.")
        else:
            df_sim2 = df.copy()
            params_sim2 = params
            resultats_sim = []
            for contrat in st.session_state["royalties_referentiel"]:
                isbn = contrat["isbn"]
                mask_v2 = df_sim2["Compte"].astype(str).str.startswith(tuple(params_sim2["ventes"]))
                mask_r2 = df_sim2["Compte"].astype(str).str.startswith(tuple(params_sim2["retours"]))
                mask_rem2 = df_sim2["Compte"].astype(str).str.startswith(tuple(params_sim2.get("remises",[])))
                df_i2 = df_sim2[df_sim2["Code_Analytique"]==isbn]
                ca_b2  = df_i2[mask_v2.reindex(df_i2.index,fill_value=False)]["Crédit"].sum()
                ret2   = df_i2[mask_r2.reindex(df_i2.index,fill_value=False)]["Débit"].sum()
                rem2   = df_i2[mask_rem2.reindex(df_i2.index,fill_value=False)]["Débit"].sum()
                ca_n2  = ca_b2 - ret2 - rem2
                base2  = max(ca_n2, 0)
                taux2  = contrat["paliers"][0]["taux"] / 100
                droits_bruts = base2 * taux2 * contrat["part"] / 100
                resultats_sim.append({
                    "Auteur": contrat["auteur"], "Titre": contrat["titre"],
                    "Statut": contrat.get("statut_fiscal",""),
                    "CA net (€)": round(ca_n2,2),
                    "Base calcul (€)": round(base2,2),
                    "Droits bruts (€)": round(droits_bruts,2),
                })
            df_sim_res = pd.DataFrame(resultats_sim)
            cols_s2 = ["CA net (€)","Base calcul (€)","Droits bruts (€)"]
            st.dataframe(df_sim_res.style.format({c:(lambda x: f"{fmt_fr(x,2)} €") for c in cols_s2}), use_container_width=True)
            total_bruts = df_sim_res["Droits bruts (€)"].sum()
            st.metric("💰 Total droits bruts estimés", f"{fmt_fr(total_bruts,2)} €")
            st.divider()
            st.markdown("**Précompte URSSAF estimé**")
            col_u1,col_u2,col_u3,col_u4 = st.columns(4)
            taux_csg  = col_u1.number_input("CSG+CRDS (%)", value=9.70, step=0.01, key="da_csg")
            taux_fp   = col_u2.number_input("Formation prof. (%)", value=1.00, step=0.01, key="da_fp")
            taux_raap = col_u3.number_input("RAAP (%)", value=0.0, step=0.01, key="da_raap")
            taux_diff = col_u4.number_input("Contribution diffuseur (%)", value=1.10, step=0.05, key="da_diff_taux")
            est_ts = df_sim_res["Statut"].str.startswith("Traitements")
            assiette = df_sim_res["Droits bruts (€)"] * 0.9825
            precompte = np.where(est_ts, assiette*(taux_csg+taux_fp+taux_raap)/100, 0.0)
            contrib   = df_sim_res["Droits bruts (€)"] * taux_diff / 100
            total_urssaf_s = precompte.sum() + contrib.sum()
            cu1,cu2,cu3 = st.columns(3)
            cu1.metric("Précompte URSSAF (TS)", f"{fmt_fr(precompte.sum(),2)} €")
            cu2.metric("Contribution diffuseur", f"{fmt_fr(contrib.sum(),2)} €")
            cu3.metric("Total à reverser URSSAF", f"{fmt_fr(total_urssaf_s,2)} €")

    with ong3:
        st.subheader("Montants réellement comptabilisés")
        df["Date"] = pd.to_datetime(df["Date"], errors="coerce")
        col1_d,col2_d = st.columns(2)
        with col1_d:
            date_debut_d = st.date_input("Période du", df["Date"].dropna().min().date() if not df["Date"].dropna().empty else _dt.date.today(), key="da_d")
            compte_db  = st.text_input("Droits bruts (charge)", value="604300000", key="da_cpt_db")
            compte_urs = st.text_input("URSSAF à payer", value="438106", key="da_cpt_urs")
        with col2_d:
            date_fin_d = st.date_input("Au", df["Date"].dropna().max().date() if not df["Date"].dropna().empty else _dt.date.today(), key="da_f")
            compte_diff_d = st.text_input("Contribution diffuseur", value="645106", key="da_cpt_diff")
            compte_net_d  = st.text_input("Droits à payer (net)", value="408106", key="da_cpt_net")
            compte_av_d   = st.text_input("À-valoirs (bilan)", value="409600", key="da_cpt_av")

        mask_p = (df["Date"]>=pd.to_datetime(date_debut_d)) & (df["Date"]<=pd.to_datetime(date_fin_d))
        df_p = df[mask_p]

        def par_isbn_d(cpt, sens):
            m = df_p["Compte"].astype(str).str.strip()==str(cpt).strip()
            if not m.any(): return pd.Series(dtype=float)
            g = df_p[m].groupby("Code_Analytique")
            return (g["Débit"].sum()-g["Crédit"].sum()) if sens=="debit" else (g["Crédit"].sum()-g["Débit"].sum())

        db_s   = par_isbn_d(compte_db,"debit")
        urs_s  = par_isbn_d(compte_urs,"credit")
        diff_s = par_isbn_d(compte_diff_d,"debit")
        net_s  = par_isbn_d(compte_net_d,"credit")
        av_s   = par_isbn_d(compte_av_d,"debit")

        isbns_d = sorted(set(db_s.index)|set(urs_s.index)|set(diff_s.index)|set(net_s.index))
        isbns_d = [i for i in isbns_d if str(i).strip() not in ("","CHARGES INDIRECTES","PRODUITS INDIRECTS")]

        if not isbns_d:
            st.info("Aucune écriture trouvée sur ces comptes pour la période.")
        else:
            lignes_d = []
            for isbn in isbns_d:
                lignes_d.append({
                    "ISBN": isbn,
                    "Droits bruts (€)":         round(float(db_s.get(isbn,0)),2),
                    "Contribution diffuseur (€)": round(float(diff_s.get(isbn,0)),2),
                    "Précompte URSSAF (€)":      round(float(urs_s.get(isbn,0)),2),
                    "Net dû auteur (€)":         round(float(net_s.get(isbn,0)),2),
                    "À-valoir restant (€)":      round(float(av_s.get(isbn,0)),2),
                })
            df_d = pd.DataFrame(lignes_d)
            cols_m = ["Droits bruts (€)","Contribution diffuseur (€)","Précompte URSSAF (€)","Net dû auteur (€)","À-valoir restant (€)"]
            st.dataframe(df_d.style.format({c:(lambda x: f"{fmt_fr(x,2)} €") for c in cols_m}), use_container_width=True)
            total_db  = df_d["Droits bruts (€)"].sum()
            total_urs = df_d["Précompte URSSAF (€)"].sum()
            total_diff_t = df_d["Contribution diffuseur (€)"].sum()
            total_net = df_d["Net dû auteur (€)"].sum()
            c1,c2,c3,c4 = st.columns(4)
            c1.metric("Droits bruts", f"{fmt_fr(total_db,2)} €")
            c2.metric("Précompte URSSAF", f"{fmt_fr(total_urs,2)} €")
            c3.metric("Contribution diffuseur", f"{fmt_fr(total_diff_t,2)} €")
            c4.metric("Net dû aux auteurs", f"{fmt_fr(total_net,2)} €")
            st.markdown(f"**🏛️ Total à reverser à l'URSSAF : {fmt_fr(total_urs+total_diff_t,2)} €**")
            st.session_state["df_royalties_reel"] = df_d
            buf_d = BytesIO()
            with pd.ExcelWriter(buf_d, engine="openpyxl") as writer:
                df_d.to_excel(writer, index=False, sheet_name="Droits_auteurs")
            buf_d.seek(0)
            st.download_button("📥 Exporter (Excel)", buf_d,
                file_name=f"Droits_auteurs_{date_debut_d}_{date_fin_d}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                key="da_export")

    with ong4:
        st.subheader("Relevés de droits par auteur")
        if "df_royalties_reel" not in st.session_state or st.session_state["df_royalties_reel"].empty:
            st.warning("⚠️ Effectuez d'abord la lecture dans l'onglet 📒 Réel (comptabilisé).")
        else:
            df_rel = st.session_state["df_royalties_reel"].copy()
            auteurs_list = ["Tous"]
            ref_map = {c["isbn"]: c["auteur"] for c in st.session_state.get("royalties_referentiel",[])}
            if ref_map:
                df_rel["Auteur"] = df_rel["ISBN"].map(ref_map).fillna("(auteur non identifié)")
                auteurs_list += sorted(df_rel["Auteur"].unique().tolist())
            auteur_sel = st.selectbox("Sélectionner un auteur", auteurs_list, key="da_auteur_sel")
            df_aff = df_rel if auteur_sel=="Tous" else df_rel[df_rel["Auteur"]==auteur_sel]
            cols_rel = [c for c in df_aff.columns if c in ["Auteur","ISBN","Droits bruts (€)","Contribution diffuseur (€)","Précompte URSSAF (€)","Net dû auteur (€)","À-valoir restant (€)"]]
            cols_m_rel = [c for c in cols_rel if "(€)" in c]
            st.dataframe(df_aff[cols_rel].style.format({c:(lambda x: f"{fmt_fr(x,2)} €") for c in cols_m_rel}), use_container_width=True, hide_index=True)
            buf_rel = BytesIO()
            with pd.ExcelWriter(buf_rel, engine="openpyxl") as writer:
                df_aff[cols_rel].to_excel(writer, index=False, sheet_name="Releve")
            buf_rel.seek(0)
            st.download_button(f"📥 Télécharger relevé — {auteur_sel}", buf_rel,
                file_name=f"Royalties_{auteur_sel.replace(' ','_')}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                key="da_export_releve")

# ── RETOURS & REMISES (EC) ──
elif role == "ec" and page == "📦 Retours & Remises":
    df, params = check_pivot()
    dos = get_dossier(st.session_state["dossier_id"])
    st.header(f"📦 Retours & Remises — {dos['nom']}")
    df["Date"] = pd.to_datetime(df["Date"], errors="coerce")
    df["Mois"] = df["Date"].dt.strftime("%Y-%m")
    seuil = st.sidebar.number_input("Seuil alerte (%)", value=25, step=5)

    df_v_  = df[df["Compte"].astype(str).str.startswith(tuple(params["ventes"]))]
    df_r_  = df[mask_retours(df,params)]
    df_rem_= df[mask_remises(df,params)]
    pref_d = tuple(params.get("ventes_distributeur") or params["ventes"])
    df_vd_ = df[df["Compte"].astype(str).str.startswith(pref_d)]

    tv = df_v_["Crédit"].sum(); vd = df_vd_["Crédit"].sum()
    tr = abs(df_r_["Débit"].sum()-df_r_["Crédit"].sum()) if not df_r_.empty else 0
    rem= abs(df_rem_["Débit"].sum()-df_rem_["Crédit"].sum()) if not df_rem_.empty else 0
    taux_r = (tr/vd*100) if vd else 0; taux_rem_pct = (rem/vd*100) if vd else 0

    c1,c2,c3,c4 = st.columns(4)
    c1.metric("CA brut",f"{fmt_fr(tv)} €")
    c2.metric("Total retours",f"{fmt_fr(tr)} €")
    c3.metric("Taux de retour",f"{taux_r:.1f} %",
              delta="⚠️ Dépasse le seuil" if taux_r>seuil else "✅ Normal",
              delta_color="inverse" if taux_r>seuil else "normal")
    c4.metric("Taux de remise",f"{taux_rem_pct:.1f} %")

    if taux_r > seuil:
        st.error(f"🚨 Taux de retour ({taux_r:.1f}%) dépasse votre seuil de {seuil}% !")

    if not df_r_.empty:
        trend_r = df_r_.groupby("Mois")["Débit"].sum().abs().reset_index()
        fig_r = px.bar(trend_r, x="Mois", y="Débit", title="Retours mensuels (€)",
                       color_discrete_sequence=["#EF4444"], height=300)
        st.plotly_chart(fig_r, use_container_width=True)

        df_ri = filtrer_isbn_reels(df_r_)
        if not df_ri.empty:
            ret_isbn = df_ri.groupby("Code_Analytique")["Débit"].sum().abs().reset_index().sort_values("Débit",ascending=False)
            st.subheader("Retours par titre")
            st.dataframe(ret_isbn, hide_index=True)

# ── SYNTHÈSE FINANCIÈRE (EC) ──
elif role == "ec" and page == "📊 Synthèse financière":
    df, params = check_pivot()
    dos = get_dossier(st.session_state["dossier_id"])
    st.header(f"📊 Synthèse financière — {dos['nom']}")

    df_v_   = df[mask_ventes(df,params)]
    df_r_   = df[mask_retours(df,params)]
    df_rem_ = df[mask_remises(df,params)]
    df_c_   = df[mask_charges(df,params)]
    df_prov = df[mask_provisions_reprises(df,params)]
    net_p   = df_prov["Crédit"].sum()-df_prov["Débit"].sum()

    ca_b  = df_v_["Crédit"].sum(); corr = df_v_["Débit"].sum()
    tr_   = df_r_["Débit"].sum()-df_r_["Crédit"].sum()
    rem_  = df_rem_["Débit"].sum()-df_rem_["Crédit"].sum()
    ca_n  = ca_b-tr_-rem_-corr
    ch_t  = (df_c_["Débit"].sum()-df_c_["Crédit"].sum())-net_p
    res_  = ca_n-ch_t
    mb_p  = (res_/ca_b*100) if ca_b else 0

    soldes   = [ca_b,-tr_,-rem_,-corr,ca_n,-ch_t,res_]
    libelles = ["CA brut","Retours","Remises","Corrections","CA net","Charges","Résultat net"]

    col1,col2 = st.columns([1.2,1])
    with col1:
        st.subheader("Compte de résultat synthétique")
        df_sum = pd.DataFrame({"Poste":libelles,"Montant (€)":soldes})
        st.dataframe(df_sum.style.format({"Montant (€)":(lambda x: fmt_fr(x))}), hide_index=True)
        st.metric("Taux de marge nette",f"{mb_p:.1f} %")
    with col2:
        fig_w = go.Figure(go.Waterfall(
            orientation="v",
            measure=["absolute","relative","relative","relative","total","relative","total"],
            x=libelles, y=soldes,
            decreasing={"marker":{"color":"#EF4444"}},
            increasing={"marker":{"color":"#10B981"}},
            totals={"marker":{"color":"#3B82F6"}}
        ))
        fig_w.update_layout(height=360, margin=dict(t=20))
        st.plotly_chart(fig_w, use_container_width=True)

    buf_s = BytesIO()
    with pd.ExcelWriter(buf_s, engine="openpyxl") as writer:
        df_sum.to_excel(writer, index=False, sheet_name="Synthese")
    buf_s.seek(0)
    st.download_button("📥 Exporter la synthèse (Excel)", buf_s,
        file_name="Synthese_Financiere.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")


# ── ASSISTANT IA (EC) ──
elif role == "ec" and page == "🤖 Assistant IA":
    df_ia, params_ia = check_pivot()
    dos = get_dossier(st.session_state["dossier_id"])
    st.header(f"🤖 Assistant IA — {dos['nom']}")
    client_ai = get_client_ai()
    if client_ai is None:
        st.error("⚠️ Clé API Anthropic non configurée (ANTHROPIC_API_KEY dans les secrets Streamlit).")
        st.stop()

    # Contexte données
    df_v_ = df_ia[mask_ventes(df_ia,params_ia)]
    df_r_ = df_ia[mask_retours(df_ia,params_ia)]
    ca_b_ = df_v_["Crédit"].sum(); tr_ = df_r_["Débit"].sum()-df_r_["Crédit"].sum()
    ca_n_ = ca_b_ - tr_
    taux_r_ = (tr_/ca_b_*100) if ca_b_ else 0
    nb_isbn_ = filtrer_isbn_reels(df_ia)["Code_Analytique"].nunique()

    data_ctx = f"""
DONNÉES ANALYTIQUES — {dos['nom']} — {dos.get('exercice','')}
CA brut : {fmt_fr(ca_b_)} EUR
Total retours : {fmt_fr(tr_)} EUR
CA net : {fmt_fr(ca_n_)} EUR
Taux de retour : {taux_r_:.1f} %
Nombre de titres actifs : {nb_isbn_}
"""
    sys_prompt = SYSTEM_PROMPT_EC + f"\n\nDONNÉES DU DOSSIER :\n{data_ctx}"

    for msg in st.session_state["messages_agent"]:
        with st.chat_message(msg["role"], avatar="🧑" if msg["role"]=="user" else "🤖"):
            st.markdown(msg["content"])

    if prompt_u := st.chat_input("Posez votre question..."):
        st.session_state["messages_agent"].append({"role":"user","content":prompt_u})
        with st.chat_message("user", avatar="🧑"):
            st.markdown(prompt_u)
        with st.chat_message("assistant", avatar="🤖"):
            with st.spinner("Analyse en cours..."):
                resp = client_ai.messages.create(
                    model="claude-sonnet-4-6", max_tokens=800,
                    system=sys_prompt, messages=st.session_state["messages_agent"]
                )
                answer = resp.content[0].text
                st.markdown(answer)
                st.session_state["messages_agent"].append({"role":"assistant","content":answer})

    if st.button("🗑️ Effacer"):
        st.session_state["messages_agent"] = []
        st.rerun()

# ── RAPPORT DE PILOTAGE (EC) ──
elif role == "ec" and page == "📄 Rapport de pilotage":
    df, params = check_pivot()
    dos = get_dossier(st.session_state["dossier_id"])
    st.header("📄 Rapport de pilotage analytique")
    st.markdown(f"*{dos['nom']} — {dos.get('exercice','')}*")

    if not DOCX_AVAILABLE:
        st.error("❌ python-docx non installé. Ajoutez `python-docx` dans requirements.txt.")
        st.stop()

    col1,col2 = st.columns(2)
    with col1:
        nom_editeur = st.text_input("Maison d'édition", value=dos.get("nom",""))
        nom_dirigeant = st.text_input("Nom du dirigeant", value=dos.get("dirigeant",""))
        exercice = st.text_input("Exercice", value=dos.get("exercice","2025/2026"))
    with col2:
        date_rapport = st.date_input("Date du rapport", value=_dt.date.today())
        periode_debut = st.text_input("Début de période", value=dos.get("periode_debut","01/04/2025"))
        periode_fin   = st.text_input("Fin de période",   value=dos.get("periode_fin","31/03/2026"))

    mode_anon = st.checkbox("🕶️ Anonymiser les titres", value=False)
    nb_top = st.slider("Nombre de titres Top / Flop", 3, 10, 5)

    titres_r = sorted(filtrer_isbn_reels(df)["Code_Analytique"].astype(str).unique().tolist())
    if not titres_r:
        st.error("Aucun ISBN détecté.")
        st.stop()

    res_r = calculer_indicateurs_titres(df, params, titres_r)
    mapping_anon = obtenir_mapping_anonymisation(df)
    res_r["Titre"] = res_r["Code_Analytique"].apply(lambda c: mapping_anon.get(c,c) if mode_anon else c)

    df_v_ = df[mask_ventes(df,params)]; df_c_ = df[mask_charges(df,params)]
    df_r_ = df[mask_retours(df,params)]; df_rem_ = df[mask_remises(df,params)]
    df_prov = df[mask_provisions_reprises(df,params)]
    net_p = df_prov["Crédit"].sum()-df_prov["Débit"].sum()
    ca_b_ = df_v_["Crédit"].sum()-df_v_["Débit"].sum()
    tr_   = df_r_["Débit"].sum()-df_r_["Crédit"].sum()
    rem_  = df_rem_["Débit"].sum()-df_rem_["Crédit"].sum()
    ca_n_ = ca_b_-tr_-rem_
    ch_t_ = (df_c_["Débit"].sum()-df_c_["Crédit"].sum())-net_p
    mn_   = ca_n_-ch_t_
    mb_   = res_r["Marge brute"].sum()
    taux_mb_ = (mb_/ca_b_*100) if ca_b_ else 0
    n_pos_ = (res_r["Marge brute"]>0).sum()
    n_neg_ = (res_r["Marge brute"]<=0).sum()
    nb_eans_ = len(titres_r)
    res_sorted_ = res_r.sort_values("Marge brute",ascending=False).reset_index(drop=True)

    st.divider()
    m1,m2,m3,m4 = st.columns(4)
    m1.metric("CA net", f"{fmt_fr(ca_n_)} €")
    m2.metric("Marge brute", f"{fmt_fr(mb_)} €", f"{taux_mb_:.1f} %")
    m3.metric("Résultat net", f"{fmt_fr(mn_)} €")
    m4.metric("Titres déficitaires", f"{n_neg_} / {nb_eans_}")

    if st.button("🖨️ Générer le rapport Word", type="primary", use_container_width=True):
        with st.spinner("Génération en cours..."):
            try:
                _MR = RGBColor(0x1F,0x4E,0x79); _TR = RGBColor(0xC0,0x52,0x2A)
                _WH = RGBColor(0xFF,0xFF,0xFF)
                _GC = "E2EFDA"; _RC = "FFDDD9"; _YC = "FFF2CC"
                _BC = "D6E4F0"; _LC = "F5E6DF"

                def _shd(cell, h):
                    tc=cell._tc; p=tc.get_or_add_tcPr()
                    for o in p.findall(qn("w:shd")): p.remove(o)
                    s=OxmlElement("w:shd"); s.set(qn("w:val"),"clear")
                    s.set(qn("w:color"),"auto"); s.set(qn("w:fill"),h); p.append(s)

                def _brd(cell, c="CCCCCC", z="4"):
                    tc=cell._tc; p=tc.get_or_add_tcPr()
                    for o in p.findall(qn("w:tcBorders")): p.remove(o)
                    b=OxmlElement("w:tcBorders")
                    for e in ["top","left","bottom","right"]:
                        x=OxmlElement(f"w:{e}"); x.set(qn("w:val"),"single")
                        x.set(qn("w:sz"),z); x.set(qn("w:color"),c); b.append(x)
                    p.append(b)

                def _wc(cell, text, bold=False, italic=False, color=None,
                         size=9, center=False, fill=None):
                    if fill: _shd(cell, fill)
                    _brd(cell); cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    p=cell.paragraphs[0]; p.clear()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after  = Pt(2)
                    r=p.add_run(str(text)); r.bold=bold; r.italic=italic
                    r.font.size=Pt(size); r.font.name="Arial"
                    r.font.color.rgb = color if color else RGBColor(0,0,0)

                def _P(d, text, bold=False, italic=False, size=10,
                        color=None, align="left", sb=6, sa=6, line=None):
                    p=d.add_paragraph()
                    p.paragraph_format.space_before = Pt(sb)
                    p.paragraph_format.space_after  = Pt(sa)
                    if line: p.paragraph_format.line_spacing = Pt(line)
                    p.alignment = {
                        "left":WD_ALIGN_PARAGRAPH.LEFT,
                        "center":WD_ALIGN_PARAGRAPH.CENTER,
                        "right":WD_ALIGN_PARAGRAPH.RIGHT,
                        "justify":WD_ALIGN_PARAGRAPH.JUSTIFY
                    }.get(align, WD_ALIGN_PARAGRAPH.LEFT)
                    r=p.add_run(text); r.bold=bold; r.italic=italic
                    r.font.size=Pt(size); r.font.name="Arial"
                    if color: r.font.color.rgb=color
                    return p

                def _sep(d, col="C0522A"):
                    p=d.add_paragraph()
                    p.paragraph_format.space_before=Pt(2); p.paragraph_format.space_after=Pt(6)
                    pr=p._p.get_or_add_pPr(); pb=OxmlElement("w:pBdr")
                    bt=OxmlElement("w:bottom"); bt.set(qn("w:val"),"single")
                    bt.set(qn("w:sz"),"6"); bt.set(qn("w:color"),col)
                    pb.append(bt); pr.append(pb)

                def _fn(n, sign=False):
                    return ("+" if (sign and n>=0) else "") + f"{n:,.0f} EUR".replace(",","\u202f")

                def _fp(n): return f"{'+'if n>=0 else ''}{n:.1f}%"

                def _hr(tbl, cols, fill="1F4E79"):
                    row = tbl.rows[0]
                    for i,(t,w) in enumerate(cols):
                        c=row.cells[i]; _wc(c,t,bold=True,color=_WH,size=9,center=True,fill=fill)
                        if w: c.width=Cm(w)

                # Calculs complémentaires pour le rapport complet
                mask_ci_rp = df["Code_Analytique"].astype(str).str.upper().isin(
                    ["CHARGES INDIRECTES","FRAIS GENERAUX"])
                mask_ci_cpt_rp = df["Compte"].astype(str) == COMPTE_CHARGES_INDIRECTES_REPARTIES
                ci_total_rp = (df[mask_ci_rp | mask_ci_cpt_rp]["Débit"].sum()
                               - df[mask_ci_rp | mask_ci_cpt_rp]["Crédit"].sum())
                ci_par_ean_rp = ci_total_rp / nb_eans_ if nb_eans_ > 0 else 0

                top_list_  = [res_sorted_.iloc[i] for i in range(min(nb_top, len(res_sorted_)))]
                flop_list_ = [res_sorted_.iloc[-(i+1)] for i in range(min(nb_top, len(res_sorted_)))]

                # ── Document ──
                d = Document()
                for s in d.sections:
                    s.top_margin=Cm(2); s.bottom_margin=Cm(2)
                    s.left_margin=Cm(2.5); s.right_margin=Cm(2)

                # En-tête cabinet
                _P(d,NOM_CABINET,bold=True,size=14,color=_MR,sb=0,sa=2)
                _P(d,"Expert-comptable — Maisons d'édition indépendantes",size=10,
                   color=RGBColor(0x55,0x55,0x55),sb=0,sa=2)
                _P(d,"Membre de l'Ordre des Experts-Comptables",size=9,italic=True,
                   color=RGBColor(0x88,0x88,0x88),sb=0,sa=4)

                _mois_fr_ = {1:"janvier",2:"février",3:"mars",4:"avril",5:"mai",6:"juin",
                              7:"juillet",8:"août",9:"septembre",10:"octobre",11:"novembre",12:"décembre"}
                date_fr_  = f"{date_rapport.day} {_mois_fr_[date_rapport.month]} {date_rapport.year}"
                pr_=d.add_paragraph(); pr_.alignment=WD_ALIGN_PARAGRAPH.RIGHT
                pr_.paragraph_format.space_before=Pt(0); pr_.paragraph_format.space_after=Pt(4)
                rr_=pr_.add_run(f"Lille, le {date_fr_}")
                rr_.font.size=Pt(10); rr_.font.name="Arial"
                _sep(d)

                _P(d,f"{nom_dirigeant}",bold=True,size=11,sb=4,sa=2)
                _P(d,nom_editeur,size=11,sb=0,sa=2)
                _P(d,"Siège social : Lille (59)",size=10,
                   color=RGBColor(0x55,0x55,0x55),sb=0,sa=10)

                # Objet
                to_=d.add_table(rows=1,cols=1); to_.style="Table Grid"
                co_=to_.rows[0].cells[0]; _shd(co_,_BC); _brd(co_)
                po_=co_.paragraphs[0]
                po_.paragraph_format.space_before=Pt(4); po_.paragraph_format.space_after=Pt(4)
                r1_=po_.add_run("Objet : "); r1_.bold=True; r1_.font.size=Pt(10)
                r1_.font.name="Arial"; r1_.font.color.rgb=_MR
                r2_=po_.add_run(f"Rapport de pilotage analytique — Exercice {exercice}")
                r2_.font.size=Pt(10); r2_.font.name="Arial"

                _P(d,"",sb=10,sa=0)
                _P(d,f"{nom_dirigeant},",bold=True,size=11,sb=0,sa=6)
                _P(d,f"Conformément aux diligences prévues dans notre mission d'accompagnement au "
                   f"pilotage analytique, nous avons l'honneur de vous adresser le présent rapport "
                   f"relatif à l'exercice {exercice}, pour la période du {periode_debut} au {periode_fin}.",
                   size=10,align="justify",sb=0,sa=6,line=14)
                _P(d,f"Ce rapport porte sur l'ensemble des {nb_eans_} ISBN actifs de votre catalogue, "
                   f"traités par notre outil Python VISION EDITION.",
                   size=10,align="justify",sb=0,sa=6,line=14)
                _P(d,f"Nous attirons votre attention sur le point central suivant : si votre maison "
                   f"d'édition dégage une marge brute globale de {_fp(taux_mb_)}, la prise en compte "
                   f"des charges indirectes de structure conduit à un résultat analytique net de "
                   f"{_fn(mn_,True)}. Cette situation appelle une analyse approfondie que nous vous "
                   f"soumettons ci-après.",
                   size=10,align="justify",sb=0,sa=10,line=14)

                # ── Section I — Synthèse ──
                _P(d,"I.   SYNTHÈSE DE LA PERFORMANCE ANALYTIQUE",bold=True,size=13,color=_MR,sb=8,sa=4)
                _sep(d)
                _P(d,"L'analyse du grand livre analytique de l'exercice fait ressortir les résultats suivants :",
                   size=10,sb=0,sa=6)

                ts_=d.add_table(rows=6,cols=3); ts_.style="Table Grid"
                _hr(ts_,[("Indicateur",8),("Montant",3),("Commentaire",5.5)])
                for i,(lbl,val,cmt,fl) in enumerate([
                    ("CA net éditeur",_fn(ca_n_),f"{nb_eans_} ISBN actifs",_BC),
                    ("Charges directes",_fn(ca_n_-mb_),"Fabrication, droits, personnel","FFFFFF"),
                    ("Marge brute sur ventes",_fn(mb_,True),f"Taux : {_fp(taux_mb_)}",_GC),
                    ("Charges indirectes de structure",_fn(ci_total_rp),
                     f"Quote-part par titre : {ci_par_ean_rp:.0f} EUR",_YC),
                    ("Résultat analytique net",_fn(mn_,True),
                     "Après affectation charges de structure",_RC if mn_<0 else _GC),
                ]):
                    rw_=ts_.rows[i+1]
                    _wc(rw_.cells[0],lbl,bold=True,size=9,fill=fl)
                    _wc(rw_.cells[1],val,bold=True,center=True,size=10,
                        color=RGBColor(0xC0,0,0) if (i==4 and mn_<0) else _MR,fill=fl)
                    _wc(rw_.cells[2],cmt,italic=True,size=8,
                        color=RGBColor(0x55,0x55,0x55),fill=fl)

                _P(d,f"L'exercice {exercice} présente une marge brute de {_fn(mb_,True)} ({_fp(taux_mb_)}). "
                   f"Toutefois, les charges indirectes ({_fn(ci_total_rp)}), réparties sur {nb_eans_} ISBN "
                   f"actifs à raison de {ci_par_ean_rp:.0f} EUR/titre, conduisent à un résultat analytique "
                   f"net de {_fn(mn_,True)}.",
                   size=10,align="justify",sb=6,sa=10,line=14)

                # ── Section II — Analyse par titre ──
                _P(d,"II.   ANALYSE DE LA RENTABILITÉ PAR TITRE",bold=True,size=13,color=_MR,sb=8,sa=4)
                _sep(d)
                _P(d,f"Sur les {nb_eans_} ISBN actifs au catalogue, {n_pos_} dégagent une marge brute "
                   f"positive et {n_neg_} affichent une marge brute négative "
                   f"({n_neg_/max(nb_eans_,1)*100:.1f}% du catalogue).",
                   size=10,align="justify",sb=0,sa=8,line=14)

                _P(d,f"A. Les {nb_top} titres les plus contributifs",bold=True,size=10,color=_MR,sb=4,sa=4)
                tt_=d.add_table(rows=nb_top+1,cols=5); tt_.style="Table Grid"
                _hr(tt_,[("Rg",0.8),("Titre",5),("CA net (EUR)",2),("Marge brute (EUR)",2.5),("Statut",1.2)])
                for i,rd_ in enumerate(top_list_):
                    bg_= _LC if i%2==0 else "FFFFFF"; rw_=tt_.rows[i+1]
                    _wc(rw_.cells[0],str(i+1),bold=True,center=True,size=9,color=_TR,fill=bg_)
                    _wc(rw_.cells[1],str(rd_["Titre"])[:44],bold=True,size=9,color=_MR,fill=bg_)
                    _wc(rw_.cells[2],_fn(rd_["CA net"]),center=True,size=9,fill=bg_)
                    _wc(rw_.cells[3],_fn(rd_["Marge brute"],True),bold=True,center=True,size=10,
                        color=RGBColor(0x21,0x73,0x46),fill=_GC)
                    _wc(rw_.cells[4],"✓ OK",bold=True,center=True,size=9,
                        color=RGBColor(0x21,0x73,0x46),fill=_GC)

                if top_list_:
                    _P(d,f"Nous attirons votre attention sur « {top_list_[0]['Titre']} » qui, avec une "
                       f"marge brute de {_fn(top_list_[0]['Marge brute'],True)}, constitue le titre "
                       f"le plus contributif de votre catalogue.",
                       size=10,align="justify",sb=6,sa=10,line=14)

                _P(d,f"B. Les {nb_top} titres les plus déficitaires — Action corrective requise",
                   bold=True,size=10,color=RGBColor(0xC0,0,0),sb=4,sa=4)
                tf_=d.add_table(rows=nb_top+1,cols=5); tf_.style="Table Grid"
                _hr(tf_,[("Rg",0.8),("Titre",5),("CA net (EUR)",2),("Marge brute (EUR)",2.5),("Statut",1.2)],
                    fill="C00000")
                for i,rd_ in enumerate(flop_list_):
                    bg_="FFF5F5" if i%2==0 else "FFFFFF"; rw_=tf_.rows[i+1]
                    _wc(rw_.cells[0],str(i+1),bold=True,center=True,size=9,
                        color=RGBColor(0xC0,0,0),fill=bg_)
                    _wc(rw_.cells[1],str(rd_["Titre"])[:44],bold=True,size=9,
                        color=RGBColor(0xC0,0,0),fill=bg_)
                    _wc(rw_.cells[2],_fn(rd_["CA net"]),center=True,size=9,fill=bg_)
                    _wc(rw_.cells[3],_fn(rd_["Marge brute"],True),bold=True,center=True,size=10,
                        color=RGBColor(0xC0,0,0),fill=_RC)
                    _wc(rw_.cells[4],"⚠ Alerte",bold=True,center=True,size=9,
                        color=RGBColor(0xC0,0,0),fill=_RC)

                if flop_list_:
                    _P(d,f"Nous appelons votre attention sur « {flop_list_[0]['Titre']} » qui accuse "
                       f"une perte de {abs(flop_list_[0]['Marge brute']):,.0f} EUR. L'ensemble de ces "
                       f"titres déficitaires appelle les mesures correctives détaillées en section IV.",
                       size=10,align="justify",sb=6,sa=10,line=14)

                # ── Section III — Charges indirectes ──
                _P(d,"III.   CHARGES INDIRECTES DE STRUCTURE",bold=True,size=13,color=_MR,sb=8,sa=4)
                _sep(d)
                _P(d,f"Les charges indirectes de structure s'élèvent à {_fn(ci_total_rp)}, soit "
                   f"{ci_total_rp/max(ca_n_,1)*100:.1f}% du chiffre d'affaires net. "
                   f"Avec {nb_eans_} ISBN actifs, chaque titre doit dégager une marge brute minimale "
                   f"de {ci_par_ean_rp:.0f} EUR pour couvrir sa quote-part de structure.",
                   size=10,align="justify",sb=0,sa=8,line=14)

                tc_=d.add_table(rows=3,cols=3); tc_.style="Table Grid"
                _hr(tc_,[("Indicateur",6),("Valeur",2.5),("Commentaire",7)])
                for i,(lbl,val,cmt,fl) in enumerate([
                    ("Charges indirectes totales",_fn(ci_total_rp),
                     "Loyer, salaires, charges de structure, honoraires transversaux",_BC),
                    ("Quote-part par titre actif",f"{ci_par_ean_rp:.0f} EUR/titre",
                     f"Clé de répartition : {nb_eans_} titres actifs — prorata catalogue","FFFFFF"),
                ]):
                    rw_=tc_.rows[i+1]
                    _wc(rw_.cells[0],lbl,bold=True,size=9,fill=fl)
                    _wc(rw_.cells[1],val,bold=True,center=True,size=10,color=_MR,fill=fl)
                    _wc(rw_.cells[2],cmt,italic=True,size=8,color=RGBColor(0x55,0x55,0x55),fill=fl)

                _P(d,f"Cette quote-part de {ci_par_ean_rp:.0f} EUR par titre constitue le seuil de "
                   f"rentabilité analytique minimal de chaque ISBN. Tout titre dont la marge brute "
                   f"est inférieure à ce seuil contribue négativement au résultat global.",
                   size=10,align="justify",sb=6,sa=10,line=14)

                # ── Section IV — Recommandations ──
                _P(d,"IV.   RECOMMANDATIONS ET PLAN D'ACTION",bold=True,size=13,color=_MR,sb=8,sa=4)
                _sep(d)
                _P(d,"Au vu de l'analyse ci-avant, nous vous soumettons les recommandations suivantes :",
                   size=10,sb=0,sa=6)

                tr_=d.add_table(rows=5,cols=4); tr_.style="Table Grid"
                _hr(tr_,[("Priorité",1.5),("Recommandation",5),("Action à engager",4),("Délai",2)])
                for i,(prio,rec,action,delai,fl) in enumerate([
                    ("1 — URGENT",f"Analyse des {n_neg_} titres déficitaires",
                     "Décision maintien ou arrêt commercial","Immédiat",_RC),
                    ("2 - URGENT","Gel des réimpressions des titres déficitaires",
                     "Suspension commandes impression en attente d'analyse","Immédiat","FFFFFF"),
                    ("3 — IMPORTANT","Réexamen des charges indirectes de structure",
                     "Audit des postes de charges - optimisation possible","T+3 mois",_YC),
                    ("4 — IMPORTANT",f"Réimpression prioritaire des {nb_top} titres porteurs",
                     "Lancement procédure réimpression","T+2 mois","FFFFFF"),
                ]):
                    rw_=tr_.rows[i+1]; urgent="URGENT" in prio
                    _wc(rw_.cells[0],prio,bold=True,center=True,size=8,
                        color=RGBColor(0xC0,0,0) if urgent else RGBColor(0xA6,0x7C,0),
                        fill=_RC if urgent else _YC)
                    _wc(rw_.cells[1],rec,bold=True,size=9,fill=fl)
                    _wc(rw_.cells[2],action,size=9,fill=fl)
                    _wc(rw_.cells[3],delai,center=True,bold=True,size=9,
                        color=RGBColor(0xC0,0,0) if urgent else RGBColor(0xA6,0x7C,0),fill=fl)

                _P(d,"",sb=10,sa=0)

                # ── Section V — Conclusion et prochaines étapes ──
                _P(d,"V.   CONCLUSION ET PROCHAINES ÉTAPES",bold=True,size=13,color=_MR,sb=8,sa=4)
                _sep(d)
                _P(d,"Le présent rapport met en évidence la nécessité d'engager rapidement une "
                   "réflexion sur la structure de votre catalogue et sur le niveau de vos charges "
                   "indirectes. La rentabilité analytique globale dépend directement de la capacité "
                   "de chaque titre à couvrir sa quote-part de charges de structure.",
                   size=10,align="justify",sb=0,sa=6,line=14)

                tv_=d.add_table(rows=4,cols=2); tv_.style="Table Grid"
                _hr(tv_,[("Prochaine étape",8.5),("Échéance",3)])
                for i,(etape,echeance,fl) in enumerate([
                    ("Réunion de restitution et analyse des titres déficitaires","Prochain trimestre",_BC),
                    ("Mise à jour des données analytiques (mois suivant)","Mensuel",_LC),
                    ("Révision du plan d'action suite décisions prises","Après réunion","FFFFFF"),
                ]):
                    rw_=tv_.rows[i+1]
                    _wc(rw_.cells[0],etape,size=9,fill=fl)
                    _wc(rw_.cells[1],echeance,bold=True,center=True,size=9,color=_MR,fill=fl)

                _P(d,"",sb=10,sa=0)
                _P(d,"Nous demeurons à votre entière disposition pour discuter de ces éléments lors "
                   "de notre prochaine réunion trimestrielle de restitution et vous prions d'agréer, "
                   f"{nom_dirigeant}, l'expression de nos salutations distinguées.",
                   size=10,align="justify",sb=0,sa=16,line=14)

                # Signatures
                ts2_=d.add_table(rows=1,cols=2); ts2_.style="Table Grid"
                for cs_,ti_,fi_,ci__ in [
                    (ts2_.rows[0].cells[0],f"L'Expert-Comptable\n{NOM_CABINET}",_BC,_MR),
                    (ts2_.rows[0].cells[1],f"Le Dirigeant\n{nom_editeur}",_LC,_TR)
                ]:
                    _shd(cs_,fi_); _brd(cs_)
                    ps_=cs_.paragraphs[0]
                    ps_.paragraph_format.space_before=Pt(6); ps_.paragraph_format.space_after=Pt(6)
                    for j,li_ in enumerate([ti_,"\n\nSignature : ______________________",
                                             "\nDate : ___________________________"]):
                        rs_=ps_.add_run(li_); rs_.font.size=Pt(9); rs_.font.name="Arial"
                        rs_.bold=(j==0); rs_.font.color.rgb=ci__

                # Note méthodologique
                _P(d,"",sb=8,sa=0)
                pn_=d.add_paragraph()
                pn_.paragraph_format.space_before=Pt(4); pn_.paragraph_format.space_after=Pt(4)
                rn1_=pn_.add_run("Note méthodologique : ")
                rn1_.bold=True; rn1_.font.size=Pt(8); rn1_.font.name="Arial"
                rn1_.font.color.rgb=RGBColor(0x55,0x55,0x55)
                sfx_ = " (titres anonymisés)" if mode_anon else ""
                rn2_=pn_.add_run(
                    f"Rapport généré automatiquement par Python VISION EDITION — "
                    f"Période : {periode_debut} au {periode_fin} — "
                    f"{nb_eans_} ISBN actifs{sfx_} — "
                    f"Charges indirectes réparties au prorata du nombre de titres actifs — "
                    f"Données provisoires à la date du {date_fr_}."
                )
                rn2_.italic=True; rn2_.font.size=Pt(8); rn2_.font.name="Arial"
                rn2_.font.color.rgb=RGBColor(0x77,0x77,0x77)

                buf_w=BytesIO(); d.save(buf_w); buf_w.seek(0)
                fn_w = (f"Rapport_Pilotage_{nom_editeur[:20].replace(' ','_')}_"
                        f"{exercice.replace('/','_')}.docx")
                st.success("✅ Rapport de pilotage généré avec succès !")
                st.download_button(
                    "⬇️ Télécharger le rapport Word",
                    data=buf_w, file_name=fn_w,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True, type="primary"
                )

            except Exception as e:
                st.error(f"❌ Erreur : {e}")
                st.exception(e)


# ── TRÉSORERIE (EC) ──
elif role == "ec" and page == "💰 Trésorerie prévisionnelle":
    df, params = check_pivot()
    dos = get_dossier(st.session_state["dossier_id"])
    st.header(f"💰 Trésorerie prévisionnelle — {dos['nom']}")
    df["Date"] = pd.to_datetime(df["Date"], errors="coerce")
    df["Débit"] = pd.to_numeric(df["Débit"], errors="coerce").fillna(0)
    df["Crédit"] = pd.to_numeric(df["Crédit"], errors="coerce").fillna(0)
    if "Journal" not in df.columns: df["Journal"] = ""
    df = df.dropna(subset=["Date"])
    if df.empty:
        st.warning("Aucune écriture datée.")
        st.stop()

    col1, col2 = st.columns(2)
    with col1:
        date_debut = st.date_input("Date de départ", df["Date"].min())
        tresorerie_ouv = st.number_input("Trésorerie à l'ouverture (€)", value=0.0, step=100.0)
    with col2:
        horizon = st.slider("Horizon projection (mois)", 0, 24, 6)
        comptes_banque = tuple(x.strip() for x in st.text_input("Comptes trésorerie","512,530,580").split(",") if x.strip())
        journaux_exclus = tuple(x.strip() for x in st.text_input("Journaux exclus (AN)","AN").split(",") if x.strip())

    mask_an = df["Journal"].isin(journaux_exclus) if journaux_exclus else pd.Series(False, index=df.index)
    df_flux = df[~mask_an & (df["Date"]>=pd.to_datetime(date_debut))].copy()
    df_flux["Mois"] = df_flux["Date"].dt.to_period("M")

    if df_flux.empty:
        st.warning("Aucune écriture après la date de départ.")
        st.stop()

    mois = sorted(df_flux["Mois"].unique())
    encaissements = df_flux[df_flux["Compte"].astype(str).str.startswith("41")].groupby("Mois")["Crédit"].sum()
    decaissements = df_flux[df_flux["Compte"].astype(str).str.startswith(("40","42","43","44"))].groupby("Mois")["Débit"].sum()
    flux_net = encaissements.sub(decaissements, fill_value=0).reindex(mois, fill_value=0)
    treso = tresorerie_ouv + flux_net.cumsum()

    st.session_state["treso_real"] = treso
    st.session_state["treso_ouverture"] = tresorerie_ouv

    m1,m2,m3 = st.columns(3)
    m1.metric("Trésorerie ouverture", f"{fmt_fr(tresorerie_ouv)} €")
    m2.metric("Trésorerie clôture (réalisé)", f"{fmt_fr(treso.iloc[-1])} €")
    m3.metric("Flux net généré", f"{fmt_fr(flux_net.sum())} €")

    _MOIS_FR = {1:"Jan",2:"Fév",3:"Mar",4:"Avr",5:"Mai",6:"Jui",7:"Jul",8:"Aoû",9:"Sep",10:"Oct",11:"Nov",12:"Déc"}
    def ml(p): return f"{_MOIS_FR.get(p.month,str(p.month))} {p.year}"

    fig_t = go.Figure()
    fig_t.add_trace(go.Scatter(x=[ml(m) for m in treso.index], y=treso.values,
                               name="Réalisé", line=dict(color="#111827",width=2.5)))

    if horizon > 0:
        base_enc = encaissements.iloc[-3:].mean() if len(encaissements)>=3 else (encaissements.mean() if len(encaissements) else 0)
        base_dec = decaissements.iloc[-3:].mean() if len(decaissements)>=3 else (decaissements.mean() if len(decaissements) else 0)
        col_sc1,col_sc2,col_sc3,col_sc4 = st.columns(4)
        t_opt  = col_sc1.number_input("Croissance encaissements optimiste (%/mois)", value=4.0, step=0.5)/100
        t_cent = col_sc2.number_input("Croissance encaissements central (%/mois)", value=2.0, step=0.5)/100
        t_pess = col_sc3.number_input("Croissance encaissements pessimiste (%/mois)", value=0.0, step=0.5)/100
        t_ch   = col_sc4.number_input("Évolution charges (%/mois)", value=1.0, step=0.5)/100
        for nom, tc, col_hex in [("Optimiste",t_opt,"#10B981"),("Central",t_cent,"#3B82F6"),("Pessimiste",t_pess,"#EF4444")]:
            futurs = [mois[-1]+i for i in range(1,horizon+1)]
            enc_f = base_enc; dec_f = base_dec; proj = []
            for m_f in futurs:
                enc_f *= (1+tc); dec_f *= (1+t_ch); proj.append(enc_f-dec_f)
            treso_proj = treso.iloc[-1] + pd.Series(proj, index=futurs).cumsum()
            treso_full = pd.concat([treso.iloc[[-1]], treso_proj])
            fig_t.add_trace(go.Scatter(x=[ml(m) for m in treso_full.index], y=treso_full.values,
                                       name=nom, line=dict(color=col_hex,width=2,dash="dot")))

    fig_t.add_hline(y=0, line_dash="dash", line_color="gray")
    fig_t.update_layout(height=380, margin=dict(t=20), xaxis_title="", yaxis_title="€",
                        legend=dict(orientation="h"))
    st.plotly_chart(fig_t, use_container_width=True)

# ── DROITS D'AUTEURS (EC) ──
elif role == "ec" and page == "✍️ Droits d'auteurs":
    df, params = check_pivot()
    dos = get_dossier(st.session_state["dossier_id"])
    st.header(f"✍️ Droits d'auteurs — {dos['nom']}")
    st.info("Module complet disponible — configurez les comptes ci-dessous.")
    st.info(
        "**Comment fonctionne ce module ?** "
        "Les droits d'auteurs déjà comptabilisés sont lus directement dans votre grand livre. "
        "Un simulateur est disponible pour estimer avant provision."
    )

    ong1, ong2, ong3, ong4 = st.tabs([
        "📋 Référentiel contrats",
        "🧮 Simulateur",
        "📒 Réel (comptabilisé)",
        "📄 Relevés par auteur"
    ])

    # ── Initialisation référentiel ──
    if "royalties_referentiel" not in st.session_state:
        st.session_state["royalties_referentiel"] = []

    with ong1:
        st.subheader("Saisie des contrats auteurs")
        col_a, col_b = st.columns(2)
        with col_a:
            inp_auteur = st.text_input("Nom de l'auteur", key="da_auteur")
            inp_isbn   = st.text_input("ISBN (code analytique exact du socle)", key="da_isbn")
            inp_titre  = st.text_input("Titre du livre", key="da_titre")
        with col_b:
            inp_part   = st.number_input("Part de cet auteur (%)", min_value=0.0, max_value=100.0, value=100.0, step=1.0, key="da_part")
            inp_statut = st.selectbox("Statut fiscal", ["Traitements et salaires (option assimilé)","BNC (droits d'auteur)"], key="da_statut")
            inp_taux   = st.number_input("Taux forfaitaire (%)", value=10.0, step=0.5, key="da_taux")
        if st.button("➕ Ajouter ce contrat", key="da_add"):
            if inp_auteur and inp_isbn:
                st.session_state["royalties_referentiel"].append({
                    "auteur": inp_auteur, "isbn": inp_isbn.strip(),
                    "titre": inp_titre or inp_isbn, "part": inp_part,
                    "statut_fiscal": inp_statut,
                    "paliers": [{"seuil": 0, "taux": inp_taux}]
                })
                st.success(f"✅ Contrat ajouté : {inp_auteur} / {inp_titre or inp_isbn}")
        st.divider()
        ref = st.session_state["royalties_referentiel"]
        if ref:
            rows_ref = [{"Auteur":c["auteur"],"ISBN":c["isbn"],"Titre":c["titre"],
                         "Part (%)":c["part"],"Taux (%)":c["paliers"][0]["taux"],
                         "Statut":c.get("statut_fiscal","")} for c in ref]
            st.dataframe(pd.DataFrame(rows_ref), use_container_width=True, hide_index=True)
            idx_s = st.number_input("Supprimer le contrat n°", min_value=0, max_value=max(0,len(ref)-1), step=1, key="da_del_idx")
            if st.button("🗑️ Supprimer", key="da_del"):
                st.session_state["royalties_referentiel"].pop(int(idx_s)); st.rerun()
        else:
            st.info("Aucun contrat enregistré.")

    with ong2:
        st.subheader("Simulateur — droits d'auteur & précompte URSSAF")
        if "df_pivot" not in st.session_state or not st.session_state["royalties_referentiel"]:
            st.warning("⚠️ Générez d'abord le socle et saisissez au moins un contrat dans l'onglet Référentiel.")
        else:
            df_sim2 = df.copy()
            params_sim2 = params
            resultats_sim = []
            for contrat in st.session_state["royalties_referentiel"]:
                isbn = contrat["isbn"]
                mask_v2 = df_sim2["Compte"].astype(str).str.startswith(tuple(params_sim2["ventes"]))
                mask_r2 = df_sim2["Compte"].astype(str).str.startswith(tuple(params_sim2["retours"]))
                mask_rem2 = df_sim2["Compte"].astype(str).str.startswith(tuple(params_sim2.get("remises",[])))
                df_i2 = df_sim2[df_sim2["Code_Analytique"]==isbn]
                ca_b2  = df_i2[mask_v2.reindex(df_i2.index,fill_value=False)]["Crédit"].sum()
                ret2   = df_i2[mask_r2.reindex(df_i2.index,fill_value=False)]["Débit"].sum()
                rem2   = df_i2[mask_rem2.reindex(df_i2.index,fill_value=False)]["Débit"].sum()
                ca_n2  = ca_b2 - ret2 - rem2
                base2  = max(ca_n2, 0)
                taux2  = contrat["paliers"][0]["taux"] / 100
                droits_bruts = base2 * taux2 * contrat["part"] / 100
                resultats_sim.append({
                    "Auteur": contrat["auteur"], "Titre": contrat["titre"],
                    "Statut": contrat.get("statut_fiscal",""),
                    "CA net (€)": round(ca_n2,2),
                    "Base calcul (€)": round(base2,2),
                    "Droits bruts (€)": round(droits_bruts,2),
                })
            df_sim_res = pd.DataFrame(resultats_sim)
            cols_s2 = ["CA net (€)","Base calcul (€)","Droits bruts (€)"]
            st.dataframe(df_sim_res.style.format({c:(lambda x: f"{fmt_fr(x,2)} €") for c in cols_s2}), use_container_width=True)
            total_bruts = df_sim_res["Droits bruts (€)"].sum()
            st.metric("💰 Total droits bruts estimés", f"{fmt_fr(total_bruts,2)} €")
            st.divider()
            st.markdown("**Précompte URSSAF estimé**")
            col_u1,col_u2,col_u3,col_u4 = st.columns(4)
            taux_csg  = col_u1.number_input("CSG+CRDS (%)", value=9.70, step=0.01, key="da_csg")
            taux_fp   = col_u2.number_input("Formation prof. (%)", value=1.00, step=0.01, key="da_fp")
            taux_raap = col_u3.number_input("RAAP (%)", value=0.0, step=0.01, key="da_raap")
            taux_diff = col_u4.number_input("Contribution diffuseur (%)", value=1.10, step=0.05, key="da_diff_taux")
            est_ts = df_sim_res["Statut"].str.startswith("Traitements")
            assiette = df_sim_res["Droits bruts (€)"] * 0.9825
            precompte = np.where(est_ts, assiette*(taux_csg+taux_fp+taux_raap)/100, 0.0)
            contrib   = df_sim_res["Droits bruts (€)"] * taux_diff / 100
            total_urssaf_s = precompte.sum() + contrib.sum()
            cu1,cu2,cu3 = st.columns(3)
            cu1.metric("Précompte URSSAF (TS)", f"{fmt_fr(precompte.sum(),2)} €")
            cu2.metric("Contribution diffuseur", f"{fmt_fr(contrib.sum(),2)} €")
            cu3.metric("Total à reverser URSSAF", f"{fmt_fr(total_urssaf_s,2)} €")

    with ong3:
        st.subheader("Montants réellement comptabilisés")
        df["Date"] = pd.to_datetime(df["Date"], errors="coerce")
        col1_d,col2_d = st.columns(2)
        with col1_d:
            date_debut_d = st.date_input("Période du", df["Date"].dropna().min().date() if not df["Date"].dropna().empty else _dt.date.today(), key="da_d")
            compte_db  = st.text_input("Droits bruts (charge)", value="604300000", key="da_cpt_db")
            compte_urs = st.text_input("URSSAF à payer", value="438106", key="da_cpt_urs")
        with col2_d:
            date_fin_d = st.date_input("Au", df["Date"].dropna().max().date() if not df["Date"].dropna().empty else _dt.date.today(), key="da_f")
            compte_diff_d = st.text_input("Contribution diffuseur", value="645106", key="da_cpt_diff")
            compte_net_d  = st.text_input("Droits à payer (net)", value="408106", key="da_cpt_net")
            compte_av_d   = st.text_input("À-valoirs (bilan)", value="409600", key="da_cpt_av")

        mask_p = (df["Date"]>=pd.to_datetime(date_debut_d)) & (df["Date"]<=pd.to_datetime(date_fin_d))
        df_p = df[mask_p]

        def par_isbn_d(cpt, sens):
            m = df_p["Compte"].astype(str).str.strip()==str(cpt).strip()
            if not m.any(): return pd.Series(dtype=float)
            g = df_p[m].groupby("Code_Analytique")
            return (g["Débit"].sum()-g["Crédit"].sum()) if sens=="debit" else (g["Crédit"].sum()-g["Débit"].sum())

        db_s   = par_isbn_d(compte_db,"debit")
        urs_s  = par_isbn_d(compte_urs,"credit")
        diff_s = par_isbn_d(compte_diff_d,"debit")
        net_s  = par_isbn_d(compte_net_d,"credit")
        av_s   = par_isbn_d(compte_av_d,"debit")

        isbns_d = sorted(set(db_s.index)|set(urs_s.index)|set(diff_s.index)|set(net_s.index))
        isbns_d = [i for i in isbns_d if str(i).strip() not in ("","CHARGES INDIRECTES","PRODUITS INDIRECTS")]

        if not isbns_d:
            st.info("Aucune écriture trouvée sur ces comptes pour la période.")
        else:
            lignes_d = []
            for isbn in isbns_d:
                lignes_d.append({
                    "ISBN": isbn,
                    "Droits bruts (€)":         round(float(db_s.get(isbn,0)),2),
                    "Contribution diffuseur (€)": round(float(diff_s.get(isbn,0)),2),
                    "Précompte URSSAF (€)":      round(float(urs_s.get(isbn,0)),2),
                    "Net dû auteur (€)":         round(float(net_s.get(isbn,0)),2),
                    "À-valoir restant (€)":      round(float(av_s.get(isbn,0)),2),
                })
            df_d = pd.DataFrame(lignes_d)
            cols_m = ["Droits bruts (€)","Contribution diffuseur (€)","Précompte URSSAF (€)","Net dû auteur (€)","À-valoir restant (€)"]
            st.dataframe(df_d.style.format({c:(lambda x: f"{fmt_fr(x,2)} €") for c in cols_m}), use_container_width=True)
            total_db  = df_d["Droits bruts (€)"].sum()
            total_urs = df_d["Précompte URSSAF (€)"].sum()
            total_diff_t = df_d["Contribution diffuseur (€)"].sum()
            total_net = df_d["Net dû auteur (€)"].sum()
            c1,c2,c3,c4 = st.columns(4)
            c1.metric("Droits bruts", f"{fmt_fr(total_db,2)} €")
            c2.metric("Précompte URSSAF", f"{fmt_fr(total_urs,2)} €")
            c3.metric("Contribution diffuseur", f"{fmt_fr(total_diff_t,2)} €")
            c4.metric("Net dû aux auteurs", f"{fmt_fr(total_net,2)} €")
            st.markdown(f"**🏛️ Total à reverser à l'URSSAF : {fmt_fr(total_urs+total_diff_t,2)} €**")
            st.session_state["df_royalties_reel"] = df_d
            buf_d = BytesIO()
            with pd.ExcelWriter(buf_d, engine="openpyxl") as writer:
                df_d.to_excel(writer, index=False, sheet_name="Droits_auteurs")
            buf_d.seek(0)
            st.download_button("📥 Exporter (Excel)", buf_d,
                file_name=f"Droits_auteurs_{date_debut_d}_{date_fin_d}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                key="da_export")

    with ong4:
        st.subheader("Relevés de droits par auteur")
        if "df_royalties_reel" not in st.session_state or st.session_state["df_royalties_reel"].empty:
            st.warning("⚠️ Effectuez d'abord la lecture dans l'onglet 📒 Réel (comptabilisé).")
        else:
            df_rel = st.session_state["df_royalties_reel"].copy()
            auteurs_list = ["Tous"]
            ref_map = {c["isbn"]: c["auteur"] for c in st.session_state.get("royalties_referentiel",[])}
            if ref_map:
                df_rel["Auteur"] = df_rel["ISBN"].map(ref_map).fillna("(auteur non identifié)")
                auteurs_list += sorted(df_rel["Auteur"].unique().tolist())
            auteur_sel = st.selectbox("Sélectionner un auteur", auteurs_list, key="da_auteur_sel")
            df_aff = df_rel if auteur_sel=="Tous" else df_rel[df_rel["Auteur"]==auteur_sel]
            cols_rel = [c for c in df_aff.columns if c in ["Auteur","ISBN","Droits bruts (€)","Contribution diffuseur (€)","Précompte URSSAF (€)","Net dû auteur (€)","À-valoir restant (€)"]]
            cols_m_rel = [c for c in cols_rel if "(€)" in c]
            st.dataframe(df_aff[cols_rel].style.format({c:(lambda x: f"{fmt_fr(x,2)} €") for c in cols_m_rel}), use_container_width=True, hide_index=True)
            buf_rel = BytesIO()
            with pd.ExcelWriter(buf_rel, engine="openpyxl") as writer:
                df_aff[cols_rel].to_excel(writer, index=False, sheet_name="Releve")
            buf_rel.seek(0)
            st.download_button(f"📥 Télécharger relevé — {auteur_sel}", buf_rel,
                file_name=f"Royalties_{auteur_sel.replace(' ','_')}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                key="da_export_releve")

# ── RETOURS & REMISES (EC) ──
elif role == "ec" and page == "📦 Retours & Remises":
    df, params = check_pivot()
    dos = get_dossier(st.session_state["dossier_id"])
    st.header(f"📦 Retours & Remises — {dos['nom']}")
    df["Date"] = pd.to_datetime(df["Date"], errors="coerce")
    df["Mois"] = df["Date"].dt.strftime("%Y-%m")
    seuil = st.sidebar.number_input("Seuil alerte (%)", value=25, step=5)

    df_v_  = df[df["Compte"].astype(str).str.startswith(tuple(params["ventes"]))]
    df_r_  = df[mask_retours(df,params)]
    df_rem_= df[mask_remises(df,params)]
    pref_d = tuple(params.get("ventes_distributeur") or params["ventes"])
    df_vd_ = df[df["Compte"].astype(str).str.startswith(pref_d)]

    tv = df_v_["Crédit"].sum(); vd = df_vd_["Crédit"].sum()
    tr = abs(df_r_["Débit"].sum()-df_r_["Crédit"].sum()) if not df_r_.empty else 0
    rem= abs(df_rem_["Débit"].sum()-df_rem_["Crédit"].sum()) if not df_rem_.empty else 0
    taux_r = (tr/vd*100) if vd else 0; taux_rem_pct = (rem/vd*100) if vd else 0

    c1,c2,c3,c4 = st.columns(4)
    c1.metric("CA brut",f"{fmt_fr(tv)} €")
    c2.metric("Total retours",f"{fmt_fr(tr)} €")
    c3.metric("Taux de retour",f"{taux_r:.1f} %",
              delta="⚠️ Dépasse le seuil" if taux_r>seuil else "✅ Normal",
              delta_color="inverse" if taux_r>seuil else "normal")
    c4.metric("Taux de remise",f"{taux_rem_pct:.1f} %")

    if taux_r > seuil:
        st.error(f"🚨 Taux de retour ({taux_r:.1f}%) dépasse votre seuil de {seuil}% !")

    if not df_r_.empty:
        trend_r = df_r_.groupby("Mois")["Débit"].sum().abs().reset_index()
        fig_r = px.bar(trend_r, x="Mois", y="Débit", title="Retours mensuels (€)",
                       color_discrete_sequence=["#EF4444"], height=300)
        st.plotly_chart(fig_r, use_container_width=True)

        df_ri = filtrer_isbn_reels(df_r_)
        if not df_ri.empty:
            ret_isbn = df_ri.groupby("Code_Analytique")["Débit"].sum().abs().reset_index().sort_values("Débit",ascending=False)
            st.subheader("Retours par titre")
            st.dataframe(ret_isbn, hide_index=True)

# ── SYNTHÈSE FINANCIÈRE (EC) ──
elif role == "ec" and page == "📊 Synthèse financière":
    df, params = check_pivot()
    dos = get_dossier(st.session_state["dossier_id"])
    st.header(f"📊 Synthèse financière — {dos['nom']}")

    df_v_   = df[mask_ventes(df,params)]
    df_r_   = df[mask_retours(df,params)]
    df_rem_ = df[mask_remises(df,params)]
    df_c_   = df[mask_charges(df,params)]
    df_prov = df[mask_provisions_reprises(df,params)]
    net_p   = df_prov["Crédit"].sum()-df_prov["Débit"].sum()

    ca_b  = df_v_["Crédit"].sum(); corr = df_v_["Débit"].sum()
    tr_   = df_r_["Débit"].sum()-df_r_["Crédit"].sum()
    rem_  = df_rem_["Débit"].sum()-df_rem_["Crédit"].sum()
    ca_n  = ca_b-tr_-rem_-corr
    ch_t  = (df_c_["Débit"].sum()-df_c_["Crédit"].sum())-net_p
    res_  = ca_n-ch_t
    mb_p  = (res_/ca_b*100) if ca_b else 0

    soldes   = [ca_b,-tr_,-rem_,-corr,ca_n,-ch_t,res_]
    libelles = ["CA brut","Retours","Remises","Corrections","CA net","Charges","Résultat net"]

    col1,col2 = st.columns([1.2,1])
    with col1:
        st.subheader("Compte de résultat synthétique")
        df_sum = pd.DataFrame({"Poste":libelles,"Montant (€)":soldes})
        st.dataframe(df_sum.style.format({"Montant (€)":(lambda x: fmt_fr(x))}), hide_index=True)
        st.metric("Taux de marge nette",f"{mb_p:.1f} %")
    with col2:
        fig_w = go.Figure(go.Waterfall(
            orientation="v",
            measure=["absolute","relative","relative","relative","total","relative","total"],
            x=libelles, y=soldes,
            decreasing={"marker":{"color":"#EF4444"}},
            increasing={"marker":{"color":"#10B981"}},
            totals={"marker":{"color":"#3B82F6"}}
        ))
        fig_w.update_layout(height=360, margin=dict(t=20))
        st.plotly_chart(fig_w, use_container_width=True)

    buf_s = BytesIO()
    with pd.ExcelWriter(buf_s, engine="openpyxl") as writer:
        df_sum.to_excel(writer, index=False, sheet_name="Synthese")
    buf_s.seek(0)
    st.download_button("📥 Exporter la synthèse (Excel)", buf_s,
        file_name="Synthese_Financiere.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")


# ── ASSISTANT IA (EC) ──
elif role == "ec" and page == "🤖 Assistant IA":
    df_ia, params_ia = check_pivot()
    dos = get_dossier(st.session_state["dossier_id"])
    st.header(f"🤖 Assistant IA — {dos['nom']}")
    client_ai = get_client_ai()
    if client_ai is None:
        st.error("⚠️ Clé API Anthropic non configurée (ANTHROPIC_API_KEY dans les secrets Streamlit).")
        st.stop()

    # Contexte données
    df_v_ = df_ia[mask_ventes(df_ia,params_ia)]
    df_r_ = df_ia[mask_retours(df_ia,params_ia)]
    ca_b_ = df_v_["Crédit"].sum(); tr_ = df_r_["Débit"].sum()-df_r_["Crédit"].sum()
    ca_n_ = ca_b_ - tr_
    taux_r_ = (tr_/ca_b_*100) if ca_b_ else 0
    nb_isbn_ = filtrer_isbn_reels(df_ia)["Code_Analytique"].nunique()

    data_ctx = f"""
DONNÉES ANALYTIQUES — {dos['nom']} — {dos.get('exercice','')}
CA brut : {fmt_fr(ca_b_)} EUR
Total retours : {fmt_fr(tr_)} EUR
CA net : {fmt_fr(ca_n_)} EUR
Taux de retour : {taux_r_:.1f} %
Nombre de titres actifs : {nb_isbn_}
"""
    sys_prompt = SYSTEM_PROMPT_EC + f"\n\nDONNÉES DU DOSSIER :\n{data_ctx}"

    for msg in st.session_state["messages_agent"]:
        with st.chat_message(msg["role"], avatar="🧑" if msg["role"]=="user" else "🤖"):
            st.markdown(msg["content"])

    if prompt_u := st.chat_input("Posez votre question..."):
        st.session_state["messages_agent"].append({"role":"user","content":prompt_u})
        with st.chat_message("user", avatar="🧑"):
            st.markdown(prompt_u)
        with st.chat_message("assistant", avatar="🤖"):
            with st.spinner("Analyse en cours..."):
                resp = client_ai.messages.create(
                    model="claude-sonnet-4-6", max_tokens=800,
                    system=sys_prompt, messages=st.session_state["messages_agent"]
                )
                answer = resp.content[0].text
                st.markdown(answer)
                st.session_state["messages_agent"].append({"role":"assistant","content":answer})

    if st.button("🗑️ Effacer"):
        st.session_state["messages_agent"] = []
        st.rerun()

# ── RAPPORT DE PILOTAGE (EC) ──
elif role == "ec" and page == "📄 Rapport de pilotage":
    df, params = check_pivot()
    dos = get_dossier(st.session_state["dossier_id"])
    st.header("📄 Rapport de pilotage analytique")
    st.markdown(f"*{dos['nom']} — {dos.get('exercice','')}*")

    if not DOCX_AVAILABLE:
        st.error("❌ python-docx non installé. Ajoutez `python-docx` dans requirements.txt.")
        st.stop()

    col1,col2 = st.columns(2)
    with col1:
        nom_editeur = st.text_input("Maison d'édition", value=dos.get("nom",""))
        nom_dirigeant = st.text_input("Nom du dirigeant", value=dos.get("dirigeant",""))
        exercice = st.text_input("Exercice", value=dos.get("exercice","2025/2026"))
    with col2:
        date_rapport = st.date_input("Date du rapport", value=_dt.date.today())
        periode_debut = st.text_input("Début de période", value=dos.get("periode_debut","01/04/2025"))
        periode_fin   = st.text_input("Fin de période",   value=dos.get("periode_fin","31/03/2026"))

    mode_anon = st.checkbox("🕶️ Anonymiser les titres", value=False)
    nb_top = st.slider("Nombre de titres Top / Flop", 3, 10, 5)

    titres_r = sorted(filtrer_isbn_reels(df)["Code_Analytique"].astype(str).unique().tolist())
    if not titres_r:
        st.error("Aucun ISBN détecté.")
        st.stop()

    res_r = calculer_indicateurs_titres(df, params, titres_r)
    mapping_anon = obtenir_mapping_anonymisation(df)
    res_r["Titre"] = res_r["Code_Analytique"].apply(lambda c: mapping_anon.get(c,c) if mode_anon else c)

    df_v_ = df[mask_ventes(df,params)]; df_c_ = df[mask_charges(df,params)]
    df_r_ = df[mask_retours(df,params)]; df_rem_ = df[mask_remises(df,params)]
    df_prov = df[mask_provisions_reprises(df,params)]
    net_p = df_prov["Crédit"].sum()-df_prov["Débit"].sum()
    ca_b_ = df_v_["Crédit"].sum()-df_v_["Débit"].sum()
    tr_   = df_r_["Débit"].sum()-df_r_["Crédit"].sum()
    rem_  = df_rem_["Débit"].sum()-df_rem_["Crédit"].sum()
    ca_n_ = ca_b_-tr_-rem_
    ch_t_ = (df_c_["Débit"].sum()-df_c_["Crédit"].sum())-net_p
    mn_   = ca_n_-ch_t_
    mb_   = res_r["Marge brute"].sum()
    taux_mb_ = (mb_/ca_b_*100) if ca_b_ else 0
    n_pos_ = (res_r["Marge brute"]>0).sum()
    n_neg_ = (res_r["Marge brute"]<=0).sum()
    nb_eans_ = len(titres_r)
    res_sorted_ = res_r.sort_values("Marge brute",ascending=False).reset_index(drop=True)

    st.divider()
    m1,m2,m3,m4 = st.columns(4)
    m1.metric("CA net", f"{fmt_fr(ca_n_)} €")
    m2.metric("Marge brute", f"{fmt_fr(mb_)} €", f"{taux_mb_:.1f} %")
    m3.metric("Résultat net", f"{fmt_fr(mn_)} €")
    m4.metric("Titres déficitaires", f"{n_neg_} / {nb_eans_}")

    if st.button("🖨️ Générer le rapport Word", type="primary", use_container_width=True):
        with st.spinner("Génération..."):
            try:
                _MR = RGBColor(0x1F,0x4E,0x79); _TR = RGBColor(0xC0,0x52,0x2A)
                _WH = RGBColor(0xFF,0xFF,0xFF)

                def _shd(cell, h):
                    tc=cell._tc; p=tc.get_or_add_tcPr()
                    for o in p.findall(qn("w:shd")): p.remove(o)
                    s=OxmlElement("w:shd"); s.set(qn("w:val"),"clear")
                    s.set(qn("w:color"),"auto"); s.set(qn("w:fill"),h); p.append(s)

                def _brd(cell,c="CCCCCC",z="4"):
                    tc=cell._tc; p=tc.get_or_add_tcPr()
                    for o in p.findall(qn("w:tcBorders")): p.remove(o)
                    b=OxmlElement("w:tcBorders")
                    for e in ["top","left","bottom","right"]:
                        x=OxmlElement(f"w:{e}"); x.set(qn("w:val"),"single")
                        x.set(qn("w:sz"),z); x.set(qn("w:color"),c); b.append(x)
                    p.append(b)

                def _wc(cell,text,bold=False,italic=False,color=None,size=9,center=False,fill=None):
                    if fill: _shd(cell,fill)
                    _brd(cell); cell.vertical_alignment=WD_ALIGN_VERTICAL.CENTER
                    p=cell.paragraphs[0]; p.clear()
                    p.alignment=WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
                    p.paragraph_format.space_before=Pt(2); p.paragraph_format.space_after=Pt(2)
                    r=p.add_run(str(text)); r.bold=bold; r.italic=italic
                    r.font.size=Pt(size); r.font.name="Arial"
                    r.font.color.rgb=color if color else RGBColor(0,0,0)

                def _P(d,text,bold=False,italic=False,size=10,color=None,align="left",sb=6,sa=6):
                    p=d.add_paragraph()
                    p.paragraph_format.space_before=Pt(sb); p.paragraph_format.space_after=Pt(sa)
                    p.alignment={"left":WD_ALIGN_PARAGRAPH.LEFT,"center":WD_ALIGN_PARAGRAPH.CENTER,
                                 "justify":WD_ALIGN_PARAGRAPH.JUSTIFY}.get(align,WD_ALIGN_PARAGRAPH.LEFT)
                    r=p.add_run(text); r.bold=bold; r.italic=italic
                    r.font.size=Pt(size); r.font.name="Arial"
                    if color: r.font.color.rgb=color
                    return p

                def _sep(d):
                    p=d.add_paragraph()
                    p.paragraph_format.space_before=Pt(2); p.paragraph_format.space_after=Pt(6)
                    pr=p._p.get_or_add_pPr(); pb=OxmlElement("w:pBdr"); bt=OxmlElement("w:bottom")
                    bt.set(qn("w:val"),"single"); bt.set(qn("w:sz"),"6"); bt.set(qn("w:color"),"C0522A")
                    pb.append(bt); pr.append(pb)

                def _fn(n,sign=False):
                    return ("+" if (sign and n>=0) else "")+f"{n:,.0f} EUR".replace(",","  ")

                top_list_ = [res_sorted_.iloc[i] for i in range(min(nb_top,len(res_sorted_)))]
                flop_list_= [res_sorted_.iloc[-(i+1)] for i in range(min(nb_top,len(res_sorted_)))]

                d=Document()
                for s in d.sections:
                    s.top_margin=Cm(2); s.bottom_margin=Cm(2)
                    s.left_margin=Cm(2.5); s.right_margin=Cm(2)

                _P(d,NOM_CABINET,bold=True,size=14,color=_MR,sb=0,sa=2)
                _P(d,"Expert-comptable — Maisons d'édition indépendantes",size=10,sb=0,sa=4)

                _mois_fr_={1:"janvier",2:"février",3:"mars",4:"avril",5:"mai",6:"juin",
                            7:"juillet",8:"août",9:"septembre",10:"octobre",11:"novembre",12:"décembre"}
                date_fr_=f"{date_rapport.day} {_mois_fr_[date_rapport.month]} {date_rapport.year}"
                pr_=d.add_paragraph(); pr_.alignment=WD_ALIGN_PARAGRAPH.RIGHT
                pr_.paragraph_format.space_before=Pt(0); pr_.paragraph_format.space_after=Pt(4)
                rr_=pr_.add_run(f"Lille, le {date_fr_}"); rr_.font.size=Pt(10); rr_.font.name="Arial"
                _sep(d)

                _P(d,f"{nom_dirigeant}",bold=True,size=11,sb=4,sa=2)
                _P(d,nom_editeur,size=11,sb=0,sa=10)

                to_=d.add_table(rows=1,cols=1); to_.style="Table Grid"
                co_=to_.rows[0].cells[0]; _shd(co_,"D6E4F0"); _brd(co_)
                po_=co_.paragraphs[0]; po_.paragraph_format.space_before=Pt(4); po_.paragraph_format.space_after=Pt(4)
                r1_=po_.add_run("Objet : "); r1_.bold=True; r1_.font.size=Pt(10); r1_.font.name="Arial"; r1_.font.color.rgb=_MR
                r2_=po_.add_run(f"Rapport de pilotage analytique — Exercice {exercice}"); r2_.font.size=Pt(10); r2_.font.name="Arial"

                _P(d,"",sb=10,sa=0)
                _P(d,f"{nom_dirigeant},",bold=True,size=11,sb=0,sa=6)
                _P(d,f"Conformément aux diligences prévues dans notre mission d'accompagnement au pilotage analytique, "
                   f"nous avons l'honneur de vous adresser le présent rapport de pilotage relatif à l'exercice {exercice}, "
                   f"pour la période du {periode_debut} au {periode_fin}.",
                   size=10,align="justify",sb=0,sa=6)
                _P(d,f"Ce rapport porte sur l'ensemble des {nb_eans_} ISBN actifs de votre catalogue, "
                   f"traités par notre outil Python VISION EDITION.",
                   size=10,align="justify",sb=0,sa=10)

                # Section I
                _P(d,"I.   SYNTHÈSE DE LA PERFORMANCE ANALYTIQUE",bold=True,size=12,color=_MR,sb=8,sa=4)
                _sep(d)
                ts_=d.add_table(rows=5,cols=3); ts_.style="Table Grid"
                ts_.rows[0].cells[0]._tc; _wc(ts_.rows[0].cells[0],"Indicateur",bold=True,color=_WH,size=9,fill="1F4E79")
                _wc(ts_.rows[0].cells[1],"Montant",bold=True,color=_WH,size=9,center=True,fill="1F4E79")
                _wc(ts_.rows[0].cells[2],"Commentaire",bold=True,color=_WH,size=9,fill="1F4E79")
                for i,(lbl,val,cmt,fl) in enumerate([
                    ("CA net éditeur",_fn(ca_n_),f"{nb_eans_} ISBN actifs","D6E4F0"),
                    ("Marge brute sur ventes",_fn(mb_,True),f"Taux : {taux_mb_:.1f}%","E2EFDA"),
                    (f"Titres déficitaires",f"{n_neg_} / {nb_eans_}","Action corrective requise","FFDDD9"),
                    ("Résultat analytique net",_fn(mn_,True),"Après charges de structure","FFDDD9" if mn_<0 else "E2EFDA"),
                ]):
                    row_=ts_.rows[i+1]
                    _wc(row_.cells[0],lbl,bold=True,size=9,fill=fl)
                    _wc(row_.cells[1],val,bold=True,center=True,size=10,fill=fl)
                    _wc(row_.cells[2],cmt,italic=True,size=8,fill=fl)

                _P(d,"",sb=10,sa=0)
                # Section II — Top & Flop
                _P(d,"II.   ANALYSE PAR TITRE",bold=True,size=12,color=_MR,sb=8,sa=4)
                _sep(d)
                _P(d,f"Top {nb_top} — titres les plus contributifs",bold=True,size=10,color=_MR,sb=4,sa=4)
                tt_=d.add_table(rows=nb_top+1,cols=4); tt_.style="Table Grid"
                for c_,h_ in enumerate(["Rg","Titre","CA net (EUR)","Marge brute (EUR)"]):
                    _wc(tt_.rows[0].cells[c_],h_,bold=True,color=_WH,size=9,fill="1F4E79")
                for i,rd_ in enumerate(top_list_):
                    bg_=("F5E6DF" if i%2==0 else "FFFFFF")
                    _wc(tt_.rows[i+1].cells[0],str(i+1),bold=True,center=True,size=9,fill=bg_)
                    _wc(tt_.rows[i+1].cells[1],str(rd_["Titre"])[:44],bold=True,size=9,fill=bg_)
                    _wc(tt_.rows[i+1].cells[2],_fn(rd_["CA net"]),center=True,size=9,fill=bg_)
                    _wc(tt_.rows[i+1].cells[3],_fn(rd_["Marge brute"],True),bold=True,center=True,size=10,
                        color=RGBColor(0x21,0x73,0x46),fill="E2EFDA")

                _P(d,"",sb=8,sa=0)
                _P(d,f"Flop {nb_top} — titres déficitaires",bold=True,size=10,color=RGBColor(0xC0,0,0),sb=4,sa=4)
                tf_=d.add_table(rows=nb_top+1,cols=4); tf_.style="Table Grid"
                for c_,h_ in enumerate(["Rg","Titre","CA net (EUR)","Marge brute (EUR)"]):
                    _wc(tf_.rows[0].cells[c_],h_,bold=True,color=_WH,size=9,fill="C00000")
                for i,rd_ in enumerate(flop_list_):
                    bg_=("FFF5F5" if i%2==0 else "FFFFFF")
                    _wc(tf_.rows[i+1].cells[0],str(i+1),bold=True,center=True,size=9,fill=bg_)
                    _wc(tf_.rows[i+1].cells[1],str(rd_["Titre"])[:44],bold=True,size=9,fill=bg_)
                    _wc(tf_.rows[i+1].cells[2],_fn(rd_["CA net"]),center=True,size=9,fill=bg_)
                    _wc(tf_.rows[i+1].cells[3],_fn(rd_["Marge brute"],True),bold=True,center=True,size=10,
                        color=RGBColor(0xC0,0,0),fill="FFDDD9")

                _P(d,"",sb=10,sa=0)
                # Conclusion
                _P(d,"III.   RECOMMANDATIONS",bold=True,size=12,color=_MR,sb=8,sa=4)
                _sep(d)
                _P(d,"Au vu de l'analyse ci-avant, nous vous recommandons d'engager rapidement une analyse "
                   "approfondie des titres déficitaires et de leur impact sur la structure de votre catalogue. "
                   "Nous demeurons à votre entière disposition pour en discuter lors de notre prochaine réunion trimestrielle.",
                   size=10,align="justify",sb=0,sa=16)
                _P(d,"Veuillez agréer, Monsieur, l'expression de nos salutations distinguées.",size=10,sb=0,sa=20)
                _P(d,NOM_CABINET + " — L'Expert-comptable",bold=True,size=10,color=_MR,sb=0,sa=4)
                _P(d,f"Date : {date_fr_}",size=9,sb=0,sa=0)

                # Note méthodologique
                _P(d,"",sb=10,sa=0)
                pn_=d.add_paragraph(); pn_.paragraph_format.space_before=Pt(4); pn_.paragraph_format.space_after=Pt(4)
                rn1_=pn_.add_run("Note méthodologique : "); rn1_.bold=True; rn1_.font.size=Pt(8); rn1_.font.name="Arial"
                rn1_.font.color.rgb=RGBColor(0x55,0x55,0x55)
                sfx=" (titres anonymisés)" if mode_anon else ""
                rn2_=pn_.add_run(f"Rapport généré par Python VISION EDITION — {periode_debut} au {periode_fin} — "
                                  f"{nb_eans_} ISBN actifs{sfx} — données provisoires.")
                rn2_.italic=True; rn2_.font.size=Pt(8); rn2_.font.name="Arial"
                rn2_.font.color.rgb=RGBColor(0x77,0x77,0x77)

                buf_w=BytesIO(); d.save(buf_w); buf_w.seek(0)
                fn_w=f"Rapport_{nom_editeur[:20].replace(' ','_')}_{exercice.replace('/','_')}.docx"
                st.success("✅ Rapport généré !")
                st.download_button("⬇️ Télécharger le rapport Word", data=buf_w, file_name=fn_w,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True, type="primary")

            except Exception as e:
                st.error(f"❌ Erreur : {e}")
                st.exception(e)


# ── TABLEAU DE BORD DIRIGEANT ──
elif role == "dirigeant" and page == "🏠 Mon tableau de bord":
    df, params, dos = get_df_dirigeant()
    st.title(f"📚 {dos['nom']}")
    st.markdown(f"*Tableau de bord de pilotage — Exercice {dos.get('exercice','')}*")

    # Affichage des notifications non lues
    did_dir = st.session_state["dossier_id"]
    notifs_dir = [n for n in dos.get("notifications", []) if not n["lu"]]
    if notifs_dir:
        st.markdown("### 🔔 Messages de votre cabinet")
        for n in notifs_dir:
            if n["type"] == "warning":
                st.warning(f"**{n['date']}** — {n['message']}")
            elif n["type"] == "error":
                st.error(f"**{n['date']}** — {n['message']}")
            else:
                st.info(f"**{n['date']}** — {n['message']}")
        if st.button("✅ Marquer comme lu(s)", key="mark_read"):
            marquer_notifications_lues(did_dir)
            st.rerun()
        st.divider()

    df_v_ = df[mask_ventes(df,params)]; df_r_ = df[mask_retours(df,params)]
    df_rem_ = df[mask_remises(df,params)]; df_c_ = df[mask_charges(df,params)]
    pref_d_ = tuple(params.get("ventes_distributeur") or params.get("ventes",[]))
    df_vd_ = df[df["Compte"].astype(str).str.startswith(pref_d_)] if pref_d_ else df_v_
    ca_d_ = df_vd_["Crédit"].sum()
    ca_b_ = df_v_["Crédit"].sum()
    tr_   = df_r_["Débit"].sum()-df_r_["Crédit"].sum()
    rem_  = df_rem_["Débit"].sum()-df_rem_["Crédit"].sum()
    ca_n_ = ca_b_-tr_-rem_
    ch_t_ = df_c_["Débit"].sum()-df_c_["Crédit"].sum()
    res_  = ca_n_-ch_t_
    taux_r_ = (tr_/ca_d_*100) if ca_d_ else 0

    st.markdown("### Vos indicateurs clés")
    c1,c2,c3,c4 = st.columns(4)
    c1.metric("💶 Chiffre d'affaires net", f"{fmt_fr(ca_n_)} €")
    c2.metric("📊 Résultat net", f"{fmt_fr(res_)} €",
              delta="✅ Bénéficiaire" if res_>0 else "⚠️ Déficitaire",
              delta_color="normal" if res_>0 else "inverse")
    c3.metric("🔁 Taux de retour", f"{taux_r_:.1f} %",
              delta="✅ Normal" if taux_r_<25 else "⚠️ Élevé",
              delta_color="normal" if taux_r_<25 else "inverse")
    titres_d = filtrer_isbn_reels(df)["Code_Analytique"].nunique()
    c4.metric("📚 Titres actifs", str(titres_d))

    st.divider()
    # Évolution mensuelle simplifiée
    df["Date"] = pd.to_datetime(df["Date"], errors="coerce")
    df_ok = df.dropna(subset=["Date"]).copy()
    df_ok["Mois"] = df_ok["Date"].dt.to_period("M").astype(str)
    df_v_ok = df_ok[mask_ventes(df_ok,params)]
    df_r_ok = df_ok[mask_retours(df_ok,params)]
    tv_ = df_v_ok.groupby("Mois")["Crédit"].sum().reset_index().rename(columns={"Crédit":"CA"})
    tr2_ = df_r_ok.groupby("Mois")["Débit"].sum().reset_index().rename(columns={"Débit":"Retours"})
    trend_ = tv_.merge(tr2_,on="Mois",how="left").fillna(0)
    trend_["CA net"] = trend_["CA"]-trend_["Retours"]

    st.markdown("### Évolution de votre chiffre d'affaires")
    fig_d = px.line(trend_, x="Mois", y=["CA","CA net"], markers=True, height=300,
                    labels={"value":"€","variable":""},
                    color_discrete_map={"CA":"#3B82F6","CA net":"#10B981"})
    fig_d.update_layout(margin=dict(t=10), legend=dict(orientation="h"))
    st.plotly_chart(fig_d, use_container_width=True)

    # Message de bienvenue personnalisé
    st.info(f"💡 Pour analyser vos titres en détail, rendez-vous dans **📖 Mes titres**. "
            f"Pour poser une question, utilisez **🤖 Mon assistant IA**.")

# ── MES TITRES (DIRIGEANT) ──
elif role == "dirigeant" and page == "📖 Mes titres":
    df, params, dos = get_df_dirigeant()
    st.header(f"📖 Mes titres — {dos['nom']}")

    titres_dir = sorted(filtrer_isbn_reels(df)["Code_Analytique"].astype(str).unique().tolist())
    if not titres_dir:
        st.warning("Aucun titre détecté dans vos données.")
        st.stop()

    indic_dir = calculer_indicateurs_titres(df, params, titres_dir)

    # Vue simplifiée pour le dirigeant
    st.markdown("### Classement de vos titres")
    indic_dir_aff = indic_dir[["Code_Analytique","Signal","CA net","Marge brute","Taux retour (%)","Résultat net"]].copy()
    indic_dir_aff = indic_dir_aff.sort_values("Marge brute", ascending=False)
    indic_dir_aff.columns = ["ISBN / Titre","Signal","CA net (€)","Marge brute (€)","Taux retour (%)","Résultat net (€)"]

    st.dataframe(
        indic_dir_aff.style.format({
            "CA net (€)":(lambda x: fmt_fr(x)),
            "Marge brute (€)":(lambda x: fmt_fr(x)),
            "Taux retour (%)":(lambda x: f"{x:.1f} %"),
            "Résultat net (€)":(lambda x: fmt_fr(x)),
        }),
        use_container_width=True, hide_index=True
    )

    st.divider()
    st.markdown("### Focus sur un titre")
    isbn_dir = st.selectbox("Choisir un titre", titres_dir)
    row_dir = indic_dir[indic_dir["Code_Analytique"]==isbn_dir].iloc[0]

    c1,c2,c3 = st.columns(3)
    c1.metric("CA net", f"{fmt_fr(row_dir['CA net'])} €")
    c2.metric("Marge brute", f"{fmt_fr(row_dir['Marge brute'])} €")
    c3.metric("Taux de retour", f"{row_dir['Taux retour (%)']:.1f} %")
    st.markdown(f"**Signal : {row_dir['Signal']}**")
    if row_dir["Signal"] == "🔴":
        st.warning("Ce titre présente des difficultés. Parlez-en à votre expert-comptable.")
    elif row_dir["Signal"] == "🟡":
        st.info("Ce titre est à surveiller. Le taux de retour mérite attention.")
    else:
        st.success("Ce titre est rentable et performant.")

# ── MA TRÉSORERIE (DIRIGEANT) ──
elif role == "dirigeant" and page == "💰 Ma trésorerie":
    df, params, dos = get_df_dirigeant()
    st.header(f"💰 Ma trésorerie — {dos['nom']}")

    if "treso_real" not in st.session_state:
        st.info("Votre expert-comptable n'a pas encore généré les projections de trésorerie. "
                "Elles seront disponibles après la prochaine mise à jour de vos données.")
        st.stop()

    treso_ = st.session_state["treso_real"]
    _MOIS = {1:"Jan",2:"Fév",3:"Mar",4:"Avr",5:"Mai",6:"Jui",7:"Jul",8:"Aoû",9:"Sep",10:"Oct",11:"Nov",12:"Déc"}
    def _ml(p): return f"{_MOIS.get(p.month,str(p.month))} {p.year}"

    st.metric("Trésorerie actuelle", f"{fmt_fr(treso_.iloc[-1])} €",
              delta="✅ Positive" if treso_.iloc[-1]>0 else "⚠️ Négative",
              delta_color="normal" if treso_.iloc[-1]>0 else "inverse")

    fig_tr = go.Figure()
    fig_tr.add_trace(go.Scatter(x=[_ml(m) for m in treso_.index], y=treso_.values,
                                fill="tozeroy", name="Trésorerie",
                                line=dict(color="#3B82F6",width=2)))
    fig_tr.add_hline(y=0, line_dash="dash", line_color="red", annotation_text="Seuil zéro")
    fig_tr.update_layout(height=320, margin=dict(t=10), xaxis_title="", yaxis_title="€")
    st.plotly_chart(fig_tr, use_container_width=True)
    st.caption("Trésorerie reconstituée à partir de votre comptabilité analytique par votre cabinet.")

# ── RETOURS (DIRIGEANT) ──
elif role == "dirigeant" and page == "📦 Retours & Remises":
    df, params, dos = get_df_dirigeant()
    st.header(f"📦 Retours & Remises — {dos['nom']}")

    df_r_ = df[mask_retours(df,params)]; df_v_ = df[mask_ventes(df,params)]
    pref_d_ = tuple(params.get("ventes_distributeur") or params.get("ventes",[]))
    df_vd_ = df[df["Compte"].astype(str).str.startswith(pref_d_)] if pref_d_ else df_v_
    ca_d_ = df_vd_["Crédit"].sum()
    tr_ = df_r_["Débit"].sum()-df_r_["Crédit"].sum() if not df_r_.empty else 0
    taux_r_ = (tr_/ca_d_*100) if ca_d_ else 0

    c1,c2 = st.columns(2)
    c1.metric("Total retours", f"{fmt_fr(tr_)} €")
    c2.metric("Taux de retour", f"{taux_r_:.1f} %",
              delta="✅ Normal" if taux_r_<25 else "⚠️ Élevé",
              delta_color="normal" if taux_r_<25 else "inverse")

    if taux_r_ > 30:
        st.error("⚠️ Votre taux de retour est élevé. Votre expert-comptable vous contactera pour en discuter.")
    elif taux_r_ > 25:
        st.warning("Votre taux de retour mérite attention. Parlez-en lors de votre prochaine réunion.")

    if not df_r_.empty:
        df_r_["Date"] = pd.to_datetime(df_r_["Date"], errors="coerce")
        df_r_["Mois"] = df_r_["Date"].dt.strftime("%Y-%m")
        trend_r_ = df_r_.groupby("Mois")["Débit"].sum().abs().reset_index()
        fig_r_ = px.bar(trend_r_, x="Mois", y="Débit", title="Retours par mois",
                        color_discrete_sequence=["#EF4444"], height=280)
        st.plotly_chart(fig_r_, use_container_width=True)

# ── DROITS D'AUTEURS (DIRIGEANT) ──
elif role == "dirigeant" and page == "✍️ Droits d'auteurs":
    df, params, dos = get_df_dirigeant()
    st.header(f"✍️ Droits d'auteurs — {dos['nom']}")
    st.info("Cette section affiche vos droits d'auteurs comptabilisés. "
            "Pour un relevé détaillé par auteur, contactez votre cabinet.")

    df["Date"] = pd.to_datetime(df["Date"], errors="coerce")
    mask_db = df["Compte"].astype(str).str.strip()=="604300000"
    mask_av = df["Compte"].astype(str).str.strip()=="409600"

    total_db = df[mask_db]["Débit"].sum()-df[mask_db]["Crédit"].sum()
    av_d = df[mask_av]["Débit"].sum(); av_c = df[mask_av]["Crédit"].sum()
    av_restant = av_d-av_c

    c1,c2 = st.columns(2)
    c1.metric("Droits d'auteurs comptabilisés", f"{fmt_fr(total_db,2)} €")
    c2.metric("À-valoirs restant à amortir", f"{fmt_fr(av_restant,2)} €",
              delta="⚠️ À surveiller" if av_restant>0 else "✅ Amortis",
              delta_color="off")

    if av_restant > 0:
        st.info(f"Un à-valoir de {fmt_fr(av_restant,2)} € reste à amortir sur votre catalogue. "
                f"Votre expert-comptable suit cet indicateur pour vous alerter si nécessaire.")

# ── ASSISTANT IA DIRIGEANT ──
elif role == "dirigeant" and page == "🤖 Mon assistant IA":
    df, params, dos = get_df_dirigeant()
    st.header(f"🤖 Mon assistant — {dos['nom']}")
    st.markdown("*Posez vos questions sur votre activité éditoriale en langage naturel.*")

    client_ai = get_client_ai()
    if client_ai is None:
        st.error("⚠️ Assistant IA non disponible. Contactez votre cabinet.")
        st.stop()

    # Contexte simplifié pour le dirigeant
    df_v_ = df[mask_ventes(df,params)]; df_r_ = df[mask_retours(df,params)]
    ca_b_ = df_v_["Crédit"].sum()
    tr_   = df_r_["Débit"].sum()-df_r_["Crédit"].sum()
    ca_n_ = ca_b_-tr_
    taux_r_ = (tr_/ca_b_*100) if ca_b_ else 0
    titres_dir_ = sorted(filtrer_isbn_reels(df)["Code_Analytique"].astype(str).unique().tolist())
    indic_dir_ = calculer_indicateurs_titres(df,params,titres_dir_) if titres_dir_ else pd.DataFrame()

    top5_ = ""
    if not indic_dir_.empty:
        top5_rows = indic_dir_.sort_values("CA net",ascending=False).head(5)
        top5_ = "\n".join([f"- {r['Code_Analytique']} : CA net {fmt_fr(r['CA net'])} EUR, "
                           f"marge brute {fmt_fr(r['Marge brute'])} EUR, retour {r['Taux retour (%)']:.1f}%"
                           for _,r in top5_rows.iterrows()])

    data_ctx_dir = f"""
DONNÉES DE {dos['nom']} — {dos.get('exercice','')}
CA brut : {fmt_fr(ca_b_)} EUR
Total retours : {fmt_fr(tr_)} EUR  
CA net : {fmt_fr(ca_n_)} EUR
Taux de retour global : {taux_r_:.1f} %
Nombre de titres actifs : {len(titres_dir_)}

Top 5 titres par CA net :
{top5_}
"""
    sys_dir = SYSTEM_PROMPT_DIRIGEANT + f"\n\nDONNÉES :\n{data_ctx_dir}"

    st.markdown("**Questions fréquentes :**")
    col_q1, col_q2, col_q3 = st.columns(3)
    questions_dir = [
        ("Quel est mon CA par titre ?", col_q1),
        ("Quel est mon taux de retour ?", col_q2),
        ("Quels sont mes titres rentables ?", col_q3),
    ]
    for q_text, col_q in questions_dir:
        if col_q.button(q_text, use_container_width=True):
            with st.spinner("Réflexion..."):
                resp_q = client_ai.messages.create(
                    model="claude-sonnet-4-6", max_tokens=500,
                    system=sys_dir, messages=[{"role":"user","content":q_text}]
                )
                st.info(f"**{q_text}**")
                st.success(resp_q.content[0].text)

    st.divider()
    for msg in st.session_state["messages_agent"]:
        with st.chat_message(msg["role"], avatar="🧑" if msg["role"]=="user" else "📚"):
            st.markdown(msg["content"])

    if prompt_dir := st.chat_input("Posez votre question..."):
        st.session_state["messages_agent"].append({"role":"user","content":prompt_dir})
        with st.chat_message("user", avatar="🧑"):
            st.markdown(prompt_dir)
        with st.chat_message("assistant", avatar="📚"):
            with st.spinner("Je cherche la réponse..."):
                resp_dir = client_ai.messages.create(
                    model="claude-sonnet-4-6", max_tokens=600,
                    system=sys_dir, messages=st.session_state["messages_agent"]
                )
                ans_dir = resp_dir.content[0].text
                st.markdown(ans_dir)
                st.session_state["messages_agent"].append({"role":"assistant","content":ans_dir})

    if st.button("🗑️ Effacer la conversation"):
        st.session_state["messages_agent"] = []
        st.rerun()




# ── PLAN D'ACTION (EC) ──
elif role == "ec" and page == "📋 Plan d'action":
    df, params = check_pivot()
    dos = get_dossier(st.session_state["dossier_id"])
    st.header(f"📋 Plan d'action — {dos['nom']}")
    st.markdown(f"*Exercice {dos.get('exercice','')} — Réunion trimestrielle de pilotage*")

    # ── Calculs depuis le pivot ──
    df_v_pa   = df[mask_ventes(df,params)]
    df_r_pa   = df[mask_retours(df,params)]
    df_rem_pa = df[mask_remises(df,params)]
    df_c_pa   = df[mask_charges(df,params)]
    df_prov_pa= df[mask_provisions_reprises(df,params)]
    net_p_pa  = df_prov_pa["Crédit"].sum()-df_prov_pa["Débit"].sum()

    ca_b_pa = df_v_pa["Crédit"].sum()-df_v_pa["Débit"].sum()
    tr_pa   = df_r_pa["Débit"].sum()-df_r_pa["Crédit"].sum()
    rem_pa  = df_rem_pa["Débit"].sum()-df_rem_pa["Crédit"].sum()
    ca_n_pa = ca_b_pa-tr_pa-rem_pa
    ch_pa   = (df_c_pa["Débit"].sum()-df_c_pa["Crédit"].sum())-net_p_pa

    mask_ci_pa  = df["Code_Analytique"].astype(str).str.upper().isin(["CHARGES INDIRECTES","FRAIS GENERAUX"])
    mask_cir_pa = df["Compte"].astype(str)==COMPTE_CHARGES_INDIRECTES_REPARTIES
    ci_pa = df[mask_ci_pa|mask_cir_pa]["Débit"].sum()-df[mask_ci_pa|mask_cir_pa]["Crédit"].sum()
    mb_pa = ca_n_pa-(ch_pa-ci_pa) if ci_pa>0 else ca_n_pa-ch_pa
    res_pa= ca_n_pa-ch_pa
    taux_mb_pa = (mb_pa/ca_n_pa*100) if ca_n_pa else 0

    titres_pa = sorted(filtrer_isbn_reels(df)["Code_Analytique"].astype(str).unique().tolist())
    nb_isbn_pa = len(titres_pa)
    ci_par_titre_pa = ci_pa/nb_isbn_pa if nb_isbn_pa>0 else 0
    indic_pa = calculer_indicateurs_titres(df,params,titres_pa) if titres_pa else pd.DataFrame()
    top5_pa  = indic_pa.sort_values("Marge brute",ascending=False).head(5) if not indic_pa.empty else pd.DataFrame()
    flop5_pa = indic_pa.sort_values("Marge brute",ascending=True).head(5)  if not indic_pa.empty else pd.DataFrame()
    n_neg_pa = int((indic_pa["Marge brute"]<0).sum()) if not indic_pa.empty else 0
    n_pos_pa = int((indic_pa["Marge brute"]>=0).sum()) if not indic_pa.empty else 0

    mode_anon_pa = st.session_state.get("mode_anonyme",False)
    mapping_pa   = obtenir_mapping_anonymisation(df) if mode_anon_pa and titres_pa else {}
    def _t(code): return mapping_pa.get(code,code) if mode_anon_pa else code

    # ── Affichage contexte ──
    st.subheader("📊 Données issues de VISION ÉDITION")
    c1,c2,c3,c4 = st.columns(4)
    c1.metric("Résultat analytique net",f"{fmt_fr(res_pa)} €",
              delta="⚠️ Déficitaire" if res_pa<0 else "✅ Bénéficiaire",
              delta_color="inverse" if res_pa<0 else "normal")
    c2.metric("Marge brute globale",f"{fmt_fr(mb_pa)} €",f"{taux_mb_pa:.1f}%")
    c3.metric("Charges indirectes",f"{fmt_fr(ci_pa)} €",
              f"{(ci_pa/ca_n_pa*100):.1f}% du CA net — {ci_par_titre_pa:.0f} EUR/titre" if ca_n_pa else "")
    c4.metric("ISBN actifs",f"{nb_isbn_pa}",f"{n_neg_pa} déficitaires / {n_pos_pa} bénéficiaires")

    col_top,col_flop = st.columns(2)
    with col_top:
        st.markdown("**🏆 Top 5 porteurs**")
        for i,(_,row) in enumerate(top5_pa.iterrows()):
            st.write(f"{i+1}. **{_t(row['Code_Analytique'])}** — marge {fmt_fr(row['Marge brute'])} € "
                     f"(retour {row['Taux retour (%)']: .1f}%)")
    with col_flop:
        st.markdown("**⚠️ Flop 5 déficitaires**")
        for i,(_,row) in enumerate(flop5_pa.iterrows()):
            st.write(f"{i+1}. **{_t(row['Code_Analytique'])}** — perte {fmt_fr(row['Marge brute'])} € "
                     f"(retour {row['Taux retour (%)']: .1f}%)")

    st.divider()
    st.subheader("⚙️ Paramètres de la réunion")
    col_p1,col_p2 = st.columns(2)
    with col_p1:
        date_reunion   = st.date_input("Date de la réunion",value=_dt.date.today(),key="pa_date_reunion")
        date_prochaine = st.date_input("Prochaine réunion",key="pa_date_prochaine")
    with col_p2:
        mode_anon_exp = st.checkbox("Anonymiser les titres dans l'export",value=mode_anon_pa,key="pa_anon")
        st.caption("Les titres seront présentés sous forme T1, T2... si coché.")

    st.divider()
    st.info("Le plan d'action sera généré avec les données réelles de VISION ÉDITION. "
            "Les zones jaunes restent modifiables dans Excel pour saisir les décisions en réunion.")

    if st.button("📋 Générer le plan d'action (Excel)", type="primary", use_container_width=True):
        with st.spinner("Génération en cours..."):
            try:
                # Helpers Excel
                def _fill(c): return _OPFill("solid",fgColor=c)
                def _font(b=False,s=10,c="000000",i=False,bold=None): return _OPFont(bold=bold if bold is not None else b,size=s,color=c,italic=i,name="Arial")
                def _align(h="left",v="center",w=True): return _OPAlign(horizontal=h,vertical=v,wrap_text=w)
                def _brd(c="CCCCCC"):
                    sd=_OPSide(style="thin",color=c); return _OPBorder(left=sd,right=sd,top=sd,bottom=sd)
                def _auto_width(ws_):
                    """Ajuste automatiquement la largeur de chaque colonne selon le contenu."""
                    from openpyxl.utils import get_column_letter
                    for col_cells in ws_.columns:
                        max_len = 0
                        col_letter = get_column_letter(col_cells[0].column)
                        for cell in col_cells:
                            if cell.value:
                                # Prendre la ligne la plus longue si contenu multi-lignes
                                lines = str(cell.value).split('\n')
                                cell_max = max(len(l) for l in lines)
                                max_len = max(max_len, cell_max)
                        # Plafonner à 60 caractères, minimum 8
                        adj = min(max(max_len + 2, 8), 60)
                        ws_.column_dimensions[col_letter].width = adj
                def _sc(ws_,row,col,value="",bg=None,bold=False,size=10,color="000000",
                        italic=False,h="left",v="center",wrap=True):
                    cell=ws_.cell(row=row,column=col,value=value)
                    if bg: cell.fill=_fill(bg)
                    cell.font=_font(bold,size,color,italic)
                    cell.alignment=_align(h,v,wrap); cell.border=_brd(); return cell
                def _mg(ws_,r1,c1,r2,c2,value="",bg=None,bold=False,size=10,
                        color="000000",italic=False,h="center",v="center"):
                    ws_.merge_cells(start_row=r1,start_column=c1,end_row=r2,end_column=c2)
                    cell=ws_.cell(row=r1,column=c1,value=value)
                    if bg: cell.fill=_fill(bg)
                    cell.font=_font(bold,size,color,italic)
                    cell.alignment=_align(h,v,True); cell.border=_brd(); return cell
                def _blank(ws_,r,nc,bg="FFFFFF",h=8):
                    ws_.row_dimensions[r].height=h
                    for c in range(1,nc+1): ws_.cell(row=r,column=c).fill=_fill(bg)
                def _sec(ws_,r,nc,let,tit,sous=""):
                    ws_.merge_cells(start_row=r,start_column=2,end_row=r,end_column=nc)
                    cell=ws_.cell(row=r,column=2,value=f"{let}.  {tit.upper()}")
                    cell.fill=_fill("1F4E79"); cell.font=_font(bold=True,s=12,c="FFFFFF")
                    cell.alignment=_align("left","center",False); cell.border=_brd()
                    ws_.row_dimensions[r].height=22; r+=1
                    if sous:
                        ws_.merge_cells(start_row=r,start_column=2,end_row=r,end_column=nc)
                        c2_=ws_.cell(row=r,column=2,value=sous)
                        c2_.fill=_fill("D6E4F0"); c2_.font=_font(s=9,c="444444",i=True)
                        c2_.alignment=_align("left","center",False); c2_.border=_brd()
                        ws_.row_dimensions[r].height=18; r+=1
                    return r
                def _hdrs(ws_,r,lst):
                    for c,h in enumerate(lst,2):
                        cell=ws_.cell(row=r,column=c,value=h)
                        cell.fill=_fill("1F4E79"); cell.font=_font(bold=True,s=9,c="FFFFFF")
                        cell.alignment=_align("center","center",True); cell.border=_brd("888888")
                    ws_.row_dimensions[r].height=22; return r+1

                NOM_PA=dos["nom"]; EXER_PA=dos.get("exercice","")
                date_r_s=date_reunion.strftime("%d/%m/%Y")
                date_p_s=date_prochaine.strftime("%d/%m/%Y") if date_prochaine else "___/___/______"
                MARINE_PA="1F4E79"; TERRA_PA="C0522A"

                wb_pa=openpyxl.Workbook()
                ws_pa=wb_pa.active; ws_pa.title="Plan d'action"
                ws_pa.sheet_view.showGridLines=False
                for col,w in [("A",2),("B",6),("C",8),("D",26),("E",28),("F",14),("G",12),("H",14),("I",22),("J",3)]:
                    ws_pa.column_dimensions[col].width=w
                for r_ in range(1,130):
                    ws_pa.row_dimensions[r_].height=16
                    for c in range(1,11): ws_pa.cell(row=r_,column=c).fill=_fill("FFFFFF")

                # Titre
                _mg(ws_pa,2,2,3,9,"OUTIL 8 — PLAN D'ACTION STRUCTURÉ",bg=MARINE_PA,bold=True,size=15,color="FFFFFF")
                ws_pa.row_dimensions[2].height=26; ws_pa.row_dimensions[3].height=26
                _mg(ws_pa,4,2,4,9,
                    f"CAB ÉDITION  /  {NOM_PA} — Exercice {EXER_PA}  |  "
                    f"Réunion de pilotage : {date_r_s}  |  Prochaine réunion : {date_p_s}",
                    bg=TERRA_PA,bold=False,size=10,color="FFFFFF",italic=True,h="left")
                ws_pa.row_dimensions[4].height=18
                _blank(ws_pa,5,10,h=6)

                # KPIs
                _mg(ws_pa,6,2,6,9,"CONTEXTE DE L'ANALYSE — DONNÉES PYTHON VISION ÉDITION",bg=MARINE_PA,bold=True,size=10,color="FFFFFF",h="left")
                ws_pa.row_dimensions[6].height=18
                kpis_=[
                    ("Résultat analytique net",f"{'+ ' if res_pa>=0 else '− '}{fmt_fr(abs(res_pa))} EUR","après charges indirectes","FFDDD9" if res_pa<0 else "E2EFDA","C00000" if res_pa<0 else "217346"),
                    ("Marge brute globale",f"+ {fmt_fr(mb_pa)} EUR",f"soit +{taux_mb_pa:.1f}% du CA net","E2EFDA","217346"),
                    ("Charges indirectes",f"{fmt_fr(ci_pa)} EUR",f"{ci_par_titre_pa:.0f} EUR/titre — {(ci_pa/ca_n_pa*100):.1f}% du CA net" if ca_n_pa else "","FFF2CC","A67C00"),
                    ("ISBN actifs",f"{nb_isbn_pa} titres",f"{n_neg_pa} déficitaires / {n_pos_pa} bénéficiaires","D6E4F0","1F4E79"),
                ]
                for i,(lbl,val,cmt,bg_k,col_k) in enumerate(kpis_):
                    cs=2+i*2
                    ws_pa.merge_cells(start_row=7,start_column=cs,end_row=8,end_column=cs)
                    ws_pa.merge_cells(start_row=7,start_column=cs+1,end_row=7,end_column=cs+1)
                    cl=ws_pa.cell(row=7,column=cs,value=lbl); cl.fill=_fill(bg_k); cl.font=_font(bold=True,s=9,c=col_k); cl.alignment=_align("center","center",False); cl.border=_brd()
                    cv=ws_pa.cell(row=7,column=cs+1,value=val); cv.fill=_fill(bg_k); cv.font=_font(bold=True,s=13,c=col_k); cv.alignment=_align("center","center",False); cv.border=_brd()
                    cc=ws_pa.cell(row=8,column=cs+1,value=cmt); cc.fill=_fill(bg_k); cc.font=_font(s=8,c=col_k,i=True); cc.alignment=_align("center","center",False); cc.border=_brd()
                ws_pa.row_dimensions[7].height=22; ws_pa.row_dimensions[8].height=16
                _blank(ws_pa,9,10,h=8)

                dv_pa=_OPDV(type="list",formula1='"Pré-rempli EC,▶ À arbitrer,✓ Acté conjointement,À faire,Reporté"',showDropDown=False,showErrorMessage=False)
                ws_pa.add_data_validation(dv_pa)
                r_=10

                # Préparer top5/flop5 strings
                top3_str=" | ".join([f"{_t(row['Code_Analytique'])} (+{fmt_fr(row['Marge brute'])} EUR)" for _,row in top5_pa.head(3).iterrows()]) if not top5_pa.empty else "Données non disponibles"
                flop5_str=" | ".join([f"{_t(row['Code_Analytique'])} ({fmt_fr(row['Marge brute'])} EUR)" for _,row in flop5_pa.iterrows()]) if not flop5_pa.empty else "Données non disponibles"
                flop1=_t(flop5_pa.iloc[0]["Code_Analytique"]) if not flop5_pa.empty else "Titre 1"
                top1=_t(top5_pa.iloc[0]["Code_Analytique"]) if not top5_pa.empty else "Titre 1"

                # SECTION A
                r_=_sec(ws_pa,r_,9,"A","PRÉPARATION EN AMONT","Réalisée par l'EC seul avant la réunion — information au dirigeant, pas de décision requise")
                r_=_hdrs(ws_pa,r_,["N°","Action préparée par l'EC","Indicateur source","Resp.","Échéance","Statut","Note EC / Information au dirigeant"])
                _pa_actions_A = [
                    ("A1", "Extraction et analyse du catalogue via Python VISION ÉDITION\n"
                            f"Top 3 porteurs : {top3_str}",
                     "Marge par titre","EC seul","Avant réunion","Pré-rempli EC",
                     "Rapport de pilotage (Outil 7) généré et transmis J-5","D6E4F0"),
                    ("A2", "Sélection des titres à analyser conjointement\n"
                            f"Flop 5 : {flop5_str}",
                     "Marge par titre","EC seul","J-5","Pré-rempli EC",
                     "Sélection présentée au dirigeant en début de réunion","FFFFFF"),
                    ("A3", "Préparation des 3 scénarios de simulation via le Simulateur de rentabilité\n"
                            "D1 : Stabilisation, D2 : Réorientation catalogue, D3 : Optimisation charges",
                     "Tous indicateurs","EC seul","J-5","Pré-rempli EC",
                     "Simulations disponibles sur Python VISION ÉDITION pendant la réunion","D6E4F0"),
                ]
                for num,action,indic,resp,ech,statut,note,bg_ in _pa_actions_A:
                    _sc(ws_pa,r_,2,num,bg="D6E4F0",bold=True,size=10,color="1F4E79",h="center")
                    _sc(ws_pa,r_,3,action,bg=bg_,size=9,color="222222",wrap=True)
                    _sc(ws_pa,r_,4,indic,bg=bg_,bold=True,size=9,color="1F4E79",italic=True)
                    _sc(ws_pa,r_,5,resp,bg=bg_,size=9,color="555555",h="center")
                    _sc(ws_pa,r_,6,ech,bg=bg_,bold=True,size=9,color="1F4E79",h="center")
                    cst=ws_pa.cell(row=r_,column=7,value=statut); cst.fill=_fill("D6E4F0"); cst.font=_font(bold=True,s=9,c="1F4E79"); cst.alignment=_align("center","center",False); cst.border=_brd(); dv_pa.add(cst)
                    _sc(ws_pa,r_,8,note,bg=bg_,size=9,color="444444",italic=True,wrap=True)
                    ws_pa.row_dimensions[r_].height=48; r_+=1
                _blank(ws_pa,r_,10,h=10); r_+=1

                # SECTION B
                r_=_sec(ws_pa,r_,9,"B","ACTIONS À COURT TERME","Décisions à prendre conjointement en réunion - cellules jaunes à compléter pendant la séance")
                r_=_hdrs(ws_pa,r_,["N°","Action","Indicateur source","Resp.","Échéance","Statut","Décision / Observation à noter en réunion"])
                _pa_b1 = (f"Arbitrage sur les {n_neg_pa} titres déficitaires\n"
                           f"Pour chaque titre : maintien, arrêt commercial — Priorité : {flop1}")
                _pa_b2 = ("Gel des réimpressions pour les titres déficitaires\n"
                           "Aucune nouvelle commande tant que la décision n\'est pas levée")
                _pa_b3 = (f"Validation du plan de réimpression des titres porteurs\n"
                           f"Priorité : {top1} — vérification stocks disponibles")
                _pa_b4 = (f"Audit des charges indirectes : {fmt_fr(ci_pa)} EUR "
                           f"({(ci_pa/ca_n_pa*100):.1f}% CA net)\nIdentifier les postes compressibles")
                for num,action,indic,resp,ech,statut,note,bg_ in [
                    ("B1",_pa_b1,"Marge par titre","EC + Dir.","Immédiat","▶ À arbitrer","Décision : ___________________________________","FCE4D6"),
                    ("B2",_pa_b2,"Marge par titre","Dirigeant","Immédiat","▶ À arbitrer","Titres concernés : ___________________________","FFF5EE"),
                    ("B3",_pa_b3,"Marge par titre","Dirigeant","T+1 mois","▶ À arbitrer","Quantités retenues : ________________________","FCE4D6"),
                    ("B4",_pa_b4,"Charges indirectes","EC","T+2 mois","▶ À arbitrer","Périmètre de l\'audit : ______________________","FFF5EE"),
                ]:
                    _sc(ws_pa,r_,2,num,bg="C0522A",bold=True,size=10,color="FFFFFF",h="center")
                    _sc(ws_pa,r_,3,action,bg=bg_,size=9,color="222222",wrap=True)
                    _sc(ws_pa,r_,4,indic,bg=bg_,bold=True,size=9,color="1F4E79",italic=True)
                    _sc(ws_pa,r_,5,resp,bg=bg_,size=9,color="555555",h="center")
                    _sc(ws_pa,r_,6,ech,bg=bg_,bold=True,size=9,color="C00000",h="center")
                    cst=ws_pa.cell(row=r_,column=7,value=statut); cst.fill=_fill("FCE4D6"); cst.font=_font(bold=True,s=9,c="A67C00"); cst.alignment=_align("center","center",False); cst.border=_brd("C0522A"); dv_pa.add(cst)
                    cn=ws_pa.cell(row=r_,column=8,value=note); cn.fill=_fill("FFF2CC"); cn.font=_font(s=9,c="333333"); cn.alignment=_align("left","center",True); cn.border=_brd("C0522A")
                    ws_pa.row_dimensions[r_].height=48; r_+=1
                _blank(ws_pa,r_,10,h=10); r_+=1

                # SECTION C
                r_=_sec(ws_pa,r_,9,"C","ACTIONS À MOYEN TERME","Actions à valider en réunion — cellules jaunes à compléter")
                r_=_hdrs(ws_pa,r_,["N°","Action","Indicateur source","Resp.","Échéance","Statut","Décision / Observation à noter en réunion"])
                ci_cible=fmt_fr(ci_pa*0.7)
                _pa_c1 = ("Révision des contrats d\'auteurs déficitaires\n"
                           "Renégociation taux droits si à-valoir non amorti après 12 mois")
                _pa_c2 = ("Optimisation du catalogue actif\n"
                           "Archivage titres sans ventes depuis 24 mois — réduction du catalogue")
                _pa_c3 = ("Préparation du dossier de subvention CNL et Région\n"
                           "Anticiper la constitution du dossier")
                _pa_c4 = f"Réduction des charges indirectes\nObjectif : CI < {ci_cible} EUR"
                for num,action,indic,resp,ech,statut,note,bg_ in [
                    ("C1",_pa_c1,"Droits d\'auteurs","Dirigeant","T+3 mois","▶ À arbitrer","Titres concernés : ___________________________","F5E6DF"),
                    ("C2",_pa_c2,"Marge par titre","EC + Dir.","T+3 mois","▶ À arbitrer","Scénario retenu : D1 / D2 / D3 (entourer)","FFFFFF"),
                    ("C3",_pa_c3,"Trésorerie prév.","Dirigeant","T+6 mois","À faire","Montant cible subvention : __________________","F5E6DF"),
                    ("C4",_pa_c4,"Charges indirectes","EC + Dir.","T+6 mois","▶ À arbitrer","Postes retenus pour réduction : _____________","FFFFFF"),
                ]:
                    _sc(ws_pa,r_,2,num,bg="C0522A",bold=True,size=10,color="FFFFFF",h="center")
                    _sc(ws_pa,r_,3,action,bg=bg_,size=9,color="222222",wrap=True)
                    _sc(ws_pa,r_,4,indic,bg=bg_,bold=True,size=9,color="1F4E79",italic=True)
                    _sc(ws_pa,r_,5,resp,bg=bg_,size=9,color="555555",h="center")
                    _sc(ws_pa,r_,6,ech,bg=bg_,bold=True,size=9,color="A67C00",h="center")
                    cst=ws_pa.cell(row=r_,column=7,value=statut); cst.fill=_fill("FCE4D6" if "arbitrer" in statut else "FFF2CC"); cst.font=_font(bold=True,s=9,c="A67C00"); cst.alignment=_align("center","center",False); cst.border=_brd(); dv_pa.add(cst)
                    cn=ws_pa.cell(row=r_,column=8,value=note); cn.fill=_fill("FFF2CC"); cn.font=_font(s=9,c="333333"); cn.alignment=_align("left","center",True); cn.border=_brd("C0522A")
                    ws_pa.row_dimensions[r_].height=48; r_+=1
                _blank(ws_pa,r_,10,h=10); r_+=1

                # SECTION D
                r_=_sec(ws_pa,r_,9,"D","ORIENTATIONS STRATÉGIQUES — Horizon 12 mois","Scénarios préparés par l'EC — arbitrage du dirigeant requis en séance")
                r_=_hdrs(ws_pa,r_,["N°","Orientation stratégique","Scénario","Impact tréso.","Impact marge","Statut","Décision dirigeant à noter en réunion"])
                _pd1 = ("Maintien de la stratégie éditoriale actuelle\n"
                         "Focus amélioration rentabilité titre par titre")
                _pd2 = ("Réorientation vers un catalogue resserré\n"
                         "Réduction des nouveautés, focus sur les titres porteurs")
                _pd3 = f"Réduction des charges indirectes\nObjectif : CI < {ci_cible} EUR - résultat net cible > 0"
                for num,orient,scen,it,im,stat,dec,bg_ in [
                    ("D1",_pd1,"Scénario 1\nStabilisation","Neutre","+3 à 5 pts","Pré-rempli EC","Décision : Retenu / Écarté / À approfondir","D6E4F0"),
                    ("D2",_pd2,"Scénario 2\nRéorientation","Positif\n(+ tréso.)","+8 à 12 pts","Pré-rempli EC","Décision : Retenu / Écarté / À approfondir","F5E6DF"),
                    ("D3",_pd3,"Scénario 3\nOptimisation","Positif","+5 à 8 pts","Pré-rempli EC","Décision : Retenu / Écarté / À approfondir","D6E4F0"),
                ]:
                    _sc(ws_pa,r_,2,num,bg="1F4E79",bold=True,size=10,color="FFFFFF",h="center")
                    _sc(ws_pa,r_,3,orient,bg=bg_,size=9,color="222222",wrap=True)
                    _sc(ws_pa,r_,4,scen,bg=bg_,bold=True,size=9,color="1F4E79",h="center",wrap=True)
                    _sc(ws_pa,r_,5,it,bg="E2EFDA",bold=True,size=9,color="217346",h="center",wrap=True)
                    _sc(ws_pa,r_,6,im,bg="E2EFDA",bold=True,size=9,color="217346",h="center")
                    cst=ws_pa.cell(row=r_,column=7,value=stat); cst.fill=_fill("D6E4F0"); cst.font=_font(bold=True,s=9,c="1F4E79"); cst.alignment=_align("center","center",False); cst.border=_brd(); dv_pa.add(cst)
                    cn=ws_pa.cell(row=r_,column=8,value=dec); cn.fill=_fill("FFF2CC"); cn.font=_font(s=9,c="333333"); cn.alignment=_align("left","center",True); cn.border=_brd("C0522A")
                    ws_pa.row_dimensions[r_].height=48; r_+=1
                _blank(ws_pa,r_,10,h=10); r_+=1

                # SECTION E
                r_=_sec(ws_pa,r_,9,"E","SUIVI DES DÉCISIONS — Compte-rendu de réunion","À compléter en fin de réunion — constitue le compte-rendu officiel signé par les deux parties")
                for rubrique,contenu,bg_,mod in [
                    ("Constats principaux",f"Marge brute {fmt_fr(mb_pa)} EUR ({taux_mb_pa:.1f}%) - Résultat net {fmt_fr(res_pa)} EUR - {n_neg_pa} titres déficitaires sur {nb_isbn_pa}","D6E4F0",False),
                    ("Décisions actées","À compléter en réunion : _______________________________________________","FFF2CC",True),
                    ("Scénario stratégique retenu","Scénario retenu (entourer) :  D1 - Stabilisation  /  D2 - Réorientation  /  D3 - Optimisation  /  Reporté","FFF2CC",True),
                    ("Points en attente","À compléter en réunion : _______________________________________________","FFF2CC",True),
                    ("Prochain RDV",f"Date : {date_p_s}   Lieu : ___________________   Ordre du jour : ___________________","FFF2CC",True),
                ]:
                    ws_pa.merge_cells(start_row=r_,start_column=2,end_row=r_,end_column=3)
                    cr_=ws_pa.cell(row=r_,column=2,value=rubrique); cr_.fill=_fill("1F4E79" if not mod else "555555"); cr_.font=_font(bold=True,s=9,c="FFFFFF"); cr_.alignment=_align("left","center",False); cr_.border=_brd()
                    ws_pa.merge_cells(start_row=r_,start_column=4,end_row=r_,end_column=9)
                    cc_=ws_pa.cell(row=r_,column=4,value=contenu); cc_.fill=_fill(bg_); cc_.font=_font(s=9,c="222222",i=not mod); cc_.alignment=_align("left","center",True); cc_.border=_brd("C0522A" if mod else "CCCCCC")
                    ws_pa.row_dimensions[r_].height=28; r_+=1
                _blank(ws_pa,r_,10,h=15); r_+=1

                # Signatures
                _mg(ws_pa,r_,2,r_,9,"SIGNATURES - Validation du plan d'action",bg="1F4E79",bold=True,size=10,color="FFFFFF",h="left")
                ws_pa.row_dimensions[r_].height=18; r_+=1
                ws_pa.merge_cells(start_row=r_,start_column=2,end_row=r_+2,end_column=5)
                _ce_val = ("L\'Expert-Comptable — CAB ÉDITION\n\n"
                            "Signature : ____________________________\n\n"
                            f"Date : {date_r_s}")
                _cd_val = (f"Le Dirigeant — {NOM_PA}\n\n"
                            "Signature : ____________________________\n\n"
                            "Date : _________________________________")
                ce_=ws_pa.cell(row=r_,column=2,value=_ce_val)
                ce_.fill=_fill("D6E4F0"); ce_.font=_font(s=10,c="1F4E79"); ce_.alignment=_align("left","top",True); ce_.border=_brd()
                ws_pa.merge_cells(start_row=r_,start_column=6,end_row=r_+2,end_column=9)
                cd_=ws_pa.cell(row=r_,column=6,value=_cd_val)
                cd_.fill=_fill("F5E6DF"); cd_.font=_font(s=10,c="1F4E79"); cd_.alignment=_align("left","top",True); cd_.border=_brd()
                for rr in range(r_,r_+3): ws_pa.row_dimensions[rr].height=22
                r_+=3

                _auto_width(ws_pa)
                buf_pa=BytesIO(); wb_pa.save(buf_pa); buf_pa.seek(0)
                fn_pa=f"Plan_Action_{NOM_PA[:20].replace(' ','_')}_{EXER_PA.replace('/','_')}.xlsx"
                st.success("✅ Plan d'action généré avec les données réelles de VISION ÉDITION !")
                st.download_button("⬇️ Télécharger le plan d'action (Excel)",
                    data=buf_pa,file_name=fn_pa,
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True,type="primary")
            except Exception as e:
                st.error(f"❌ Erreur : {e}")
                st.exception(e)


# ============================================================
# FOOTER
# ============================================================
st.divider()
st.markdown(
    "<div style='text-align:center;color:#888;font-size:12px'>"
    "© 2025 Nicolas CUISSET — VISION EDITION — Mémoire d'expertise comptable"
    "</div>",
    unsafe_allow_html=True
)
